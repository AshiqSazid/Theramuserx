import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import date, datetime
from pathlib import Path
import json
from typing import Dict, List, Optional, Tuple
import re
import io
import base64
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn
import toml
import sqlite3
import os

# Load color schema from config
def load_color_schema():
    """Load color schema from config.toml file"""
    try:
        config_path = Path(__file__).parent / ".streamlit" / "config.toml"
        if config_path.exists():
            config = toml.load(config_path)
            return config.get('color_schema', {})
    except Exception as e:
        st.warning(f"Could not load color schema: {e}")

    # Fallback colors if config fails to load
    return {
        'slider_primary': '#a8b894',
        'slider_secondary': '#e8ddb5',
        'slider_track': '#5B9C96',
        'table_header': '#e8ddb5',
        'table_cells': '#a8b894',
        'table_border': '#d4c4a8',
        'table_hover': '#A3B18A',
        'spider_primary': '#5B9C96',
        'spider_secondary': '#a8b894',
        'spider_fill': '#e8ddb580',
        'spider_text': '#000000',
        'spider_grid': '#A3B18A50',
        'accent_teal': '#5B9C96',
        'accent_olive': '#a8b894',
        'accent_beige': '#e8ddb5',
        'accent_dark': '#344E41'
    }

# Import from ml.py - Added robust error handling for Streamlit Cloud
import sys
import os
from pathlib import Path

# Add current directory to Python path if not already there
current_dir = Path(__file__).parent.absolute()
if str(current_dir) not in sys.path:
    sys.path.insert(0, str(current_dir))

# Debug: Print current path for troubleshooting
if "STREAMLIT_SERVER" in os.environ:
    print(f"Streamlit Cloud - Current directory: {current_dir}")
    print(f"Python path: {sys.path[:3]}")
    print(f"Files in directory: {list(current_dir.glob('*.py'))[:5]}")

try:
    import ml
    # Extract classes from ml module
    TheraMuse = ml.TheraMuse
    DementiaTherapy = ml.DementiaTherapy
    DownSyndromeTherapy = ml.DownSyndromeTherapy
    ADHDTherapy = ml.ADHDTherapy
    YouTubeAPI = ml.YouTubeAPI
    BangladeshiGenerationalMatrix = ml.BangladeshiGenerationalMatrix
    BigFivePersonalityMapping = ml.BigFivePersonalityMapping
    LinearThompsonSampling = ml.LinearThompsonSampling
    DatabaseManager = ml.DatabaseManager

    if "STREAMLIT_SERVER" in os.environ:
        print("✓ Successfully imported all classes from ml module")

except ImportError as e:
    error_msg = f"Import Error: Failed to import required modules from ml.py"
    if "STREAMLIT_SERVER" in os.environ:
        error_msg += f"\n\nDebug info:\n- Current directory: {current_dir}\n- Python path: {sys.path}\n- Error: {str(e)}"

    st.error(error_msg)
    st.error("Please ensure ml.py exists in the same directory as streamlit.py")
    st.stop()
except Exception as e:
    st.error(f"Unexpected error during import: {str(e)}")
    if "STREAMLIT_SERVER" in os.environ:
        import traceback
        st.error(f"Full traceback: {traceback.format_exc()}")
    st.stop()

# JavaScript to prevent Enter key from submitting forms in text input fields
st.markdown("""
<script>
document.addEventListener('DOMContentLoaded', function() {
    // Function to prevent Enter key submission in text inputs
    function preventEnterSubmission() {
        // Get all text input elements
        const textInputs = document.querySelectorAll('input[type="text"]');

        textInputs.forEach(function(input) {
            // Remove existing event listeners to avoid duplicates
            input.removeEventListener('keydown', preventEnterHandler);

            // Add new event listener
            input.addEventListener('keydown', preventEnterHandler);
        });
    }

    function preventEnterHandler(event) {
        // Check if the key pressed is Enter
        if (event.key === 'Enter' || event.keyCode === 13) {
            // Prevent the default form submission behavior
            event.preventDefault();
            event.stopPropagation();

            // Move focus to the next element or blur (lose focus)
            // This simulates typical "Tab" behavior when Enter is pressed
            const formElements = Array.from(document.querySelectorAll('input, select, button, textarea'));
            const currentIndex = formElements.indexOf(event.target);

            if (currentIndex !== -1 && currentIndex < formElements.length - 1) {
                // Try to focus on the next form element
                const nextElement = formElements[currentIndex + 1];
                if (nextElement && nextElement.type !== 'submit') {
                    nextElement.focus();
                } else {
                    // If next element is submit button, just blur current input
                    event.target.blur();
                }
            } else {
                // If no next element, just blur current input
                event.target.blur();
            }

            return false;
        }
    }

    // Initial setup
    preventEnterSubmission();

    // Set up a MutationObserver to handle dynamically added elements
    const observer = new MutationObserver(function(mutations) {
        mutations.forEach(function(mutation) {
            if (mutation.addedNodes.length) {
                preventEnterSubmission();
            }
        });
    });

    // Start observing the document for changes
    observer.observe(document.body, {
        childList: true,
        subtree: true
    });

    // Also run periodically as a fallback
    setInterval(preventEnterSubmission, 1000);
});
</script>
""", unsafe_allow_html=True)

# PAGE CONFIG (MUST BE FIRST STREAMLIT COMMAND)
LOGO_PATH = next((p for p in [Path("p.png"), Path.cwd() / "p.png"] if p.exists()), None)

st.set_page_config(
    page_title="TheraMuse - Music Therapy",
    page_icon=str(LOGO_PATH) if LOGO_PATH else "",
    layout="wide",
    initial_sidebar_state="expanded"
)
##########

st.markdown("""
<style>
/* Column headers */
[data-testid="stDataFrame"] [role="columnheader"] {
    background-color: #a8b894 ;   /* same as your background */
    color: #e8dfb0 ;              /* warm beige header text */
    font-weight: 600 ;
    border: none ;
}

/* Table cells */
[data-testid="stDataFrame"] [role="gridcell"] {
    background-color: #a8b894 ;   /* same olive background */
    color: #e8dfb0 ;              /* beige text for cells */
    border: none ;
}

/* Rounded corners and hover highlight */
[data-testid="stDataFrame"] {
    border-radius: 12px ;
    overflow: hidden ;
}
[data-testid="stDataFrame"] [role="row"]:hover [role="gridcell"] {
    background-color: #b7c3a2 ;   /* gentle hover highlight */
}
</style>

<script>
// Prevent Enter key from submitting forms in text inputs
document.addEventListener('DOMContentLoaded', function() {
    const textInputs = document.querySelectorAll('input[type="text"]');
    textInputs.forEach(function(input) {
        input.addEventListener('keydown', function(event) {
            if (event.key === 'Enter') {
                event.preventDefault();
                event.stopPropagation();
                // Just move focus out of the input field
                this.blur();
                return false;
            }
        });
    });
});
</script>
""", unsafe_allow_html=True)



#STYLING
st.markdown("""
<style>
    /* Global Premium Styling */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    .main {
        background: linear-gradient(135deg, #0a0a0a 0%, #1a1a1a 100%);
        animation: fadeIn 0.8s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    /* Premium Header */
    .main-header {
        font-size: 3.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #A3B18A 0%, #E8DDB5 50%, #A3B18A 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-align: center;
        margin: 2rem 0;
        letter-spacing: -0.02em;
        animation: shimmer 3s ease-in-out infinite;
    }
    
    @keyframes shimmer {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    /* Glassmorphism Card */
    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px) saturate(180%);
        -webkit-backdrop-filter: blur(20px) saturate(180%);
        border-radius: 20px;
        border: 1px solid rgba(255, 255, 255, 0.08);
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 
            0 8px 32px 0 rgba(0, 0, 0, 0.37),
            inset 0 1px 0 0 rgba(255, 255, 255, 0.05);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: slideUp 0.6s ease-out;
    }
    
    @keyframes slideUp {
        from { opacity: 0; transform: translateY(30px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .glass-card:hover {
        transform: translateY(-8px);
        box-shadow: 
            0 16px 48px 0 rgba(163, 177, 138, 0.2),
            inset 0 1px 0 0 rgba(255, 255, 255, 0.1);
        border-color: rgba(163, 177, 138, 0.3);
    }
    
    /* Premium Song Card */
    .song-card {
        background: linear-gradient(135deg, rgba(26, 26, 26, 0.95) 0%, rgba(20, 20, 20, 0.95) 100%);
        backdrop-filter: blur(20px);
        padding: 1.8rem;
        border-radius: 16px;
        border-left: 4px solid transparent;
        border-image: linear-gradient(180deg, #A3B18A 0%, #588157 100%);
        border-image-slice: 1;
        margin: 1rem 0;
        position: relative;
        overflow: hidden;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: fadeInCard 0.5s ease-out backwards;
    }
    
    @keyframes fadeInCard {
        from { 
            opacity: 0; 
            transform: translateX(-20px);
        }
        to { 
            opacity: 1; 
            transform: translateX(0);
        }
    }
    
    .song-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.05) 0%, transparent 100%);
        opacity: 0;
        transition: opacity 0.4s ease;
        pointer-events: none; /* allow clicks on links/buttons above overlay */
    }
    
    .song-card:hover {
        transform: translateX(8px) scale(1.02);
        border-left-width: 6px;
        box-shadow: 
            -4px 0 24px rgba(163, 177, 138, 0.3),
            0 8px 32px rgba(0, 0, 0, 0.4);
    }
    
    .song-card:hover::before {
        opacity: 1;
    }
    
    .song-card h4 {
        color: #FFFFFF;
        font-size: 1.25rem;
        font-weight: 600;
        margin-bottom: 0.75rem;
        letter-spacing: -0.01em;
        transition: color 0.3s ease;
    }
    
    .song-card:hover h4 {
        color: #E8DDB5;
    }
    
    .song-card p {
        color: rgba(255, 255, 255, 0.7);
        font-size: 0.95rem;
        line-height: 1.6;
        margin: 0.5rem 0;
    }
    
    .song-card strong {
        color: #A3B18A;
        font-weight: 600;
    }
    .song-card a { pointer-events: auto; }
    
    /* YouTube Link Styling */
    .youtube-link {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.75rem 1.5rem;
        background: linear-gradient(135deg, #FF0000 0%, #CC0000 100%);
        color: white;
        text-decoration: none;
        border-radius: 12px;
        font-weight: 600;
        font-size: 0.95rem;
        margin-top: 1rem;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 16px rgba(255, 0, 0, 0.3);
        position: relative;
        z-index: 3; /* ensure above decorative overlays */
        cursor: pointer;
    }
    
    .youtube-link:hover {
        transform: translateY(-2px) scale(1.05);
        box-shadow: 0 8px 24px rgba(255, 0, 0, 0.5);
        background: linear-gradient(135deg, #FF1a1a 0%, #DD0000 100%);
        text-decoration: none;
        color: white;
    }
    
    .youtube-link::before {
        content: '▶';
        font-size: 1.1rem;
    }
    
    /* Embedded YouTube Player Container */
    .youtube-embed-container {
        position: relative;
        width: 100%;
        padding-top: 56.25%; /* 16:9 Aspect Ratio */
        margin-top: 1rem;
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.5);
        transition: all 0.3s ease;
    }
    
    .youtube-embed-container:hover {
        transform: scale(1.02);
        box-shadow: 0 12px 48px rgba(163, 177, 138, 0.3);
    }
    
    .youtube-embed-container iframe {
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        border: none;
        border-radius: 12px;
    }
    
    /* Sidebar Navigation Colors */
    [data-testid="stSidebar"] [data-testid="stRadio"] label p {
        color: #FDFBF7 ;
    }
    [data-testid="stSidebar"] [data-testid="stRadio"] label {
        color: #FDFBF7 ;
    }
    
    /* Metric Cards */
    .metric-card {
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.15) 0%, rgba(88, 129, 87, 0.15) 100%);
        backdrop-filter: blur(20px);
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        border: 1px solid rgba(163, 177, 138, 0.2);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: scaleIn 0.5s ease-out backwards;
    }
    
    @keyframes scaleIn {
        from { 
            opacity: 0; 
            transform: scale(0.9);
        }
        to { 
            opacity: 1; 
            transform: scale(1);
        }
    }
    
    .metric-card:hover {
        transform: translateY(-8px) scale(1.05);
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.25) 0%, rgba(88, 129, 87, 0.25) 100%);
        box-shadow: 0 16px 48px rgba(163, 177, 138, 0.3);
        border-color: rgba(163, 177, 138, 0.4);
    }
    
    .metric-card h2, .metric-card h3 {
        color: #E8DDB5;
        margin: 0.5rem 0;
    }
    
    /* Premium Buttons - Enhanced for better functionality */
    .stButton>button {
        background: linear-gradient(135deg, #338AFF 0%, #1E6FFF 100%) ;
        color: white ;
        border: none ;
        border-radius: 12px ;
        padding: 0.75rem 1.5rem ;
        font-weight: 600 ;
        font-size: 0.95rem ;
        letter-spacing: 0.02em ;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) ;
        box-shadow: 0 4px 16px rgba(51, 138, 255, 0.3) ;
        cursor: pointer ;
        pointer-events: auto ;
        position: relative ;
        z-index: 10 ;
    }

    .stButton>button:hover {
        background: linear-gradient(135deg, #1E6FFF 0%, #0D5BFF 100%) ;
        transform: translateY(-2px) ;
        box-shadow: 0 8px 24px rgba(51, 138, 255, 0.5) ;
    }

    .stButton>button:active {
        transform: translateY(0) ;
        box-shadow: 0 4px 16px rgba(51, 138, 255, 0.3) ;
    }

    /* Ensure buttons are clickable and not blocked */
    .stButton {
        pointer-events: auto ;
        z-index: 10 ;
        position: relative ;
    }

    /* Fix for buttons in containers */
    div.stButton > button {
        pointer-events: auto ;
        z-index: 10 ;
        position: relative ;
    }

    
    /* Feedback Buttons */
    .feedback-button {
        width: 100%;
        margin: 0.3rem 0;
    }
    
    /* Expander Premium Style */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, rgba(51, 138, 255, 0.1) 0%, rgba(13, 91, 255, 0.1) 100%);
        border-radius: 12px;
        border: 1px solid rgba(51, 138, 255, 0.2);
        transition: all 0.3s ease;
    }

    .streamlit-expanderHeader:hover {
        background: linear-gradient(135deg, rgba(51, 138, 255, 0.2) 0%, rgba(13, 91, 255, 0.2) 100%);
        border-color: rgba(51, 138, 255, 0.4);
    }
    
    /* Category Badge */
    .category-badge {
        display: inline-block;
        padding: 0.5rem 1rem;
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%);
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        color: #E8DDB5;
        margin: 0.5rem 0;
        border: 1px solid rgba(163, 177, 138, 0.3);
    }
    
    /* Rank Badge */
    .rank-badge {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 32px;
        height: 32px;
        background: linear-gradient(135deg, #A3B18A 0%, #588157 100%);
        border-radius: 50%;
        font-weight: 700;
        font-size: 0.9rem;
        color: white;
        margin-right: 0.75rem;
        box-shadow: 0 2px 8px rgba(163, 177, 138, 0.4);
    }
    
    /* Smooth Scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.03);
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, #A3B18A 0%, #588157 100%);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(180deg, #588157 0%, #3a5a40 100%);
    }
    
    /* Tab Styling - Updated to #338AFF theme */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: rgba(51, 138, 255, 0.1);
        border-radius: 12px;
        border: 1px solid rgba(51, 138, 255, 0.2);
        padding: 4px;
    }

    .stTabs [data-baseweb="tab"] {
        background: linear-gradient(135deg, #338AFF 0%, #1E6FFF 100%);
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        color: white;
        transition: all 0.3s ease;
        border: none;
        font-weight: 600;
        box-shadow: 0 2px 8px rgba(51, 138, 255, 0.2);
    }

    .stTabs [aria-selected="true"] [data-baseweb="tab"] {
        background: linear-gradient(135deg, #1E6FFF 0%, #0D5BFF 100%);
        box-shadow: 0 4px 12px rgba(51, 138, 255, 0.4);
        transform: translateY(-1px);
    }

    .stTabs [data-baseweb="tab"]:hover {
        background: linear-gradient(135deg, #1E6FFF 0%, #0D5BFF 100%);
        color: white;
        box-shadow: 0 6px 16px rgba(51, 138, 255, 0.5);
        transform: translateY(-1px);
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, rgba(51, 138, 255, 0.2) 0%, rgba(13, 91, 255, 0.2) 100%);
        color: white;
        border-color: rgba(51, 138, 255, 0.4);
    }
    
    /* Form Elements */
    .stTextInput>div>div>input,
    .stSelectbox>div>div>select {
        background: rgba(255, 255, 255, 0.05) ;
        border: 1px solid rgba(255, 255, 255, 0.1) ;
        border-radius: 12px ;
        color: white ;
        transition: all 0.3s ease ;
    }
    
    .stTextInput>div>div>input:focus,
    .stSelectbox>div>div>select:focus {
        border-color: rgba(163, 177, 138, 0.5) ;
        box-shadow: 0 0 0 2px rgba(163, 177, 138, 0.2) ;
    }

    /* Success/Error Messages */
    .stSuccess, .stError, .stInfo, .stWarning {
        border-radius: 12px;
        backdrop-filter: blur(20px);
        animation: slideInRight 0.4s ease-out;
    }
    
    @keyframes slideInRight {
        from { 
            opacity: 0; 
            transform: translateX(20px);
        }
        to { 
            opacity: 1; 
            transform: translateX(0);
        }
    }
    
    /* Loading Spinner */
    .stSpinner > div {
        border-color: #A3B18A transparent transparent transparent ;
    }
    
    /* Premium Section Headers */
    h1, h2, h3 {
        color: #E8DDB5 ;
        font-weight: 700 ;
        letter-spacing: -0.02em ;
    }
    
    /* Sidebar Styling */
    [data-testid="stSidebar"] {
        background: #338aff !important;
        backdrop-filter: blur(20px);
    }
    
    [data-testid="stSidebar"] .stRadio > label {
        background: rgba(255, 255, 255, 0.1);
        border-radius: 12px;
        padding: 0.75rem;
        margin: 0.3rem 0;
        transition: all 0.3s ease;
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    
    [data-testid="stSidebar"] .stRadio > label:hover {
        background: rgba(255, 255, 255, 0.2);
        border-color: rgba(255, 255, 255, 0.4);
    }
</style>
""", unsafe_allow_html=True)

# HELPER FUNCTIONS

def render_logo(size=150):
    """Render logo with premium animation"""
    if LOGO_PATH and LOGO_PATH.exists():
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image(str(LOGO_PATH), width=size)
    else:
         st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&display=swap');

.typing-header {
    text-align: center;
    font-family: 'Playfair Display', serif;
    font-size: 3.5rem;
    color: #E8DDB5;
    letter-spacing: 2px;
    white-space: nowrap;
    overflow: hidden;
    border-right: 3px solid #E8DDB5;
    width: 0;
    margin: 2rem auto 1rem auto;
    animation: typing 3s steps(20, end) forwards, blink 0.7s step-end infinite;
}

/* typing animation */
@keyframes typing {
    from { width: 0; }
    to { width: 16ch; } /* length of "🎵 TheraMuse" */
}

/* blinking cursor */
@keyframes blink {
    50% { border-color: transparent; }
}

/* optional subtitle */
.sub-header {
    text-align: center;
    font-family: 'Poppins', sans-serif;
    font-size: 1rem;
    letter-spacing: 1px;
    color: rgba(232, 221, 181, 0.8);
    opacity: 0;
    animation: fadeIn 1s ease-in-out 3s forwards;
}

@keyframes fadeIn {
    from { opacity: 0; }
    to { opacity: 1; }
}
</style>


""", unsafe_allow_html=True)



def compute_age_from_dob(dob: date) -> int:
    """Calculate age from date of birth"""
    today = date.today()
    return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))

def reverse_1to7(score: float) -> float:
    """Reverse score for negatively keyed items"""
    return 8.0 - score

def slider_with_ticks(label: str, key: str) -> float:
    """Render a standard Streamlit slider."""
    return st.slider(label, min_value=1, max_value=7, value=4, key=key)

def get_condition_code(condition: str) -> str:
    """Map condition names to codes used in main.py"""
    mapping = {
        "Dementia / Alzheimer's": "dementia",
        "Down Syndrome": "down_syndrome",
        "ADHD": "adhd"
    }
    return mapping.get(condition, "dementia")

# HELPER FUNCTIONS
def get_database_path():
    """Get the correct database path for different environments"""
    import os

    # For Streamlit Cloud, use a temporary directory or relative path
    if "STREAMLIT_SERVER" in os.environ or not os.path.exists("/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app"):
        # Use the current working directory for Streamlit Cloud
        return Path("theramuse.db")
    else:
        # Local development path
        return Path("/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app/theramuse.db")

# PATIENT DATABASE FUNCTIONS
def get_patient_db_connection():
    """Get connection to patient database"""
    import tempfile
    import os

    # Get the correct database path
    db_path = get_database_path()

    # Ensure the directory exists
    db_path.parent.mkdir(parents=True, exist_ok=True)

    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()

    # Create patients table if it doesn't exist
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patients (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            age INTEGER,
            sex TEXT,
            birthplace_city TEXT,
            birthplace_country TEXT,
            favorite_genre TEXT,
            favorite_musician TEXT,
            favorite_season TEXT,
            instruments TEXT,
            natural_elements TEXT,
            condition TEXT,
            difficulty_sleeping BOOLEAN,
            trouble_remembering BOOLEAN,
            forgets_everyday_things BOOLEAN,
            difficulty_recalling_old_memories BOOLEAN,
            memory_worse_than_year_ago BOOLEAN,
            visited_mental_health_professional BOOLEAN,
            extraversion REAL,
            agreeableness REAL,
            conscientiousness REAL,
            neuroticism REAL,
            openness REAL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Create sessions table if it doesn't exist
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS therapy_sessions (
            id TEXT PRIMARY KEY,
            patient_id TEXT,
            session_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            recommendations_count INTEGER,
            session_data TEXT,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )
    ''')

    # Check if table exists and has correct structure
    cursor.execute("PRAGMA table_info(patients)")
    columns = cursor.fetchall()

    if not columns:
        # Table doesn't exist, commit the creation
        conn.commit()

    return conn

def save_patient_to_database(patient_info: Dict, big5_scores: Dict, recommendations: Dict, session_id: str):
    """Save patient information to database"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    patient_id = f"patient_{datetime.now().strftime('%Y%m%d%H%M%S')}"

    # Insert patient data
    cursor.execute('''
        INSERT OR REPLACE INTO patients (
            id, name, age, sex, birthplace_city, birthplace_country,
            favorite_genre, favorite_musician, favorite_season,
            instruments, natural_elements, condition,
            difficulty_sleeping, trouble_remembering, forgets_everyday_things,
            difficulty_recalling_old_memories, memory_worse_than_year_ago,
            visited_mental_health_professional, extraversion, agreeableness,
            conscientiousness, neuroticism, openness, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        patient_id,
        patient_info.get('name', 'Anonymous'),
        patient_info.get('age', 0),
        patient_info.get('sex', ''),
        patient_info.get('birthplace_city', ''),
        patient_info.get('birthplace_country', ''),
        patient_info.get('favorite_genre', ''),
        patient_info.get('favorite_musician', ''),
        patient_info.get('favorite_season', ''),
        json.dumps(patient_info.get('instruments', [])),
        json.dumps(patient_info.get('natural_elements', [])),
        patient_info.get('condition', ''),
        patient_info.get('difficulty_sleeping', False),
        patient_info.get('trouble_remembering', False),
        patient_info.get('forgets_everyday_things', False),
        patient_info.get('difficulty_recalling_old_memories', False),
        patient_info.get('memory_worse_than_year_ago', False),
        patient_info.get('visited_mental_health_professional', False),
        big5_scores.get('extraversion', 0),
        big5_scores.get('agreeableness', 0),
        big5_scores.get('conscientiousness', 0),
        big5_scores.get('neuroticism', 0),
        big5_scores.get('openness', 0),
        datetime.now()
    ))

    # Insert session data
    cursor.execute('''
        INSERT INTO therapy_sessions (
            id, patient_id, recommendations_count, session_data
        ) VALUES (?, ?, ?, ?)
    ''', (
        session_id,
        patient_id,
        recommendations.get('total_songs', 0),
        json.dumps(recommendations)
    ))

    conn.commit()
    conn.close()
    return patient_id

def get_all_patients():
    """Get all patients from database with Big Five scores and reinforcement learning"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # First try the enhanced schema with id column
        cursor.execute('''
            SELECT p.id, p.name, p.age, p.condition, p.created_at,
                   p.extraversion, p.agreeableness, p.conscientiousness,
                   p.neuroticism, p.openness, p.reinforcement_learning
            FROM patients p
            ORDER BY p.created_at DESC
        ''')
        patients = cursor.fetchall()
        conn.close()
        return patients
    except sqlite3.OperationalError as e:
        # If the enhanced schema doesn't exist, try the basic schema
        if "no such column: p.id" in str(e):
            try:
                cursor.execute('''
                    SELECT p.patient_id, p.name, p.age, p.condition, p.created_at,
                           0 as extraversion, 0 as agreeableness, 0 as conscientiousness,
                           0 as neuroticism, 0 as openness, 0 as reinforcement_learning
                    FROM patients p
                    ORDER BY p.created_at DESC
                ''')
                patients = cursor.fetchall()
                conn.close()
                return patients
            except sqlite3.Error as e2:
                st.error(f"Database query error (basic schema): {str(e2)}")
                conn.close()
                return []
        else:
            st.error(f"Database query error: {str(e)}")
            conn.close()
            return []
    except sqlite3.Error as e:
        st.error(f"Database query error: {str(e)}")
        conn.close()
        return []

def get_patient_details(patient_id: str):
    """Get detailed information for a specific patient with Big Five scores and reinforcement learning"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # Try to get patient with enhanced schema first
        try:
            cursor.execute('SELECT * FROM patients WHERE id = ?', (patient_id,))
            patient = cursor.fetchone()
        except sqlite3.OperationalError:
            # Fall back to basic schema
            cursor.execute('SELECT * FROM patients WHERE patient_id = ?', (patient_id,))
            patient = cursor.fetchone()

        # Get Big Five scores with reinforcement learning
        cursor.execute('''
            SELECT * FROM big5_scores
            WHERE patient_id = ?
            ORDER BY created_at DESC
            LIMIT 1
        ''', (patient_id,))
        big5_scores = cursor.fetchone()

        cursor.execute('SELECT * FROM therapy_sessions WHERE patient_id = ? ORDER BY session_date DESC', (patient_id,))
        sessions = cursor.fetchall()

        # Combine patient data with Big Five scores if available
        if patient and big5_scores:
            patient = list(patient) + list(big5_scores[2:])  # Skip patient_id and session_id from big5_scores

        conn.close()
        return patient, sessions
    except sqlite3.Error as e:
        st.error(f"Database query error: {str(e)}")
        conn.close()
        return None, []

def delete_patient(patient_id: str):
    """Delete a patient from database"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    cursor.execute('DELETE FROM therapy_sessions WHERE patient_id = ?', (patient_id,))

    # Try both possible column names
    try:
        cursor.execute('DELETE FROM patients WHERE id = ?', (patient_id,))
    except sqlite3.OperationalError:
        cursor.execute('DELETE FROM patients WHERE patient_id = ?', (patient_id,))

    conn.commit()
    conn.close()

def get_comprehensive_patient_data():
    """Get comprehensive patient data with therapy recommendations"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # Get basic patient info with backward compatibility
        try:
            cursor.execute('''
                SELECT id, name, age, condition, created_at
                FROM patients ORDER BY created_at DESC
            ''')
            patients = cursor.fetchall()
            id_col = "id"
        except sqlite3.OperationalError:
            cursor.execute('''
                SELECT patient_id, name, age, condition, created_at
                FROM patients ORDER BY created_at DESC
            ''')
            patients = cursor.fetchall()
            id_col = "patient_id"

        patient_data = []

        for patient in patients:
            patient_id = patient[0]

            # Get therapy sessions and recommendations
            cursor.execute('''
                SELECT id, session_date, recommendations_count
                FROM therapy_sessions WHERE patient_id = ?
                ORDER BY session_date DESC
            ''', (patient_id,))
            sessions = cursor.fetchall()

            # Get therapy recommendations with song details
            cursor.execute('''
                SELECT category, song_title, video_id, channel, rank
                FROM therapy_recommendations WHERE patient_id = ?
                ORDER BY category, rank
            ''', (patient_id,))
            recommendations = cursor.fetchall()

            # Get feedback data
            cursor.execute('''
                SELECT feedback_type, reward, created_at
                FROM therapy_feedback WHERE patient_id = ?
                ORDER BY created_at DESC
            ''', (patient_id,))
            feedback = cursor.fetchall()

            # Get Big Five scores if available
            cursor.execute('''
                SELECT openness, conscientiousness, extraversion,
                       agreeableness, neuroticism, reinforcement_learning
                FROM big5_scores WHERE patient_id = ?
                ORDER BY created_at DESC LIMIT 1
            ''', (patient_id,))
            big5_scores = cursor.fetchone()

            patient_data.append({
                'patient_info': patient,
                'sessions': sessions,
                'recommendations': recommendations,
                'feedback': feedback,
                'big5_scores': big5_scores or (0, 0, 0, 0, 0, 0)
            })

        conn.close()
        return patient_data

    except Exception as e:
        st.error(f"Database error: {str(e)}")
        conn.close()
        return []

def page_patient_database():
    """Advanced Patient Database Management Page"""
    render_logo()

    st.markdown("""
    <div class='glass-card'>
        <h2> Advanced Patient Database</h2>
        <p style='color: rgba(255,255,255,0.7);'>Comprehensive patient records with therapy analytics and recommendations</p>
    </div>
    """, unsafe_allow_html=True)

    # Database connection status
    try:
        conn = get_patient_db_connection()
        cursor = conn.cursor()

        # Get database statistics
        cursor.execute("SELECT COUNT(*) FROM patients")
        total_patients = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_sessions")
        total_sessions = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_recommendations")
        total_recommendations = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_feedback")
        total_feedback = cursor.fetchone()[0]

        conn.close()

        # Display statistics
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown(f"""
            <div class='metric-card'>
                <h3 style='color:#364153;'>Total Patients</h3>
                <h2 style='color:#364153;'>{total_patients}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown(f"""
            <div class='metric-card'>
                <h3 style='color:#364153;'>Therapy Sessions</h3>
                <h2 style='color:#364153;'>{total_sessions}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col3:
            st.markdown(f"""
            <div class='metric-card'>
                <h3 style='color:#364153;'>Song Recommendations</h3>
                <h2 style='color:#364153;'>{total_recommendations}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col4:
            st.markdown(f"""
            <div class='metric-card'>
                <h3 style='color:#364153;'>Feedback Recorded</h3>
                <h2 style='color:#364153;'>{total_feedback}</h2>
            </div>
            """, unsafe_allow_html=True)

    except Exception as e:
        st.error(f"Database connection error: {str(e)}")
        return

    # Advanced filters and search
    st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Advanced Search & Filters</h3>',
    unsafe_allow_html=True
)

    col1, col2, col3 = st.columns(3)

    with col1:
        search_term = st.text_input(" Search patients by name...", placeholder="Enter patient name...")

    with col2:
        condition_filter = st.selectbox(" Filter by condition",
                                       ["All Conditions", "Dementia / Alzheimer's", "Down Syndrome", "ADHD"])

    with col3:
        sort_by = st.selectbox(" Sort by",
                              ["Latest First", "Oldest First", "Most Sessions", "Most Feedback"])

    # Get comprehensive patient data
    patient_data = get_comprehensive_patient_data()

    if not patient_data:
        st.info("No patient records found in the database.")
        return

    # Apply filters
    filtered_data = patient_data.copy()

    if search_term:
        filtered_data = [p for p in filtered_data
                        if search_term.lower() in p['patient_info'][1].lower()]

    if condition_filter != "All Conditions":
        condition_map = {
            "Dementia / Alzheimer's": "dementia",
            "Down Syndrome": "down_syndrome",
            "ADHD": "adhd"
        }
        filtered_data = [p for p in filtered_data
                        if condition_map.get(condition_filter, "") in str(p['patient_info'][3]).lower()]

    # Apply sorting
    if sort_by == "Latest First":
        filtered_data.sort(key=lambda x: x['patient_info'][4] if len(x['patient_info']) > 4 else "", reverse=True)
    elif sort_by == "Oldest First":
        filtered_data.sort(key=lambda x: x['patient_info'][4] if len(x['patient_info']) > 4 else "")
    elif sort_by == "Most Sessions":
        filtered_data.sort(key=lambda x: len(x['sessions']), reverse=True)
    elif sort_by == "Most Feedback":
        filtered_data.sort(key=lambda x: len(x['feedback']), reverse=True)

    if not filtered_data:
        st.warning("No patients match the selected filters.")
        return

    st.success(f"Found {len(filtered_data)} patient(s) matching your criteria")

    # Enhanced patient display
    for idx, patient in enumerate(filtered_data):
        patient_info = patient['patient_info']
        sessions = patient['sessions']
        recommendations = patient['recommendations']
        feedback = patient['feedback']
        big5_scores = patient['big5_scores']

        with st.expander(f"👤 {patient_info[1]} ({patient_info[2]} years) - {patient_info[3].title()}", expanded=False):

            # Patient Overview
            col1, col2, col3 = st.columns([2, 1, 1])

            with col1:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Patient Information</h3>',
    unsafe_allow_html=True
)
                st.write(f"**ID:** {patient_info[0]}")
                st.write(f"**Age:** {patient_info[2]}")
                st.write(f"**Condition:** {patient_info[3].title()}")
                if len(patient_info) > 4:
                    st.write(f"**Registered:** {patient_info[4]}")

            with col2:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Personality Profile</h3>',
    unsafe_allow_html=True
)
                if big5_scores:
                    for trait, score in zip(['Openness', 'Conscientiousness', 'Extraversion', 'Agreeableness', 'Neuroticism'], big5_scores[:5]):
                        if score > 0:
                            st.metric(trait, f"{score:.1f}/7.0")
                    st.metric("RL Interactions", big5_scores[5] if len(big5_scores) > 5 else 0)
                else:
                    st.info("No personality data")

            with col3:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Activity Summary</h3>',
    unsafe_allow_html=True
)
                st.metric("Sessions", len(sessions))
                st.metric("Songs", len(recommendations))
                st.metric("Feedback", len(feedback))

            # Therapy Sessions
            if sessions:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Therapy Sessions</h3>',
    unsafe_allow_html=True
)
                session_df = pd.DataFrame(sessions,
                                         columns=['Session ID', 'Date', 'Songs Recommended'])
                if not session_df.empty:
                    session_df['Date'] = pd.to_datetime(session_df['Date']).dt.strftime('%Y-%m-%d %H:%M')
                    st.dataframe(session_df, hide_index=True, use_container_width=True)

            # Song Recommendations
            if recommendations:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">🎵 Song Recommendations</h3>',
    unsafe_allow_html=True
)

                # Group recommendations by category
                categories = {}
                for rec in recommendations:
                    category = rec[0]
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(rec)

                for category, songs in categories.items():
                    with st.expander(f"🎼 {category.replace('_', ' ').title()} ({len(songs)} songs)", expanded=False):
                        for song in songs[:5]:  # Show first 5 songs per category
                            col_song, col_video = st.columns([3, 1])
                            with col_song:
                                st.write(f"**{song[1]}**")
                                if song[2]:
                                    st.caption(f"Channel: {song[3] if len(song) > 3 else 'N/A'}")
                            with col_video:
                                if song[2]:
                                    st.markdown(f"[ Watch](https://www.youtube.com/watch?v={song[2]})")

                        if len(songs) > 5:
                            st.caption(f"... and {len(songs) - 5} more songs in this category")

            # Feedback Analysis
            if feedback:
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Feedback Analysis</h3>',
    unsafe_allow_html=True
)

                feedback_counts = {}
                for fb in feedback:
                    fb_type = fb[0]
                    feedback_counts[fb_type] = feedback_counts.get(fb_type, 0) + 1

                col1, col2, col3 = st.columns(3)
                with col1:
                    if 'like' in feedback_counts:
                        st.metric(" Likes", feedback_counts['like'])
                with col2:
                    if 'dislike' in feedback_counts:
                        st.metric(" Dislikes", feedback_counts['dislike'])
                with col3:
                    if 'skip' in feedback_counts:
                        st.metric("⏭ Skips", feedback_counts['skip'])

                # Recent feedback
                st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Recent Feedback:</h3>',
    unsafe_allow_html=True
)
                for fb in feedback[:3]:
                    emoji = {"like": "", "dislike": "", "skip": "", "neutral": ""}.get(fb[0], "")
                    st.write(f"{emoji} {fb[0].title()} - {fb[2] if len(fb) > 2 else 'N/A'}")

            # Action buttons
            st.markdown("---")
            col1, col2, col3 = st.columns(3)

            with col1:
                # Create a unique key that handles None patient IDs
                unique_key = f"export_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
                if st.button(f" Export Data", key=unique_key, use_container_width=True):
                    export_data = {
                        'patient_info': patient_info,
                        'sessions': sessions,
                        'recommendations': recommendations,
                        'feedback': feedback,
                        'big5_scores': big5_scores
                    }
                    st.download_button(
                        label=" Download JSON",
                        data=json.dumps(export_data, indent=2, default=str),
                        file_name=f"patient_{patient_info[1]}_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json",
                        use_container_width=True
                    )

            with col2:
                refresh_key = f"refresh_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
                if st.button(f" Refresh Data", key=refresh_key, use_container_width=True):
                    st.rerun()

            with col3:
                delete_key = f"delete_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
                if st.button(f" Delete Patient", key=delete_key, type="secondary", use_container_width=True):
                    st.session_state[f"confirm_delete_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"] = True

            # Confirmation dialog
            confirm_key = f"confirm_delete_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
            if st.session_state.get(confirm_key, False):
                st.warning(" **Confirm Deletion**: This will permanently delete all patient data including sessions and recommendations.")
                col1, col2 = st.columns(2)
                with col1:
                    yes_key = f"confirm_yes_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
                    if st.button(" Yes, Delete", key=yes_key, type="primary"):
                        delete_patient(patient_info[0])
                        st.success("Patient deleted successfully!")
                        st.session_state[confirm_key] = False
                        st.rerun()
                with col2:
                    no_key = f"confirm_no_{patient_info[0] if patient_info[0] is not None else 'unknown'}_{hash(str(patient_info))}"
                    if st.button(" Cancel", key=no_key):
                        st.session_state[confirm_key] = False
                        st.rerun()

def create_docx_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> bytes:
    """Create a DOCX file with patient information and recommendations"""
    doc = Document()

    # Add title
    title = doc.add_heading('TheraMuse Therapy Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add patient information
    doc.add_heading('Patient Information', level=1)
    patient_table = doc.add_table(rows=0, cols=2)
    patient_table.style = 'Table Grid'

    patient_data = [
        ('Name', patient_info.get('name', 'N/A')),
        ('Age', str(patient_info.get('age', 'N/A'))),
        ('Sex', patient_info.get('sex', 'N/A')),
        ('Birthplace City', patient_info.get('birthplace_city', 'N/A')),
        ('Birthplace Country', patient_info.get('birthplace_country', 'N/A')),
        ('Condition', str(patient_info.get('condition', 'N/A')).title()),
        ('Favorite Genres', patient_info.get('favorite_genre', 'N/A')),
        ('Favorite Musician', patient_info.get('favorite_musician', 'N/A')),
        ('Favorite Season', patient_info.get('favorite_season', 'N/A')),
        ('Preferred Instruments', ', '.join(patient_info.get('instruments', []))),
        ('Natural Elements', ', '.join(patient_info.get('natural_elements', [])))
    ]

    for key, value in patient_data:
        row_cells = patient_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    # Add Big-5 personality scores
    doc.add_heading('Personality Profile (Big-5)', level=1)
    personality_table = doc.add_table(rows=0, cols=2)
    personality_table.style = 'Table Grid'

    personality_labels = {
        'extraversion': 'Extraversion',
        'agreeableness': 'Agreeableness',
        'conscientiousness': 'Conscientiousness',
        'neuroticism': 'Neuroticism',
        'openness': 'Openness'
    }

    for key, label in personality_labels.items():
        score = big5_scores.get(key, 0)
        row_cells = personality_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{score:.2f}/7.00"

    # Add recommendations
    doc.add_heading('Music Recommendations', level=1)

    if 'categories' in recommendations:
        category_labels = {
            "birthplace_country": "From Your Country",
            "birthplace_city": "From Your City",
            "instruments": "Instrumental Favorites",
            "seasonal": "Seasonal Music",
            "natural_elements": "Nature-Inspired",
            "favorite_genre": "Favorite Genres",
            "favorite_musician": "Favorite Musician",
            "therapeutic": "Therapeutic Selections",
            "personality_based": "Personality Match",
            "calming_sensory": "Calming Sensory",
            "concentration": "Focus & Concentration",
            "binaural_beats": "Binaural Beats",
            "relief_study": "Study & Relief",
            "additional_calm": "Additional Calming",
            "additional_focus": "Additional Focus"
        }

        for category, data in recommendations['categories'].items():
            label = category_labels.get(category, category.replace('_', ' ').title())
            songs = data.get('songs', [])

            if songs:
                doc.add_heading(label, level=2)
                for idx, song in enumerate(songs, 1):
                    title = song.get('title', 'Unknown Title')
                    channel = song.get('channel', 'Unknown Channel')
                    url = song.get('url', 'No URL available')

                    # Add song paragraph with title and channel
                    song_para = doc.add_paragraph()
                    song_para.add_run(f"{idx}. {title} - {channel}").bold = True

                    # Add YouTube URL
                    if url and url.startswith('https://www.youtube.com/'):
                        url_para = doc.add_paragraph()
                        url_para.add_run("YouTube Link: ").bold = True
                        url_para.add_run(f"{url}")
                    else:
                        url_para = doc.add_paragraph()
                        url_para.add_run("URL: ").bold = True
                        url_para.add_run(url if url else "No URL available")

                    doc.add_paragraph()  # Add spacing between songs

    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Generated on: ').bold = True
    footer.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Save document to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def create_pdf_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> str:
    """Create a simple HTML-based PDF content"""
    html_content = f"""
    <html>
    <head>
        <title>TheraMuse Therapy Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ color: #A3B18A; text-align: center; }}
            h2 {{ color: #588157; border-bottom: 2px solid #A3B18A; }}
            h3 {{ color: #E8DDB5; }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            .song {{ margin: 5px 0; padding: 5px; background-color: #f9f9f9; }}
        </style>
    </head>
    <body>
        <h1>TheraMuse Therapy Report</h1>
        <p><strong>Generated on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

        <h2>Patient Information</h2>
        <table>
            <tr><th>Field</th><th>Value</th></tr>
            <tr><td>Name</td><td>{patient_info.get('name', 'N/A')}</td></tr>
            <tr><td>Age</td><td>{patient_info.get('age', 'N/A')}</td></tr>
            <tr><td>Sex</td><td>{patient_info.get('sex', 'N/A')}</td></tr>
            <tr><td>Birthplace City</td><td>{patient_info.get('birthplace_city', 'N/A')}</td></tr>
            <tr><td>Birthplace Country</td><td>{patient_info.get('birthplace_country', 'N/A')}</td></tr>
            <tr><td>Condition</td><td>{patient_info.get('condition', 'N/A').title()}</td></tr>
            <tr><td>Favorite Genres</td><td>{patient_info.get('favorite_genre', 'N/A')}</td></tr>
            <tr><td>Favorite Musician</td><td>{patient_info.get('favorite_musician', 'N/A')}</td></tr>
            <tr><td>Favorite Season</td><td>{patient_info.get('favorite_season', 'N/A')}</td></tr>
            <tr><td>Preferred Instruments</td><td>{', '.join(patient_info.get('instruments', []))}</td></tr>
            <tr><td>Natural Elements</td><td>{', '.join(patient_info.get('natural_elements', []))}</td></tr>
        </table>

        <h2>Personality Profile (Big-5)</h2>
        <table>
            <tr><th>Trait</th><th>Score</th></tr>
            <tr><td>Extraversion</td><td>{big5_scores.get('extraversion', 0):.2f}/7.00</td></tr>
            <tr><td>Agreeableness</td><td>{big5_scores.get('agreeableness', 0):.2f}/7.00</td></tr>
            <tr><td>Conscientiousness</td><td>{big5_scores.get('conscientiousness', 0):.2f}/7.00</td></tr>
            <tr><td>Neuroticism</td><td>{big5_scores.get('neuroticism', 0):.2f}/7.00</td></tr>
            <tr><td>Openness</td><td>{big5_scores.get('openness', 0):.2f}/7.00</td></tr>
        </table>

        <h2>Music Recommendations</h2>
    """

    if 'categories' in recommendations:
        category_labels = {
            "birthplace_country": "From Your Country",
            "birthplace_city": "From Your City",
            "instruments": "Instrumental Favorites",
            "seasonal": "Seasonal Music",
            "natural_elements": "Nature-Inspired",
            "favorite_genre": "Favorite Genres",
            "favorite_musician": "Favorite Musician",
            "therapeutic": "Therapeutic Selections",
            "personality_based": "Personality Match",
            "calming_sensory": "Calming Sensory",
            "concentration": "Focus & Concentration",
            "binaural_beats": "Binaural Beats",
            "relief_study": "Study & Relief",
            "additional_calm": "Additional Calming",
            "additional_focus": "Additional Focus"
        }

        for category, data in recommendations['categories'].items():
            label = category_labels.get(category, category.replace('_', ' ').title())
            songs = data.get('songs', [])

            if songs:
                html_content += f"<h3>{label}</h3>"
                for idx, song in enumerate(songs, 1):
                    title = song.get('title', 'Unknown Title')
                    channel = song.get('channel', 'Unknown Channel')
                    url = song.get('url', 'No URL available')

                    html_content += f'<div class="song">'
                    html_content += f'<strong>{idx}. {title} - {channel}</strong><br>'

                    # Add YouTube link if available
                    if url and url.startswith('https://www.youtube.com/'):
                        html_content += f'<a href="{url}" target="_blank" style="color: #007bff; text-decoration: underline;">📺 YouTube Link</a><br>'
                    else:
                        html_content += f'<span style="color: #666;">URL: {url}</span><br>'

                    html_content += '</div>'

    html_content += """
    </body>
    </html>
    """

    return html_content

def create_json_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> str:
    """Create a JSON file with all data"""
    export_data = {
        "patient_info": patient_info,
        "recommendations": recommendations,
        "big5_scores": big5_scores,
        "export_timestamp": datetime.now().isoformat(),
        "export_version": "TheraMuse v9.0"
    }
    return json.dumps(export_data, indent=2, ensure_ascii=False)

def render_download_options(patient_info: Dict, recommendations: Dict, big5_scores: Dict):
    """Render download options at the end of the page"""
    st.markdown("---")

    # Scroll to Top Button - Goes to Personalized Music Recommendations
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("⬆ Scroll to Top", key="btn_scroll_to_recommendations", use_container_width=True):
            # Set session state flag to scroll to recommendations
            st.session_state._scroll_to_recommendations_flag = True
            # Also add a URL fragment for fallback
            st.query_params.scroll = "recommendations"
            st.rerun()

    # Research Evidence Section
    st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Research Evidence</h3>',
    unsafe_allow_html=True
)
    st.write("Scientific evidence supporting music therapy for various conditions and personality-based music preferences")

    if st.button(" View Research Evidence", key="research_evidence_btn", width='stretch'):
        st.session_state.page = "Research Evidence"
        st.rerun()

    st.markdown("""
    <div class='glass-card' style='margin-top: 1rem;'>
        <h2>Download Your Therapy Report</h2>
        <p style='color: rgba(255,255,255,0.7);'>Save your personalized music therapy recommendations in different formats</p>
    </div>
    """, unsafe_allow_html=True)

    # Create download data
    docx_data = create_docx_download(patient_info, recommendations, big5_scores)
    json_data = create_json_download(patient_info, recommendations, big5_scores)
    pdf_data = create_pdf_download(patient_info, recommendations, big5_scores)

    # Download buttons in columns
    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            label=" Download DOCX",
            data=docx_data,
            file_name=f"theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width='stretch'
        )

    with col2:
        st.markdown(f"""
        <a href="data:text/html;base64,{base64.b64encode(pdf_data.encode()).decode()}"
           download="theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
           style="
               display: block;
               width: 100%;
               padding: 0.75rem 1.5rem;
               background: linear-gradient(135deg, #338AFF 0%, #1E6FFF 100%);
               color: white;
               text-decoration: none;
               border-radius: 12px;
               font-weight: 600;
               text-align: center;
               transition: all 0.3s ease;
               box-shadow: 0 4px 16px rgba(51, 138, 255, 0.3);
               margin: 0;
        ">
             Download PDF
        </a>
        """, unsafe_allow_html=True)

    with col3:
        st.download_button(
            label=" Download JSON",
            data=json_data,
            file_name=f"theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            width='stretch'
        )

    st.markdown("""
    <div style='text-align: center; margin-top: 1rem; color: rgba(255,255,255,0.6); font-size: 0.9rem;'>
         <em>DOCX format includes formatted tables and is best for printing. JSON contains all raw data for developers. PDF (HTML) can be printed to PDF from your browser.</em>
    </div>
    """, unsafe_allow_html=True)

    # PERSONALITY & RAGA FUNCTIONS

def create_personality_radar(base_values: Dict, adjusted_values: Dict = None):
    """Create premium radar chart for Big-5 personality"""
    categories = ['Openness', 'Conscientiousness', 'Extraversion', 'Agreeableness', 'Neuroticism']
    
    fig = go.Figure()
    
    # Base personality with premium colors
    base_scores = [base_values.get(k.lower(), 4) for k in categories]
    fig.add_trace(go.Scatterpolar(
        r=base_scores,
        theta=categories,
        fill='toself',
        name='Base Personality',
        line=dict(color='#364153', width=3),
        fillcolor='rgba(163, 177, 138, 0.2)'
    ))
    
    # Adjusted personality
    if adjusted_values:
        adj_scores = [adjusted_values.get(k.lower(), 4) for k in categories]
        fig.add_trace(go.Scatterpolar(
            r=adj_scores,
            theta=categories,
            fill='toself',
            name='Adjusted (Nature)',
            line=dict(color='#364153', width=3),
            fillcolor='rgba(232, 221, 181, 0.2)'
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True, 
                range=[0, 7],
                gridcolor='rgba(255, 255, 255, 0.1)',
                linecolor='rgba(255, 255, 255, 0.2)'
            ),
            angularaxis=dict(
                gridcolor='rgba(255, 255, 255, 0.1)',
                linecolor='rgba(255, 255, 255, 0.2)'
            ),
            bgcolor='rgba(0, 0, 0, 0)'
        ),
        showlegend=True,
        title={
            'text': "Big-5 Personality Profile",
            'font': {'size': 24, 'color': "#0A55CE", 'family': 'Inter'}
        },
        paper_bgcolor='#0A55CE',
        plot_bgcolor='#0A55CE',
        font=dict(color='white', family='Inter')
    )
    
    return fig

def get_generational_ragas(birth_year: int) -> List[str]:
    """Get therapeutic ragas based on birth year"""
    matrix = BangladeshiGenerationalMatrix()
    context = matrix.get_generational_context(birth_year)
    return context.get("therapeutic_ragas", [])

# PREMIUM RECOMMENDATION DISPLAY

def extract_youtube_id(url: str) -> Optional[str]:
    """Best-effort extraction of a YouTube video ID from various URL shapes."""
    if not url:
        return None
    patterns = [
        r"youtu\.be/([A-Za-z0-9_-]{6,})",
        r"v=([A-Za-z0-9_-]{6,})",
        r"/embed/([A-Za-z0-9_-]{6,})",
        r"/shorts/([A-Za-z0-9_-]{6,})",
    ]
    for pat in patterns:
        m = re.search(pat, url)
        if m:
            vid = m.group(1)
            # Strip any trailing params just in case
            for sep in ("?", "&", "#"):
                vid = vid.split(sep)[0]
            return vid
    return None

def normalize_youtube_from_song(song: Dict) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Return (video_id, youtube_url, embed_url) using id dict, id str, or URL fallbacks."""
    video_id: Optional[str] = None
    # 1) id might be dict with videoId
    _id = song.get('id')
    if isinstance(_id, dict):
        video_id = _id.get('videoId') or _id.get('video_id')
    # 2) sometimes videoId is at top level
    if not video_id:
        video_id = song.get('videoId') or song.get('video_id')
    # 3) id might be a string containing either the id or a URL
    if not video_id and isinstance(_id, str):
        if re.fullmatch(r"[A-Za-z0-9_-]{6,}", _id):
            video_id = _id
        else:
            video_id = extract_youtube_id(_id)
    # 4) URL candidates
    url_candidates = [
        song.get('url'), song.get('youtube_url'), song.get('link'),
        song.get('webpage_url'), song.get('watch_url')
    ]
    youtube_url: Optional[str] = None
    for u in url_candidates:
        if not u:
            continue
        if not youtube_url and ("youtube" in u or "youtu.be" in u):
            youtube_url = u
        if not video_id:
            video_id = extract_youtube_id(u)
    # 5) Construct URLs from ID if needed
    if video_id and not youtube_url:
        youtube_url = f"https://www.youtube.com/watch?v={video_id}"
    embed_url = f"https://www.youtube.com/embed/{video_id}" if video_id else None
    return video_id, youtube_url, embed_url

def display_song_card(song: Dict, category: str, rank: int):
    """Display a premium song card with embedded YouTube player and link below title."""
    title = song.get('title', 'Unknown Title')
    channel = song.get('channel', 'Unknown Channel')
    description = song.get('description', '')[:200] + '...' if song.get('description') else ''

    video_id, youtube_url, youtube_embed = normalize_youtube_from_song(song)

    # Build link markup if we have any usable URL
    youtube_link_markup = (
        f'<a href="{youtube_url}" target="_blank" class="youtube-link" '
        f'style="font-size: 0.9rem; margin-top: 0.25rem; display: inline-block;"> Watch on YouTube</a>'
        if youtube_url else ''
    )

    # YouTube embed component - will be added separately
    youtube_embed_html = ''

    st.markdown(f"""
    <div class='song-card' style='animation-delay: {rank * 0.05}s;'>
        <div style='display: flex; align-items: center; margin-bottom: 1rem;'>
            <span class='rank-badge'>#{rank}</span>
            <div>
                <h4 style='margin: 0;'>{title}</h4>
                {youtube_link_markup}
            </div>
        </div>
        <p><strong>Channel:</strong> {channel}</p>
        <p><strong>Category:</strong> <span class='category-badge'>{category}</span></p>
        {f'<p style="font-size: 0.85rem; color: rgba(255,255,255,0.5);">{description}</p>' if description else ''}
    </div>
    """, unsafe_allow_html=True)

    # Add YouTube embed separately if we have a video
    if video_id and youtube_embed:
        st.components.v1.iframe(youtube_embed, height=315, width=560)

def render_recommendations_with_feedback(recommendations: Dict, patient_info: Dict, 
                                        session_id: str, patient_id: str):
    """Render premium recommendations with feedback options"""
    
    if not recommendations or 'categories' not in recommendations:
        st.warning("No recommendations available.")
        return
    
    # Add anchor for scrolling with style to make it visible for debugging
    st.markdown('<div id="personalized-recommendations" style="scroll-margin-top: 20px; min-height: 1px;"></div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class='glass-card'>
        <h2 style='margin-bottom: 0.5rem;'> Personalized Music Recommendations</h2>
        <p style='color: rgba(255,255,255,0.7); font-size: 1.1rem;'>
            <strong style='color: #A3B18A;'>{recommendations.get('total_songs', 0)}</strong> songs curated for your therapy
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Get TheraMuse instance
    if 'theramuse' not in st.session_state:
        st.session_state.theramuse = TheraMuse(db_path=str(get_database_path()))

    theramuse = st.session_state.theramuse
    
    # Display by category
    category_labels = {
        "birthplace_country": " From Your Country",
        "birthplace_city": " From Your City",
        "instruments": " Instrumental Favorites",
        "seasonal": " Seasonal Music",
        "natural_elements": " Nature-Inspired",
        "favorite_genre": " Favorite Genres",
        "favorite_musician": " Favorite Musician",
        "therapeutic": " Therapeutic Selections",
        "personality_based": " Personality Match",
        "calming_sensory": " Calming Sensory",
        "concentration": " Focus & Concentration",
        "binaural_beats": " Binaural Beats",
        "relief_study": " Study & Relief",
        "additional_calm": " Additional Calming",
        "additional_focus": "Additional Focus"
    }
    
    for category, data in recommendations['categories'].items():
        label = category_labels.get(category, category.replace('_', ' ').title())
        songs = data.get('songs', [])
        
        if not songs:
            continue
            
        with st.expander(f"{label} ({len(songs)} songs)", expanded=True):
                  
            for idx, song in enumerate(songs, 1):
                col1, col2 = st.columns([5, 1])
                
                with col1:
                    display_song_card(song, label, idx)
                
                with col2:
                    feedback_key = f"feedback_{session_id}_{category}_{idx}"
                    
                    st.markdown("<div class='feedback-button'>", unsafe_allow_html=True)
                    
                    # Feedback buttons
                    if st.button(" Like", key=f"like_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id, 
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "like", patient_info
                        )
                        st.success(" Feedback recorded!")
                        st.toast("Thank you for your feedback!")
                    
                    if st.button(" Dislike", key=f"dislike_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id,
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "dislike", patient_info
                        )
                        st.info("Feedback recorded!")
                    
                    if st.button(" Skip", key=f"skip_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id,
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "skip", patient_info
                        )
                        st.info("Skipped!")
                    
                    st.markdown("</div>", unsafe_allow_html=True)

# STREAMLIT PAGES

def page_intake():
    """Patient intake page with premium design"""
    render_logo(size=220)

    st.markdown("""
    <div class='glass-card'>
        <h2>Patient Intake</h2>
        <p style='color: #364153; font-size: 1.05rem;'>
            Complete the assessment to receive TheramuseRX personalized music recommendations.
        </p>
    </div>
    """, unsafe_allow_html=True)


    # Initialize session state flags
    if 'show_results' not in st.session_state:
        st.session_state.show_results = False
    if 'processing_complete' not in st.session_state:
        st.session_state.processing_complete = False

    # Only show form if not currently showing results
    if not st.session_state.show_results:
        # Initialize session state for form data if not present
        if 'form_data' not in st.session_state:
            st.session_state.form_data = {}

        # Text inputs outside of form to prevent Enter submission
        st.markdown('<h3 style="color:#364153; font-size:28px; font-weight:700;">Demographics</h3>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)

        # Use session state to preserve input values
        name = col1.text_input(
            "Name",
            placeholder="John Doe",
            value=st.session_state.form_data.get('name', ''),
            key="name_input"
        )
        st.session_state.form_data['name'] = name

        birthplace_city = col1.text_input(
            "Birthplace City",
            placeholder="e.g., Dhaka",
            value=st.session_state.form_data.get('birthplace_city', ''),
            key="birthplace_city_input"
        )
        st.session_state.form_data['birthplace_city'] = birthplace_city

        birthplace_country = col2.text_input(
            "Birthplace Country *",
            placeholder="e.g., Bangladesh, India, USA",
            value=st.session_state.form_data.get('birthplace_country', ''),
            key="birthplace_country_input"
        )
        st.session_state.form_data['birthplace_country'] = birthplace_country

        dob = col1.date_input(
            "Date of Birth *",
            value=st.session_state.form_data.get('dob', None),
            min_value=date(1900, 1, 1),
            max_value=date.today(),
            key="dob_input"
        )
        st.session_state.form_data['dob'] = dob

        sex = col1.selectbox(
            "Sex",
            ["Male", "Female", "Other"],
            index=["Male", "Female", "Other"].index(st.session_state.form_data.get('sex', 'Male')) if st.session_state.form_data.get('sex') in ["Male", "Female", "Other"] else 0,
            key="sex_input"
        )
        st.session_state.form_data['sex'] = sex

        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Musical Preferences</h3>',
    unsafe_allow_html=True
)
        preferred_instruments = st.multiselect(
            "Preferred Instruments",
            ["Piano", "Guitar", "Violin", "Flute", "Sitar", "Tabla", "Drums", "Saxophone"],
            default=st.session_state.form_data.get('instruments', []),
            key="instruments_input"
        )
        st.session_state.form_data['instruments'] = preferred_instruments

        col1, col2 = st.columns(2)
        with col1:
            # Support multiple favorite genres
            st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Favorite Music Genres (select up to 3)</h3>',
    unsafe_allow_html=True
)
            favorite_genre1 = st.selectbox(
                "Favorite Genre 1 *",
                ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                index=["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"].index(st.session_state.form_data.get('favorite_genre1', '')) if st.session_state.form_data.get('favorite_genre1') in ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"] else 0,
                key="favorite_genre1"
            )
            st.session_state.form_data['favorite_genre1'] = favorite_genre1

            favorite_genre2 = st.selectbox(
                "Favorite Genre 2 (optional)",
                ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                index=["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"].index(st.session_state.form_data.get('favorite_genre2', '')) if st.session_state.form_data.get('favorite_genre2') in ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"] else 0,
                key="favorite_genre2"
            )
            st.session_state.form_data['favorite_genre2'] = favorite_genre2

            favorite_genre3 = st.selectbox(
                "Favorite Genre 3 (optional)",
                ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                index=["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"].index(st.session_state.form_data.get('favorite_genre3', '')) if st.session_state.form_data.get('favorite_genre3') in ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"] else 0,
                key="favorite_genre3"
            )
            st.session_state.form_data['favorite_genre3'] = favorite_genre3

            # Combine favorite genres into a comma-separated string
            favorite_genres = [genre for genre in [favorite_genre1, favorite_genre2, favorite_genre3] if genre]
            favorite_genre = ", ".join(favorite_genres) if favorite_genres else ""
            st.session_state.form_data['favorite_genre'] = favorite_genre

        with col2:
            favorite_musician = st.text_input(
                "Favorite Musician/Artist",
                placeholder="e.g., Opeth, Thom Yorke, Artcell",
                value=st.session_state.form_data.get('favorite_musician', ''),
                key="favorite_musician_input"
            )
            st.session_state.form_data['favorite_musician'] = favorite_musician

        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Environmental Preferences</h3>',
    unsafe_allow_html=True
)
        col1, col2 = st.columns(2)
        with col1:
            favorite_season = st.selectbox(
                "Favorite Season",
                ["", "Spring", "Summer", "Monsoon", "Autumn", "Winter"],
                index=["", "Spring", "Summer", "Monsoon", "Autumn", "Winter"].index(st.session_state.form_data.get('favorite_season', '')) if st.session_state.form_data.get('favorite_season') in ["", "Spring", "Summer", "Monsoon", "Autumn", "Winter"] else 0,
                key="favorite_season"
            )
            st.session_state.form_data['favorite_season'] = favorite_season
        with col2:
            natural_elements = st.multiselect(
                "Natural Elements You Connect With",
                ["Rain", "Forest", "Ocean", "Mountains", "Desert", "Rivers", "Sunset", "Sunrise"],
                default=st.session_state.form_data.get('natural_elements', []),
                key="natural_elements"
            )
            st.session_state.form_data['natural_elements'] = natural_elements

        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Health & Wellness Profile</h3>',
    unsafe_allow_html=True
)
        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Select the condition for therapy:</h3>',
    unsafe_allow_html=True
)

        condition = st.selectbox(
            "Primary Condition *",
            ["Dementia / Alzheimer's", "Down Syndrome", "ADHD"],
            index=["Dementia / Alzheimer's", "Down Syndrome", "ADHD"].index(st.session_state.form_data.get('condition', 'Dementia / Alzheimer\'s')) if st.session_state.form_data.get('condition') in ["Dementia / Alzheimer's", "Down Syndrome", "ADHD"] else 0,
            key="condition"
        )
        st.session_state.form_data['condition'] = condition

        # Memory & Sleep Assessment - shown for all conditions
        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Memory & Sleep Assessment</h3>',
    unsafe_allow_html=True
)
        difficulty_sleeping = st.checkbox(
            "Difficulty sleeping?",
            value=st.session_state.form_data.get('difficulty_sleeping', False),
            key="difficulty_sleeping"
        )
        st.session_state.form_data['difficulty_sleeping'] = difficulty_sleeping

        trouble_remembering = st.checkbox(
            "Trouble remembering recent things?",
            value=st.session_state.form_data.get('trouble_remembering', False),
            key="trouble_remembering"
        )
        st.session_state.form_data['trouble_remembering'] = trouble_remembering

        forgets_everyday_things = st.checkbox(
            "Often forgets everyday things?",
            value=st.session_state.form_data.get('forgets_everyday_things', False),
            key="forgets_everyday_things"
        )
        st.session_state.form_data['forgets_everyday_things'] = forgets_everyday_things

        difficulty_recalling_old_memories = st.checkbox(
            "Difficulty recalling older memories?",
            value=st.session_state.form_data.get('difficulty_recalling_old_memories', False),
            key="difficulty_recalling_old_memories"
        )
        st.session_state.form_data['difficulty_recalling_old_memories'] = difficulty_recalling_old_memories

        memory_worse_than_year_ago = st.checkbox(
            "Memory worse than a year ago?",
            value=st.session_state.form_data.get('memory_worse_than_year_ago', False),
            key="memory_worse_than_year_ago"
        )
        st.session_state.form_data['memory_worse_than_year_ago'] = memory_worse_than_year_ago

        visited_mental_health = st.checkbox(
            "Visited mental health professional?",
            value=st.session_state.form_data.get('visited_mental_health', False),
            key="visited_mental_health"
        )
        st.session_state.form_data['visited_mental_health'] = visited_mental_health

        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Personality Assessment (Big-5)</h3>',
    unsafe_allow_html=True
)
        st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Rate each statement from 1 (Strongly Disagree) to 7 (Strongly Agree)</h3>',
    unsafe_allow_html=True
)

        statements = [
            "I see myself as extraverted, enthusiastic",
            "I see myself as critical, quarrelsome",
            "I see myself as dependable, self-disciplined",
            "I see myself as anxious, easily upset",
            "I see myself as open to new experiences, complex",
            "I see myself as reserved, quiet",
            "I see myself as sympathetic, warm",
            "I see myself as disorganized, careless",
            "I see myself as calm, emotionally stable",
            "I see myself as conventional, uncreative"
        ]

        big5_items = [slider_with_ticks(f"{i+1}. {s}", key=f"big5_{i+1}") for i, s in enumerate(statements)]

        # Now create a simple form just for the submission button
        with st.form("submission_form"):
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                submitted = st.form_submit_button("Theramuse Recommendations", type="primary", width='stretch')

        # ONLY process when form is explicitly submitted
        if submitted:
            # Get data from session state
            form_data = st.session_state.form_data

            # Validate required fields
            if not form_data.get('dob') or not form_data.get('birthplace_country'):
                st.error(" Please fill in Date of Birth and Birthplace Country.")
                st.stop()  # Stop execution here

            # Validate at least one favorite genre is selected
            if not form_data.get('favorite_genre1'):
                st.error(" Please select at least one favorite music genre.")
                st.stop()  # Stop execution here

            # Calculate age and birth year
            dob = form_data.get('dob')
            age = compute_age_from_dob(dob)
            birth_year = dob.year

            # Calculate Big-5 scores
            E_base = (big5_items[0] + reverse_1to7(big5_items[5])) / 2.0
            A_base = (reverse_1to7(big5_items[1]) + big5_items[6]) / 2.0
            C_base = (big5_items[2] + reverse_1to7(big5_items[7])) / 2.0
            N_base = (big5_items[3] + reverse_1to7(big5_items[8])) / 2.0
            O_base = (big5_items[4] + reverse_1to7(big5_items[9])) / 2.0

            big5_scores = {
                "extraversion": E_base,
                "agreeableness": A_base,
                "conscientiousness": C_base,
                "neuroticism": N_base,
                "openness": O_base
            }

            # Determine actual condition
            actual_condition = get_condition_code(form_data.get('condition'))

            # Build patient info dictionary from session state
            patient_info = {
                "name": form_data.get('name') or "Anonymous",
                "age": age,
                "birth_year": birth_year,
                "sex": form_data.get('sex'),
                "birthplace_city": form_data.get('birthplace_city', '').strip(),
                "birthplace_country": form_data.get('birthplace_country', '').strip(),
                "instruments": form_data.get('instruments', []),
                "favorite_genre": form_data.get('favorite_genre', ''),
                "favorite_musician": form_data.get('favorite_musician', '').strip(),
                "favorite_season": form_data.get('favorite_season', ''),
                "natural_elements": form_data.get('natural_elements', []),
                "condition": actual_condition,
                "difficulty_sleeping": form_data.get('difficulty_sleeping', False),
                "trouble_remembering": form_data.get('trouble_remembering', False),
                "forgets_everyday_things": form_data.get('forgets_everyday_things', False),
                "difficulty_recalling_old_memories": form_data.get('difficulty_recalling_old_memories', False),
                "memory_worse_than_year_ago": form_data.get('memory_worse_than_year_ago', False),
                "visited_mental_health_professional": form_data.get('visited_mental_health', False),
                "big5_scores": big5_scores
            }

            # Generate patient ID
            patient_id = f"patient_{datetime.now().strftime('%Y%m%d%H%M%S')}"

            # Initialize TheraMuse
            with st.spinner(f" TheramuseRX is generating personalized recommendations "):
                try:
                    if 'theramuse' not in st.session_state:
                        st.session_state.theramuse = TheraMuse(db_path=str(get_database_path()))

                    theramuse = st.session_state.theramuse

                    # Get recommendations
                    recommendations = theramuse.get_therapy_recommendations(
                        patient_info, actual_condition, patient_id
                    )
                except Exception as db_error:
                    st.error(f"Database initialization error: {str(db_error)}")
                    st.error("Please try again or check database permissions.")
                    st.stop()

                # Save to our enhanced patient database
                try:
                    db_patient_id = save_patient_to_database(
                        patient_info, big5_scores, recommendations,
                        recommendations.get('session_id', f"session_{datetime.now().strftime('%Y%m%d%H%M%S')}")
                    )
                    st.success(f"  Patient data saved to database! ID: {db_patient_id}")
                    st.info(" You can view all patient records in the 'Patient Database' section")
                except Exception as e:
                    st.warning(f"  Could not save to enhanced database: {str(e)}")
                    st.info("Recommendations will still be displayed, but enhanced patient features won't be available.")
                    st.info("Check the 'Patient Database' section to see if other patients are stored there.")

                # Store in session state
                st.session_state.tm_recs = recommendations
                st.session_state.tm_patient_data = patient_info
                st.session_state.tm_patient_id = patient_id
                st.session_state.tm_session_id = recommendations.get('session_id')
                st.session_state.tm_b5_scores = big5_scores
                st.session_state.show_results = True
                st.session_state.processing_complete = True

            st.success(f" Patient intake completed! Patient ID: {patient_id}")
            st.snow()
            # Set flag to scroll to top after rerun
            st.session_state.scroll_to_top = True
            st.rerun()  # Rerun to show results
    
    # Display results ONLY if processing is complete and show_results is True
    if st.session_state.show_results and st.session_state.processing_complete and 'tm_recs' in st.session_state:
        # Add anchor at top for scrolling
        st.markdown('<div id="results-top"></div>', unsafe_allow_html=True)

        st.markdown(
    f"""
    <div style="background-color:#E8F1FF; padding:10px; border-radius:10px;">
        <p style="color:#338AFF; font-weight:600; margin:0;">
            TheraMuseRX's {st.session_state.tm_recs.get('total_songs', 0)} recommendations!
        </p>
    </div>
    """,
    unsafe_allow_html=True
)


        # Start new intake button at top
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(" Start New Patient Intake", type="secondary", width='stretch'):
                # Clear all session state
                for key in ['tm_recs', 'tm_patient_data', 'tm_patient_id', 'tm_session_id', 'tm_b5_scores', 'show_results', 'processing_complete', 'form_data']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

        tab1, tab2, tab3 = st.tabs([" Recommendations", " Personality Profile", " Bandit Statistics"])

        with tab1:
            render_recommendations_with_feedback(
                st.session_state.tm_recs,
                st.session_state.tm_patient_data,
                st.session_state.tm_session_id,
                st.session_state.tm_patient_id
            )

        with tab2:
            st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Big-5 Personality Profile</h3>',
    unsafe_allow_html=True
)
            fig = create_personality_radar(st.session_state.tm_b5_scores)
            st.plotly_chart(fig, use_container_width=True)

            # Show generational context if dementia
            if st.session_state.tm_patient_data.get('condition') == 'dementia':
                birth_year = st.session_state.tm_patient_data.get('birth_year')
                if birth_year:
                    ragas = get_generational_ragas(birth_year)
                    if ragas:
                        st.markdown("""
                        <div class='glass-card'>
                            <h3 style='color:#364153;'> Therapeutic Ragas for Your Generation</h3>
                        </div>
                        """, unsafe_allow_html=True)
                        for raga in ragas:
                            st.markdown(f"<span class='category-badge' style='color:#364153;'>{raga}</span>", unsafe_allow_html=True)

        with tab3:
            st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Thompson Sampling Statistics</h3>',
    unsafe_allow_html=True
)
            bandit_stats = st.session_state.tm_recs.get('bandit_stats', {})

            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3 style='color:#364153;'>Total Interactions</h3>
                    <h2>{bandit_stats.get('n_interactions', 0)}</h2>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3 style='color:#364153;'>Average Reward</h3>
                    <h2>{bandit_stats.get('avg_reward', 0):.3f}</h2>
                </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3 style='color:#364153;'>Exploration Rate</h3>
                    <h2>{bandit_stats.get('exploration_rate', 0):.3f}</h2>
                </div>
                """, unsafe_allow_html=True)

            st.info(" The system uses Linear Thompson Sampling to learn from your feedback and improve recommendations over time.")

        # Add download options at the end
        render_download_options(
            st.session_state.tm_patient_data,
            st.session_state.tm_recs,
            st.session_state.tm_b5_scores
        )

def page_analytics():
    """Analytics dashboard page with premium design"""
    render_logo()
    st.markdown("""
    <div class='glass-card'>
        <h2> Analytics Dashboard</h2>
        <p style='color: rgba(255,255,255,0.7);'>Monitor system performance and user engagement</p>
    </div>
    """, unsafe_allow_html=True)
    
    if 'theramuse' not in st.session_state:
        st.session_state.theramuse = TheraMuse(db_path=str(get_database_path()))

    theramuse = st.session_state.theramuse
    analytics = theramuse.get_analytics()
    
    # Top metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f"""
        <div class='metric-card'>
            <h3 style='color:#364153;'>Total Sessions</h3>
            <h2 style='color:#364153;'>{analytics.get('total_sessions', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class='metric-card'>
            <h3 style='color:#364153;'>Total Feedback</h3>
            <h2 style='color:#364153;'>{analytics.get('total_feedback', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown(f"""
        <div class='metric-card'>
            <h3 style='color:#364153;'>Total Patients</h3>
            <h2 style='color:#364153;'>{analytics.get('total_patients', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    # Rewards by condition
    st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Average Reward by Condition</h3>',
    unsafe_allow_html=True
)
    rewards_data = analytics.get('rewards_by_condition', [])
    
    if rewards_data:
        df_rewards = pd.DataFrame(rewards_data)
        fig = px.bar(df_rewards, x='condition', y='avg_reward', 
                    title='Average Reward by Medical Condition',
                    color='count', color_continuous_scale='Viridis')
        fig.update_layout(
            paper_bgcolor='rgba(0, 0, 0, 0)',
            plot_bgcolor='rgba(0, 0, 0, 0)',
            font=dict(color='white', family='Inter')
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No feedback data available yet. Start collecting feedback to see analytics!")
    
    # YouTube API Health
    st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">System Health</h3>',
    unsafe_allow_html=True
)
    health = theramuse.check_api_health()
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.json(health.get('youtube_api', {}))
        st.markdown("</div>", unsafe_allow_html=True)
    with col2:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.json(health.get('database', {}))
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Cache management
    st.markdown(
    '<h3 style="color:#364153; font-size:28px; font-weight:700;">Cache Management</h3>',
    unsafe_allow_html=True
)
    cache_status = theramuse.get_youtube_cache_status()
    st.metric("YouTube Cache Size", cache_status.get('cache_size', 0))
    
    if st.button("Clear YouTube Cache", type="secondary"):
        theramuse.clear_youtube_cache()
        st.success(" YouTube cache cleared!")
        st.snow()

def page_research_evidence():
    """Research Evidence page with comprehensive scientific literature"""
    render_logo()

    st.markdown("""
    <div class='glass-card'>
        <h2>Research Evidence</h2>
        <p style='color: rgba(255,255,255,0.7); font-size: 1.05rem;'>
            Comprehensive scientific evidence supporting music therapy for various conditions and personality-based music preferences
        </p>
    </div>
    """, unsafe_allow_html=True)

    # # Back button
    if st.button("← Back to Previous Page", type="secondary"):
        st.session_state.page = " Patient Intake"
        st.rerun()

    st.markdown("## Research Evidence of Music Therapy for Dementia, ADHD, Down Syndrome, and Big Five Personality Traits & Music Genre Preferences")

    # Updated long-form compilation provided by user
    with st.expander("**Click Here to Read  Research Evidence (Detailed Compilation**)", expanded=False):
        st.markdown(
            """
            <div style='color: #000000; line-height: 1.6;'>
            <h3 style='color: #000000; margin-bottom: 1rem;'>Research Papers on Music Therapy for Dementia and Big Five Personality Traits</h3>

            1. Music Therapy in the Treatment of Dementia: A Systematic Review and Meta-Analysis (Moreno-Morales et al., 2020)
            - Link: https://pubmed.ncbi.nlm.nih.gov/32509790/
            - Study Type: Systematic review and meta-analysis
            - Key Findings: Cognitive function: Music therapy significantly improves cognitive function in people with dementia; Quality of life: Improvements in quality of life after intervention; Depression: Long-term depression improvements demonstrated; Limitations: No evidence of improvement in quality of life long-term or short-term depression
            - Databases Searched: Medline, PubMed Central, Embase, PsycINFO, and Cochrane Library
            - Conclusion: Music therapy shows promise as a powerful treatment strategy, but standardized protocols need to be developed depending on dementia stage.

            2. The Effects of Music-Based Interventions on Behavioural and Psychological Symptoms of Dementia (de Witte et al., 2024)
            - Link: https://www.tandfonline.com/doi/full/10.1080/13607863.2024.2373969
            - Study Type: Systematic review and network meta-analysis protocol
            - Key Findings: Previous meta-analyses showed music therapy reduces agitation (g = −0.66; 9 RCTs) and anxiety (g = −0.51; 5 RCTs). Active music therapy by music therapists showed positive effects on global cognition (SMD = 0.29; 3 RCTs). Music therapist-delivered active therapy was more effective (SMD −3.00; 15 RCTs) than music listening by other healthcare professionals (SMD −2.06; 15 RCTs) for depression. Substantial heterogeneity exists due to intervention types, study design, and instruments.
            - Significance: First network meta-analysis comparing different types of music-based interventions for BPSD.

            3. Cochrane Review: Music-Based Therapeutic Interventions for People with Dementia (van der Steen et al., 2018/Updated 2024)
            - Link: https://www.cochrane.org/evidence/CD003477_does-music-based-therapy-help-people-dementia
            - Alternative Link: https://pubmed.ncbi.nlm.nih.gov/30033623/
            - Study Type: Cochrane systematic review with meta-analysis
            - Sample: 30 studies with 1,720 randomized participants across 15 countries
            - Key Findings: Depressive symptoms: Moderate-certainty evidence of slight improvement (SMD −0.23); Overall behavioral problems: Low-certainty evidence of improvement (SMD −0.31); Agitation/aggression: Likely no improvement (SMD −0.05); No significant improvements in emotional well-being, anxiety, social behavior, or cognition at end of treatment; Long-term effects (4+ weeks after) may be smaller and remain uncertain.
            - Lead Author Quote: “Music therapy is a drug-free way of helping people feel less sad and less anxious... reasonable alternative to pharmacological approaches.”

            4. HOMESIDE: Home-Based Family Caregiver-Delivered Music Intervention (Baker et al., 2023)
            - Link: https://www.thelancet.com/journals/eclinm/article/PIIS2589-5370(23)00401-7/fulltext
            - Study Type: Randomized controlled trial protocol (2 × 2 factorial cluster design)
            - Innovation: Tests whether family caregivers can deliver music interventions at home
            - Primary Outcome: Reduction in BPSD after 90 days compared to standard care
            - Mechanisms Identified: Activation of neuroplastic and neurochemical processes; Auditory-motor coupling; Neural entrainment; Arousal-mood pathways; Autobiographical and implicit memory activation
            - Significance: Could enable longer home living and reduce admissions.

            5. MIDDEL Trial: Clinical Effectiveness of Music Interventions for Dementia and Depression in Elderly Care (Baker et al., 2022)
            - Link: https://www.sciencedirect.com/science/article/pii/S2666756822000277
            - Alternative: https://www.thelancet.com/journals/lanhl/article/PIIS2666-7568(22)00027-7/fulltext
            - Study Type: 2 × 2 factorial cluster-RCT comparing Group Music Therapy (GMT), Recreational Choir Singing (RCS), GMT+RCS, and standard care
            - Key Findings: RCS reduced depressive symptoms at end of 6 months; Positive effects on neuropsychiatric symptoms and quality of life; Singing effects on depression sustained at 12 months; Interventions more intensive and standardized than prior work.
            - Significance: First large-scale direct comparison of music therapy vs recreational singing; singing shows durable benefits.

            6. ALMUTH Study: 12-Month Randomised Pilot Trial (Matziorinis et al., 2023)
            - Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC10114372/
            - Study Type: 12-month three-arm RCT (music therapy via choir singing vs physical activity vs control) in mild-to-moderate AD
            - Unique: Longest duration music therapy RCT to date; includes structural/functional MRI and DTI; 45–60 min choir sessions, 4x/month for 12 months
            - Neurobiological Rationale: Musical memory networks are spared late in AD; music activates broad networks; preserved musical memory supports familiar recognition.

            7. The Promise of Music Therapy for Alzheimer’s Disease: A Review (Fang et al., 2022)
            - Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC9796133/
            - Focus: Music-Evoked Autobiographical Memories (MEAMs) remain preserved; self-selected music evokes faster, more specific, and more emotional memories; unfamiliar music can also enhance recall and reduce trait anxiety; anterior hippocampus links emotion and autobiographical memory.

            8. Individualized Music Listening RCT Protocol (Jakob et al., 2024)
            - Link: https://pubmed.ncbi.nlm.nih.gov/38532365/
            - Intervention: App-based individualized music listening delivered by family caregivers at home (20 min every other day × 6 weeks)
            - Outcomes: Well-being, physiological stress (hair cortisol), QoL, BPSD, resistance during care, caregiver burden.

            9. Personalized Playlists and Individual Differences (Garrido et al., 2018)
            - Link: https://pubmed.ncbi.nlm.nih.gov/29966193/
            - Finding: Personalization is necessary but not sufficient; depression/apathy severity and cognitive impairment level moderate responses; high depression may increase sadness.

            10. How and Why Music Therapy Reduces Distress in Advanced Dementia (Thompson et al., 2024)
            - Link: https://www.nature.com/articles/s44220-024-00342-x
            - Study Type: Realist review developing program theory
            - Mechanisms: Core intervention elements; individual/interpersonal/institutional contexts; hidden mechanisms include meeting unmet needs and improving communication; integrates neurological, social, psychological models.

            <h4 style='color: #000000; margin-bottom: 0.75rem; margin-top: 1.5rem;'>Big Five Personality Traits and Music Preferences: Key Research Papers</h4>
            11. Associations Between Personality Traits and Music Preference (Boccia, 2020)
            - Link: https://digitalcommons.lindenwood.edu/cgi/viewcontent.cgi?article=1000&context=psych_journals
            - Finding: Overall enjoyment correlates positively with Openness, Extraversion, Agreeableness; negatively with Neuroticism; Conscientiousness not significant. Genre specifics: Openness→rap/classical; Extraversion→country/pop/electronic; Agreeableness→rap/country/pop; Conscientiousness→country; Neuroticism negatively with new age/classical.

            12. The Structure of Musical Preferences: MUSIC Model (Rentfrow, Goldberg & Levitin, 2011)
            - Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC3138530/
            - Factors: Mellow, Urban, Sophisticated, Intense, Campestral.

            13. Big Five and Embodied Musical Emotions (Mendoza Garay, 2023)
            - Link: https://journals.sagepub.com/doi/10.1177/03057356221135355
            - Finding: Openness and Agreeableness most strongly linked to embodied musical emotions.

            14. Global Personality-Music Links (Greenberg et al., 2022)
            - Link: https://www.technologynetworks.com/neuroscience/news/nirvana-for-neuroticism-how-musical-preferences-match-personality-traits-around-the-world-358429
            - Findings: Extraversion→contemporary/upbeat; Openness→sophisticated; Agreeableness→upbeat/conventional; Conscientiousness→negative with intense; Neuroticism→prefers intense styles.

            15. Music Listening and High Neuroticism (Chan et al., 2024)
            - Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC11129041/
            - Findings: Music reduces physiological stress in both high/low neuroticism; high neuroticism group reports smaller subjective improvements (cognitive bias); use physiological as well as self-report measures.

            16. Personality Computing with Naturalistic Music Listening (Hansen et al., 2023)
            - Link: https://online.ucpress.edu/collabra/article/9/1/75214/196347/Personality-Computing-With-Naturalistic-Music
            - Findings: Openness most related to music listening (r=.25); Conscientiousness second (r=.13); audio features best for Openness, lyrics for Conscientiousness.

            17. Big Five Domains and Music Listening (Hansen et al., 2025)
            - Link: https://www.nature.com/articles/s41598-025-95661-z
            - Findings: Open-mindedness most strongly associated; Agreeableness and Neuroticism also strongly associated.

            18. Musical Capacity and Big Five (2023)
            - Link: https://ijip.in/wp-content/uploads/2023/02/18.01.087.20231101.pdf
            - Findings: Listening sophistication correlated with all Big Five; all musical capacity factors correlated with Openness; indifference to music negatively with Openness/Conscientiousness.

            19. Personality, Music Listening, and Well-Being (Martarelli et al., 2024)
            - Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC11064775/
            - Mechanisms: Emotion regulation, empathy, state absorption; stable traits involved: Neuroticism, absorption, Openness; listening to sad music often reflects immersion/absorption, not mood.

            ---
            <h3 style='color: #000000; margin-bottom: 0.75rem; margin-top: 1.5rem;'>How These Papers Apply to Dementia and Big Five Personality</h3>
            - Individualized vs Generic Music: Personalization matters but must consider depression, apathy, severity of impairment.
            - Active vs Passive Engagement: Active music-making by trained therapists more effective than passive listening, especially for depression.
            - Autobiographical Memory Activation: MEAMs explain why familiar music works; anterior hippocampus links emotion and memory.
            - Caregiver-Delivered Interventions: Family-delivered interventions at home are promising and scalable.
            - BPSD Focus: Strongest effects for depression and behavioral problems; modest for anxiety; limited for agitation/aggression.

            <h3 style='color: #000000; margin-bottom: 0.75rem; margin-top: 1.5rem;'>Personality-Matched Music Selection</h3>
            - Openness: Diverse/sophisticated genres; world/jazz/classical/experimental.
            - Extraversion: Contemporary, rhythmic, upbeat; pop/electronic/country.
            - Agreeableness: Positive-valence, conventional, communal; pop/country/soul.
            - Neuroticism: Intense styles for catharsis; monitor relaxation attempts; use physiological measures.
            - Conscientiousness: Focus on lyrics/meaning rather than audio features.

            <h3 style='color: #000000; margin-bottom: 0.75rem; margin-top: 1.5rem;'>Combined Application (Dementia + Personality)</h3>
            <div style='color: #000000; margin-bottom: 1rem;'>
            Tailor selections to lifelong personality and autobiographical era (youth/young adult music), prioritize active engagement, and monitor outcomes behaviorally and physiologically when neuroticism is high.
            </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("###  Music Therapy for Dementia: Key Research Papers")
    st.markdown("#### 1. Systematic Reviews and Meta-Analyses")

    with st.expander(" **Music Therapy in the Treatment of Dementia: A Systematic Review and Meta-Analysis** (Moreno-Morales et al., 2020)", expanded=False):
        st.markdown("""
        <a href="https://pubmed.ncbi.nlm.nih.gov/32509790/" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Systematic review and meta-analysis of 8 studies with 816 participants

        **Key Findings:**
        -  **Cognitive function**: Music therapy significantly improves cognitive function (SMD = -0.23, 95% CI: -0.44, -0.02)
        -  **Quality of life**: Improvements in quality of life after intervention
        - **Depression**: Long-term depression improvements demonstrated
        - **Limitations**: No evidence of improvement in quality of life long-term or short-term depression

        **Databases Searched:** Medline, PubMed Central, Embase, PsycINFO, and Cochrane Library

        **Conclusion:** Music therapy shows promise as a powerful treatment strategy, but standardized protocols need to be developed depending on dementia stage.
        """, unsafe_allow_html=True)

    with st.expander(" **The Effects of Music-Based Interventions on Behavioural and Psychological Symptoms of Dementia** (de Witte et al., 2024)", expanded=False):
        st.markdown("""
        <a href="https://www.tandfonline.com/doi/full/10.1080/13607863.2024.2373969" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Systematic review and network meta-analysis protocol

        **Key Findings:**
        - Previous meta-analyses showed music therapy reduces **agitation** (effect size g = −0.66; 9 RCTs) and **anxiety** (g = −0.51; 5 RCTs)
        - Active music therapy delivered by music therapists showed positive effects on **global cognition** (SMD = 0.29; 3 RCTs)
        - Music therapist-delivered active therapy was more effective (SMD −3.00; 15 RCTs) than music listening by other healthcare professionals (SMD −2.06; 15 RCTs) for **depression**
        - **Substantial heterogeneity** exists due to intervention types, study design, and measurement instruments

        **Significance:** First network meta-analysis comparing different types of music-based interventions for BPSD.
        """, unsafe_allow_html=True)

    with st.expander(" **Cochrane Review: Music-Based Therapeutic Interventions for People with Dementia** (van der Steen et al., 2018/Updated 2024)", expanded=False):
        st.markdown("""
        <a href="https://www.cochrane.org/evidence/CD003477_does-music-based-therapy-help-people-dementia" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>
        <br/>
        <a href="https://pubmed.ncbi.nlm.nih.gov/30033623/" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Alternative Link</a>

        **Study Design:** Cochrane systematic review with meta-analysis
        **Sample:** 30 studies with 1,720 randomized participants across 15 countries

        **Key Findings:**
        - **Depressive symptoms**: Moderate-certainty evidence that music therapy probably improved depressive symptoms slightly (SMD −0.23, 95% CI −0.42 to −0.04; 9 studies, 441 participants)
        - **Overall behavioral problems**: Low-certainty evidence it may improve overall behavioral problems (SMD −0.31, 95% CI −0.60 to −0.02; 10 studies, 385 participants)
        - **Agitation/aggression**: Moderate-certainty evidence that music therapy likely did NOT improve agitation or aggression (SMD −0.05)
        - No significant improvements in emotional well-being, anxiety, social behavior, or cognition at end of treatment
        - **Long-term effects** (4+ weeks after treatment): May be smaller and remain uncertain

        **Lead Author Quote:** *"Music therapy is a drug-free way of helping people feel less sad and less anxious. Looking at the effect sizes, music therapy is a reasonable alternative to pharmacological approaches and is much more person-centered."*
        """, unsafe_allow_html=True)

    with st.expander("**HOMESIDE: Home-Based Family Caregiver-Delivered Music Intervention** (Baker et al., 2023)", expanded=False):
        st.markdown("""
        <a href="https://www.thelancet.com/journals/eclinm/article/PIIS2589-5370(23)00401-7/fulltext" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Randomized controlled trial protocol (2 × 2 factorial cluster design)

        **Innovation:** Tests whether **family caregivers (not professionals)** can deliver music interventions at home

        **Primary Outcome:** Reduction in BPSD after 90 days compared to standard care

        **Mechanisms Identified:**
        - **Activation of neuroplastic and neurochemical processes**
        - **Auditory-motor coupling**
        - **Neural entrainment**
        - **Arousal-mood pathways**
        - **Autobiographical and implicit memory activation**

        **Significance:** Could enable people with dementia to live at home longer and reduce hospital/care home admissions.
        """, unsafe_allow_html=True)

    with st.expander(" **MIDDEL Trial: Clinical Effectiveness of Music Interventions for Dementia and Depression** (Baker et al., 2022)", expanded=False):
        st.markdown("""
        <a href="https://www.sciencedirect.com/science/article/pii/S2666756822000277" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>
        <br/>
        <a href="https://www.thelancet.com/journals/lanhl/article/PIIS2666-7568(22)00027-7/fulltext" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Alternative Link</a>

        **Study Design:** 2 × 2 factorial cluster-randomized controlled trial

        **Interventions Compared:**
        - **Group Music Therapy (GMT)** - delivered by music therapists
        -  **Recreational Choir Singing (RCS)** - delivered by community musicians
        -  **GMT + RCS combined**
        -  **Standard care**

        **Key Findings:**
        -  **RCS (singing)** reduced depressive symptoms at end of 6-month intervention
        - **Positive effects** on neuropsychiatric symptoms and quality of life
        -  **Singing effects on depression** were sustained long-term (12 months)
        - **More intensive and standardized** than previous studies

        **Significance:** First large-scale trial directly comparing music therapy vs. recreational singing, showing singing has durable benefits.
        """, unsafe_allow_html=True)

    st.markdown("---")

    st.markdown("### Music Therapy for ADHD: Key Research Papers")

    with st.expander(" **Modulation in background music influences sustained attention** (Woods et al., 2019)", expanded=False):
        st.markdown("""
        <a href="https://arxiv.org/abs/1907.06909" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Large-scale study (N=677) examining how amplitude modulation in music affects sustained attention

        **Key Finding:** Music with **16 Hz beta-band modulation** significantly improved sustained attention performance. Participants with higher ADHD traits benefited more from stronger modulation depths.

        **Application:** This is the most rigorous experimental evidence for specific frequency modulation in ADHD. The 16 Hz modulation rate works by entraining brain oscillations in the beta band, which is associated with active concentration and alertness.
        """, unsafe_allow_html=True)

    with st.expander(" **The Effect of 40 Hz Binaural Beats on Working Memory** (Wang, Zhang, Li & Yang, 2022)", expanded=False):
        st.markdown("""
        <a href="https://ieeexplore.ieee.org/document/9802990" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Controlled experiment with 40 healthy volunteers using EEG monitoring

        **Key Findings:**
        - **40 Hz binaural beats** significantly improved working memory task performance
        -  Induced "frequency-following response" where brain waves synchronized to the 40 Hz stimulus
        - Increased Higuchi fractal dimension (HFD) in temporal and parietal lobes, indicating enhanced neural complexity
        -  Increased duration and coverage of EEG microstate D (associated with attention networks)
        - Decreased microstate A (associated with mind-wandering)

        **Application for ADHD:** Working memory deficits are core features of ADHD. The 40 Hz gamma frequency appears to entrain brain oscillations in the gamma band, which is crucial for cognitive processing, attention, and memory consolidation.
        """, unsafe_allow_html=True)

    with st.expander(" **Pilot add-on Randomized-Controlled Trial evaluating binaural beats in adult ADHD** (Malandrone et al., 2022)", expanded=False):
        st.markdown("""
        <a href="https://pmc.ncbi.nlm.nih.gov/articles/PMC9564012/" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>

        **Study Design:** Randomized controlled trial with adult ADHD patients
        **Frequency Used:** **15 Hz binaural beats** (beta range: 415 Hz right ear, 400 Hz left ear)

        **Key Findings:**
        -  **Significant improvement** in subjective studying performance (mean difference=2.7, p<0.001)
        -  **Improvements maintained** across fortnightly follow-ups
        -  No significant changes in standardized ADHD rating scales, but subjective performance clearly improved

        **Application for ADHD:** This is the only RCT specifically testing binaural beats in diagnosed ADHD patients. The 15 Hz beta frequency appears to help with focus during studying tasks. Beta frequencies (13-30 Hz) are associated with active concentration and alert states.
        """, unsafe_allow_html=True)

    st.markdown("---")

    st.markdown("###  Music Therapy for Down Syndrome: Key Research Papers")

    with st.expander(" **40Hz Sensory Stimulation Improves Cognition in Down Syndrome Mice** (Islam, Jackson et al., 2025)", expanded=False):
        st.markdown("""
        <a href="https://www.technologynetworks.com/genomics/news/40hz-sensory-stimulation-improves-cognition-in-down-syndrome-mice-398950" target="_blank" style="color:#000000; font-weight:700; text-decoration: underline;">Link to Study</a>
        **[Full Research](https://picower.mit.edu/news/down-syndrome-mice-40hz-light-and-sound-improve-cognition-neurogenesis-connectivity)**

        **Study Design:** Preclinical study using Ts65Dn mouse model of Down syndrome at MIT Picower Institute
        
        **Intervention:** 40 Hz combined light and sound stimulation (GENUS), 1 hour daily for 3 weeks

        **Key Findings:**
        -  **Improved performance** on three different short-term memory tasks
        -  **Enhanced hippocampal activity** and connectivity
        -  **Promoted neurogenesis** (growth of new neurons)
        -  **Increased indicators** of neural activity in the hippocampus

        **Application for Down Syndrome:** This is groundbreaking research showing 40 Hz stimulation specifically helps cognitive deficits in Down syndrome. A human clinical trial is currently underway at MIT to test this in people with Down syndrome.
        """)

    with st.expander(" **Effect of 528 Hz Music on the Endocrine System and Autonomic Nervous System** (Akimoto et al., 2018)", expanded=False):
        st.markdown("""
        <a href=\"https://www.scirp.org/journal/paperinformation?paperid=87146\" target=\"_blank\" style=\"color:#000000; font-weight:700; text-decoration: underline;\">Link to Study</a>

        **Study Design:** Controlled study with 9 healthy participants comparing 528 Hz vs 440 Hz music

        **Intervention:** 5 minutes of soothing piano music at each frequency

        **Key Findings (528 Hz condition):**
        -  **Cortisol levels significantly decreased** (stress reduction)
        -  **Chromogranin A tended to decrease** (stress marker)
        - **Oxytocin significantly increased** (bonding/happiness hormone)
        -  **Tension-anxiety scores significantly reduced** (p<0.0091)
        - **Total Mood Disturbance scores significantly reduced** (p<0.0487)
        -  **Autonomic nervous system** showed relaxation responses

        -  **440 Hz condition**: No significant changes in any biomarkers

        **Application:** This provides the strongest scientific evidence for 528 Hz as a healing frequency. The study used objective biomarkers rather than just subjective reports. The effect occurred after only 5 minutes, suggesting powerful physiological impact.
        """, unsafe_allow_html=True)

    with st.expander(" **Effect of Music Engagement and Movement on Children with Down Syndrome** (Orff-Schulwerk Study, 2024)", expanded=False):
        st.markdown("""
        <a href=\"https://archive.conscientiabeam.com/index.php/61/article/download/3626/7913\" target=\"_blank\" style=\"color:#000000; font-weight:700; text-decoration: underline;\">Link to Study</a>

        **Study Design:** Experimental study with children aged 7-10 with Down syndrome

        **Intervention:** 8-week Orff-based Music Engagement and Movement (MEM) program

        **Key Findings:**
        - **Significantly improved attention span** in Down syndrome group compared to controls
        -  **Significantly improved memory retention**
        -  **Orff-Schulwerk approach** (rhythm, movement, singing, instruments) particularly effective

        **Application for Down Syndrome:** Demonstrates that active music participation (not just listening) with rhythmic and movement components produces measurable cognitive improvements in children with Down syndrome.
        """, unsafe_allow_html=True)

    st.markdown("---")

    st.markdown("###  Big Five Personality Traits and Music Preferences: Key Research Papers")

    with st.expander("**Associations Between Personality Traits and Music Preference** (Boccia, 2020)", expanded=False):
        st.markdown("""
        <a href=\"https://digitalcommons.lindenwood.edu/cgi/viewcontent.cgi?article=1000&context=psych_journals\" target=\"_blank\" style=\"color:#000000; font-weight:700; text-decoration: underline;\">Link to Study</a>

        **Study Design:** Empirical study with 175 participants
        **Method:** Participants completed Big Five personality assessment and rated enjoyment of 15-second audio clips from 7 genres

        **Key Findings - Overall Music Enjoyment:**
        -  **Openness**: Significant positive correlation (r=.169, p=.013) - high Openness = enjoy MORE genres
        -  **Extraversion**: Strong positive correlation (r=.275, p<.001) - high Extraversion = enjoy MORE genres
        - **Agreeableness**: Positive correlation (r=.233, p=.001) - high Agreeableness = enjoy MORE genres
        - **Neuroticism**: Significant negative correlation (r=−.186, p=.007) - LOW Neuroticism = enjoy MORE genres
        - **Conscientiousness**: No significant relationship (r=.061, p=.212) - no difference in genre enjoyment

        **Most Preferred Genres Overall:** Pop (130 participants), Rock (91), Rap (82)
        """, unsafe_allow_html=True)

    with st.expander(" **The Structure of Musical Preferences: A Five-Factor Model** (Rentfrow, Goldberg & Levitin, 2011)", expanded=False):
        st.markdown("""
        <a href=\"https://pmc.ncbi.nlm.nih.gov/articles/PMC3138530/\" target=\"_blank\" style=\"color:#000000; font-weight:700; text-decoration: underline;\">Link to Study</a>

        **Study Design:** Multi-study factor analysis developing the MUSIC model

        **Major Contribution:** Moved beyond genre-based assessment to use actual music excerpts

        **Five Music Preference Factors Identified:**
        -  **Mellow**: Smooth and relaxing styles
        -  **Urban**: Rhythmic and percussive (rap, funk, acid jazz)
        -  **Sophisticated**: Classical, operatic, world, jazz
        - **Intense**: Loud, forceful, energetic
        -  **Campestral**: Direct, rootsy music (country, singer-songwriter)
        """, unsafe_allow_html=True)

    with st.expander(" **Nirvana for Neuroticism: How Musical Preferences Match Personality Traits Around the World** (Greenberg et al., 2022)", expanded=False):
        st.markdown("""
        <a href=\"https://www.technologynetworks.com/neuroscience/news/nirvana-for-neuroticism-how-musical-preferences-match-personality-traits-around-the-world-358429\" target=\"_blank\" style=\"color:#000000; font-weight:700; text-decoration: underline;\">Link to Study</a>

        **Study Design:** Large-scale international study (350,000+ participants from 50+ countries)

        **Key Findings - Universal Patterns:**
        -  **Extraversion**: Correlated with Contemporary/upbeat music (Ed Sheeran's "Shivers") - particularly strong around equator and Central/South America
        -  **Openness**: Correlated with "Sophisticated music" (David Bowie's "Space Oddity", Nina Simone)
        -  **Agreeableness**: Correlated with upbeat/conventional music (Marvin Gaye's "What's Going On", Lady Gaga's "Shallow")
        -  **Conscientiousness**: Negative correlation with intense music (unlikely to enjoy Rage Against the Machine)
        -  **Neuroticism**: Correlated with **Intense musical styles** (Nirvana's "Smells Like Teen Spirit"), reflecting "inner angst and frustration"

        **Surprising Finding:** Neuroticism did not clearly favor either sad music (for catharsis) or upbeat music (mood shift), but instead preferred intense styles.
        """, unsafe_allow_html=True)

    st.markdown("---")

    st.markdown("###  Integrated Model: How to Apply This Research")

    with st.expander(" **Application to Dementia Music Therapy**", expanded=False):
        st.markdown("""
        ** Individualized vs. Generic Music:**
        -  Personalization matters, BUT account for depression levels, apathy, and cognitive impairment severity
        -  Someone with high depression might experience increased sadness even with their favorite music

        ** Active vs. Passive Music Engagement:**
        -  Active music-making (singing, playing instruments) is more effective than passive listening
        -  Especially for depression - therapist-delivered active therapy more effective than music listening

        **Autobiographical Memory Activation:**
        -  Music-evoked autobiographical memories (MEAMs) remain intact in AD
        -  Even unfamiliar music (Vivaldi's "Spring") enhanced memory recall and reduced trait anxiety

        ** Caregiver-Delivered Interventions:**
        -  Family caregivers can deliver effective music interventions at home
        -  Mechanisms: neural entrainment, auditory-motor coupling, activating arousal-mood pathways
        """)

    with st.expander("**Application to Personality-Matched Music Selection**", expanded=False):
        st.markdown("""
        ** Openness - The "Omnivore" Trait:**
        -  Offer diverse genres, sophisticated music, world music, jazz, classical, experimental styles
        -  High Openness predicts enjoying MORE genres overall (r=.169-.25)

        ** Extraversion - The "Energetic/Social" Trait:**
        -  Provide contemporary, rhythmic, danceable, upbeat music (pop, electronic, country)
        -  Extraversion strongly predicts enjoyment of contemporary, upbeat music (r=.275)

        ** Agreeableness - The "Harmonious" Trait:**
        -  Select positive-valence, conventional, communal music (pop, country, soul)
        -  Agreeable individuals prefer upbeat, conventional, positive-valence music

        ** Neuroticism - The "Intense/Anxious" Trait:**
        -  **For catharsis**: Intense, emotionally expressive music (alternative, intense rock)
        -  **For relaxation**: Monitor carefully—high-neurotic individuals may avoid mellow music
        -  **Use physiological measures** (heart rate, cortisol) not just self-reports

        ** Conscientiousness - The "Neutral" Trait:**
        -  Focus on lyrics/meaning rather than sound characteristics
        -  Relatively neutral for music enjoyment
        """)

    st.markdown("""
    <div style='background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%);
                padding: 1.5rem;
                border-radius: 12px;
                border-left: 4px solid #A3B18A;
                margin: 1.5rem 0;'>
        <h3 style='color: #364153; margin-top: 0; margin-bottom: 1rem; font-weight: 600;'>Research Summary</h3>
        <p style='color: style='color:#364153;'; line-height: 1.6; margin: 0;'>
            The research evidence demonstrates <strong style='color: #A3B18A;'>consistent beneficial effects</strong> of music therapy across all three populations studied.
            For dementia patients, music therapy <strong style='color: #588157;'>improves cognitive function, reduces distress, and enhances quality of life</strong>.
            For individuals with ADHD, music therapy <strong style='color: #364153;'>shows promise in improving attention, reducing hyperactivity, and enhancing social skills</strong>.
            For children with Down syndrome, music therapy <strong style='color: #A3B18A;'>significantly improves socialization, communication, and cognitive development</strong>.
            Regarding personality and music preferences, research consistently shows that the Big Five traits predict musical preferences, with <strong style='color: #588157;'>Openness and Extraversion</strong> showing the strongest correlations with diverse musical tastes, while <strong style='color: #364153;'>Conscientiousness and Neuroticism</strong> are associated with more specific genre preferences.
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='color:#364153; font-size: 0.85rem; margin: 1rem 0; font-style: italic;'>
        This evidence base, comprising systematic reviews, meta-analyses, randomized controlled trials, and large-scale observational studies,
        provides strong support for the therapeutic applications of music across these populations and confirms the relationship between
        personality traits and musical preferences.
    </div>
    """, unsafe_allow_html=True)

    # Research Links Section
    st.markdown("""
    <div style='margin: 2rem 0;'>
        <h3 style='color: #364153; margin-bottom: 1rem; font-size: 1.5rem;'>🔗 Research Paper Links</h3>
        <p style='color: #A3B18A; font-size: 1.1rem; margin-bottom: 1rem; font-weight: 500;'>
            Access to full research papers and additional resources:
        </p>
    </div>
    """, unsafe_allow_html=True)

    with st.expander(" View Research Links", expanded=False):
        research_links = [
            "https://www.frontiersin.org/journals/medicine/articles/10.3389/fmed.2020.00160/full",
            "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1618324/full",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12307461/",
            "https://www.nature.com/articles/s44220-024-00342-x",
            "https://jpalliativecare.com/effect-of-music-therapy-on-quality-of-life-in-geriatric-population-a-systematic-review-and-meta-analysis/",
            "https://pubmed.ncbi.nlm.nih.gov/40680190/",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC10221503/",
            "https://journals.sapienzaeditorial.com/index.php/SIJIS/article/view/e25044",
            "https://journals.plos.org/plosone/article?id=10.1371%2Fjournal.pone.0324369",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12316199/",
            "https://archive.conscientiabeam.com/index.php/61/article/download/3626/7913",
            "https://www.scholarlyreview.org/api/v1/articles/121698-music-as-a-language-assessing-the-extent-to-which-active-music-therapy-promotes-socialization-development-for-children-under-12-with-down-syndrome.pdf",
            "https://journals.sagepub.com/doi/abs/10.1080/03080188.2020.1755556",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12191269/",
            "https://digitalcommons.lindenwood.edu/cgi/viewcontent.cgi?article=1000&context=psych_journals",
            "https://www.cp.jku.at/people/schedl/Research/Publications/pdf/ferwerda_umap_2017.pdf",
            "https://www.davidmgreenberg.com/wp-content/uploads/2018/11/Nave-et-al-2018-music-preferences-from-fb-likes.pdf",
            "https://www.cam.ac.uk/stories/musical-preferences-unite-personalities-worldwide",
            "https://research.atspotify.com/just-the-way-you-are-music-listening-and-personality",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC3138530/",
            "https://ir.library.illinoisstate.edu/cgi/viewcontent.cgi?article=2574&context=etd",
            "https://www.pioneerpublisher.com/jrssh/article/download/1045/948/1097",
            "https://www.nature.com/articles/s41598-025-95661-z",
            "https://journals.sagepub.com/doi/10.1177/03057356221135355",
            "https://journals.sagepub.com/doi/pdf/10.1177/0305735616658957",
            "https://www.nature.com/articles/s41598-025-93795-8",
            "https://www.sciencedirect.com/science/article/abs/pii/S0190740925001343",
            "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2022.905113/full",
            "https://journals.sagepub.com/doi/10.1177/23320249241265240",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC3280156/",
            "https://www.auctoresonline.org/article/use-and-effectiveness-of-musical-social-story-therapy-in-children-with-developmental-disorders-down-syndrome-autism-spectrum-disorder-fragile-x-syndrome-fetal-alcohol-spectrum-disorder-cerebral-palsy-and-adhd",
            "https://baylor-ir.tdl.org/bitstreams/81c66bd2-202f-4051-b3b5-3aaee0e48fa9/download",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12235852/",
            "https://files.eric.ed.gov/fulltext/EJ976663.pdf",
            "https://www.sciencedirect.com/science/article/pii/S002239562200231X",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC6481398/",
            "https://www.pagepress.org/medicine/gimle/article/view/595"
        ]

        # Display all research links as black colored text
        st.markdown("""
        <div style='color: #000000; line-height: 1.8;'>
        """ + "<br>".join([f"• <a href='{link}' target='_blank' style='color: #000000; text-decoration: underline;'>{link}</a>" for link in research_links]) + """
        </div>
        """, unsafe_allow_html=True)

def page_about():
    """About page with premium design"""
    render_logo()

    st.markdown("## About TheramuseRX")

    st.markdown("""
    **TheramuseRX is an our own AI powered model music therapy recommendation system**
    designed to provide personalized and therapeutically relevant music selections from YouTube with
    **automatic reinforcement learning**.
    """)

    st.markdown("### Key Features")

    features = [
        "**Condition-Specific Therapy:** Dedicated modules for Dementia, Down Syndrome, and ADHD",
        "**Big-5 Personality:** Personality-based music genre recommendations",
        "**Reinforcement Learnin:** Reinforcement learning algorithm for personalized learning",

    ]

    for feature in features:
        st.markdown(f"• {feature}")

    st.markdown("### Therapy Modules")

    with st.container():
        st.markdown("#### Developed for:")
        dementia_features = [
            "**Working on:** Dementia, Down Sysndrome,ADHD",
            "**Memory Enhancement:** Music to stimulate memory recall and cognitive function",
            "**Sleep Improvement:** Calming tracks to aid sleep quality",
            "**Emotional Well-being:** Uplifting music to enhance mood and reduce anxiety",
            "**Personalized Selection:** Incorporates patient's birthplace, age, and preferences"

        ]
        for feature in dementia_features:
            st.markdown(f"  - {feature}")



    # st.markdown("### Reinforcement Learning")

    # rl_features = [
    #     "**Linear Thompson Sampling:** Bayesian approach for exploration-exploitation",
    #     "**Contextual Features:** 20-dimensional feature space",
    #     "**Continuous Learning:** Updates with every feedback submission",
    #     "**Non-stationary:** Decay factor for adapting to changing preferences"
    # ]

    # for feature in rl_features:
    #     st.markdown(f"• {feature}")
    st.markdown("### Developer")
    st.markdown("**TheramuseRX-Ashiq Sazid**")
 
# MAIN APP

def main():
    """Main application"""
    
    # Sidebar navigation
    # Sidebar navigation
    # Sidebar navigation with logo
    # st.sidebar.image("b.png", width=200)
    # st.sidebar.markdown("<h1 style='text-align: center; margin-bottom: 2rem; color:  #338AFF !important;'></h1>", unsafe_allow_html=True)

    pages = {
    " Patient Intake": page_intake,
    "Patient Database": page_patient_database,
    " Analytics": page_analytics,
    "Research Evidence": page_research_evidence,
    "About": page_about
}

    # Always show the sidebar navigation
    selection = st.sidebar.radio("Navigation", list(pages.keys()))

  
    # Check for session state navigation override
    if 'page' in st.session_state and st.session_state.page in pages:
        selection = st.session_state.page
        # Clear the session state after using it
        del st.session_state.page

    # Centered image at the top of the main screen
    logo_col_left, logo_col_center, logo_col_right = st.columns([2, 1, 2])
    with logo_col_center:
        st.image("b.png", width=500)

    # Scroll to top if flag is set (after generating results)
    if 'scroll_to_top' in st.session_state and st.session_state.scroll_to_top:
        st.markdown("""
        <script>
            // Scroll to top immediately and also after a short delay
            window.scrollTo(0, 0);
            setTimeout(() => {
                window.scrollTo(0, 0);
            }, 100);
        </script>
        """, unsafe_allow_html=True)
        # Clear the flag
        st.session_state.scroll_to_top = False

    # Scroll to Personalized Recommendations if flag is set
    if '_scroll_to_recommendations_flag' in st.session_state and st.session_state._scroll_to_recommendations_flag:
        st.markdown("""
        <script>
            // Multiple attempts to scroll to the Personalized Music Recommendations section
            function scrollToRecommendations() {
                console.log('Attempting to scroll to personalized-recommendations...');
                const element = document.getElementById('personalized-recommendations');
                if (element) {
                    console.log('Found element, scrolling...');
                    element.scrollIntoView({ behavior: 'smooth', block: 'start' });
                    return true;
                } else {
                    console.log('Element not found, trying alternative approach...');
                    // Try finding by text content
                    const headers = document.querySelectorAll('h2');
                    for (let h2 of headers) {
                        if (h2.textContent && h2.textContent.includes('Personalized Music Recommendations')) {
                            console.log('Found by text content, scrolling...');
                            h2.scrollIntoView({ behavior: 'smooth', block: 'start' });
                            return true;
                        }
                    }
                }
                return false;
            }

            // Try immediately
            if (!scrollToRecommendations()) {
                // Try after 100ms
                setTimeout(() => scrollToRecommendations(), 100);
            }

            // Try after 300ms
            setTimeout(() => scrollToRecommendations(), 300);

            // Fallback: scroll to top if element not found after 500ms
            setTimeout(() => {
                if (!scrollToRecommendations()) {
                    console.log('Fallback: scrolling to top');
                    window.scrollTo({ top: 0, behavior: 'smooth' });
                }
            }, 500);
        </script>
        """, unsafe_allow_html=True)
        # Clear the flag
        st.session_state._scroll_to_recommendations_flag = False
        # Clear URL parameter
        if 'scroll' in st.query_params:
            del st.query_params['scroll']

    # Display selected page
    pages[selection]()
    
    # Footer
    st.sidebar.markdown("---")
    st.sidebar.markdown("")
    st.sidebar.markdown("")
    st.sidebar.markdown("")
    st.sidebar.markdown("")
    st.sidebar.markdown("")
    st.sidebar.markdown("""
     <style>
    /* Ensure sidebar layout allows absolute positioning */
    [data-testid="stSidebar"] {
        position: relative !important;
        padding-bottom: 0 !important;
    }

    /* Fix the footer firmly at the bottom edge */
    .sidebar-footer {
        position: fixed;
        bottom: 0.2 rem;
        left: 0;
        width: inherit;
        text-align: center;
        padding: 1rem 1rem 1.5rem 1rem;
        background-color: #338AFF; /* match sidebar color if needed */
    }

    .sidebar-footer p {
        margin: 0;
        line-height: 1.3;
    }
    </style>

    <div class='sidebar-footer'>
        <p style='font-weight:700; font-size:1.1rem; color:#FDFBF7; margin-bottom:0.5rem;'>
            “Where words fail, music speaks — not to the mind, but to the soul.”
        </p>
        <p style='color:rgba(255,255,255,0.7); font-size:0.9rem;'>
            — Hans Christian Andersen
        </p>
    </div>
    
    """, unsafe_allow_html=True)
    
    # System status in sidebar

if __name__ == "__main__":
    main()

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import date, datetime
from pathlib import Path
import json
from typing import Dict, List, Optional, Tuple
import re
import io
import base64
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import qn
import toml
import sqlite3
import os

# Load color schema from config
def load_color_schema():
    """Load color schema from config.toml file"""
    try:
        config_path = Path(__file__).parent / ".streamlit" / "config.toml"
        if config_path.exists():
            config = toml.load(config_path)
            return config.get('color_schema', {})
    except Exception as e:
        st.warning(f"Could not load color schema: {e}")

    # Fallback colors if config fails to load
    return {
        'slider_primary': '#a8b894',
        'slider_secondary': '#e8ddb5',
        'slider_track': '#5B9C96',
        'table_header': '#e8ddb5',
        'table_cells': '#a8b894',
        'table_border': '#d4c4a8',
        'table_hover': '#A3B18A',
        'spider_primary': '#5B9C96',
        'spider_secondary': '#a8b894',
        'spider_fill': '#e8ddb580',
        'spider_text': '#000000',
        'spider_grid': '#A3B18A50',
        'accent_teal': '#5B9C96',
        'accent_olive': '#a8b894',
        'accent_beige': '#e8ddb5',
        'accent_dark': '#344E41'
    }

# Import from main.py
from ml import (
    TheraMuse, 
    DementiaTherapy, 
    DownSyndromeTherapy, 
    ADHDTherapy,
    YouTubeAPI,
    BangladeshiGenerationalMatrix,
    BigFivePersonalityMapping,
    LinearThompsonSampling,
    DatabaseManager
)

# PAGE CONFIG (MUST BE FIRST STREAMLIT COMMAND)
LOGO_PATH = next((p for p in [Path("p.png"), Path.cwd() / "p.png"] if p.exists()), None)

st.set_page_config(
    page_title="TheraMuse - Music Therapy",
    page_icon=str(LOGO_PATH) if LOGO_PATH else "",
    layout="wide",
    initial_sidebar_state="expanded"
)
##########

st.markdown("""
<style>
/* Column headers */
[data-testid="stDataFrame"] [role="columnheader"] {
    background-color: #a8b894 !important;   /* same as your background */
    color: #e8dfb0 !important;              /* warm beige header text */
    font-weight: 600 !important;
    border: none !important;
}

/* Table cells */
[data-testid="stDataFrame"] [role="gridcell"] {
    background-color: #a8b894 !important;   /* same olive background */
    color: #e8dfb0 !important;              /* beige text for cells */
    border: none !important;
}

/* Rounded corners and hover highlight */
[data-testid="stDataFrame"] {
    border-radius: 12px !important;
    overflow: hidden !important;
}
[data-testid="stDataFrame"] [role="row"]:hover [role="gridcell"] {
    background-color: #b7c3a2 !important;   /* gentle hover highlight */
}
</style>

<script>
// Prevent Enter key from submitting forms in text inputs
document.addEventListener('DOMContentLoaded', function() {
    const textInputs = document.querySelectorAll('input[type="text"]');
    textInputs.forEach(function(input) {
        input.addEventListener('keydown', function(event) {
            if (event.key === 'Enter') {
                event.preventDefault();
                event.stopPropagation();
                // Just move focus out of the input field
                this.blur();
                return false;
            }
        });
    });
});
</script>
""", unsafe_allow_html=True)



#STYLING
st.markdown("""
<style>
    /* Global Premium Styling */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    .main {
        background: linear-gradient(135deg, #0a0a0a 0%, #1a1a1a 100%);
        animation: fadeIn 0.8s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    /* Premium Header */
    .main-header {
        font-size: 3.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #A3B18A 0%, #E8DDB5 50%, #A3B18A 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-align: center;
        margin: 2rem 0;
        letter-spacing: -0.02em;
        animation: shimmer 3s ease-in-out infinite;
    }
    
    @keyframes shimmer {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    /* Glassmorphism Card */
    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px) saturate(180%);
        -webkit-backdrop-filter: blur(20px) saturate(180%);
        border-radius: 20px;
        border: 1px solid rgba(255, 255, 255, 0.08);
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 
            0 8px 32px 0 rgba(0, 0, 0, 0.37),
            inset 0 1px 0 0 rgba(255, 255, 255, 0.05);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: slideUp 0.6s ease-out;
    }
    
    @keyframes slideUp {
        from { opacity: 0; transform: translateY(30px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .glass-card:hover {
        transform: translateY(-8px);
        box-shadow: 
            0 16px 48px 0 rgba(163, 177, 138, 0.2),
            inset 0 1px 0 0 rgba(255, 255, 255, 0.1);
        border-color: rgba(163, 177, 138, 0.3);
    }
    
    /* Premium Song Card */
    .song-card {
        background: linear-gradient(135deg, rgba(26, 26, 26, 0.95) 0%, rgba(20, 20, 20, 0.95) 100%);
        backdrop-filter: blur(20px);
        padding: 1.8rem;
        border-radius: 16px;
        border-left: 4px solid transparent;
        border-image: linear-gradient(180deg, #A3B18A 0%, #588157 100%);
        border-image-slice: 1;
        margin: 1rem 0;
        position: relative;
        overflow: hidden;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: fadeInCard 0.5s ease-out backwards;
    }
    
    @keyframes fadeInCard {
        from { 
            opacity: 0; 
            transform: translateX(-20px);
        }
        to { 
            opacity: 1; 
            transform: translateX(0);
        }
    }
    
    .song-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.05) 0%, transparent 100%);
        opacity: 0;
        transition: opacity 0.4s ease;
        pointer-events: none; /* allow clicks on links/buttons above overlay */
    }
    
    .song-card:hover {
        transform: translateX(8px) scale(1.02);
        border-left-width: 6px;
        box-shadow: 
            -4px 0 24px rgba(163, 177, 138, 0.3),
            0 8px 32px rgba(0, 0, 0, 0.4);
    }
    
    .song-card:hover::before {
        opacity: 1;
    }
    
    .song-card h4 {
        color: #FFFFFF;
        font-size: 1.25rem;
        font-weight: 600;
        margin-bottom: 0.75rem;
        letter-spacing: -0.01em;
        transition: color 0.3s ease;
    }
    
    .song-card:hover h4 {
        color: #E8DDB5;
    }
    
    .song-card p {
        color: rgba(255, 255, 255, 0.7);
        font-size: 0.95rem;
        line-height: 1.6;
        margin: 0.5rem 0;
    }
    
    .song-card strong {
        color: #A3B18A;
        font-weight: 600;
    }
    .song-card a { pointer-events: auto; }
    
    /* YouTube Link Styling */
    .youtube-link {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.75rem 1.5rem;
        background: linear-gradient(135deg, #FF0000 0%, #CC0000 100%);
        color: white;
        text-decoration: none;
        border-radius: 12px;
        font-weight: 600;
        font-size: 0.95rem;
        margin-top: 1rem;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 16px rgba(255, 0, 0, 0.3);
        position: relative;
        z-index: 3; /* ensure above decorative overlays */
        cursor: pointer;
    }
    
    .youtube-link:hover {
        transform: translateY(-2px) scale(1.05);
        box-shadow: 0 8px 24px rgba(255, 0, 0, 0.5);
        background: linear-gradient(135deg, #FF1a1a 0%, #DD0000 100%);
        text-decoration: none;
        color: white;
    }
    
    .youtube-link::before {
        content: '▶';
        font-size: 1.1rem;
    }
    
    /* Embedded YouTube Player Container */
    .youtube-embed-container {
        position: relative;
        width: 100%;
        padding-top: 56.25%; /* 16:9 Aspect Ratio */
        margin-top: 1rem;
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.5);
        transition: all 0.3s ease;
    }
    
    .youtube-embed-container:hover {
        transform: scale(1.02);
        box-shadow: 0 12px 48px rgba(163, 177, 138, 0.3);
    }
    
    .youtube-embed-container iframe {
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        border: none;
        border-radius: 12px;
    }
    
    /* Sidebar Navigation Colors */
    [data-testid="stSidebar"] [data-testid="stRadio"] label p {
        color: #E8DDB5 !important;
    }
    [data-testid="stSidebar"] [data-testid="stRadio"] label {
        color: #E8DDB5 !important;
    }
    
    /* Metric Cards */
    .metric-card {
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.15) 0%, rgba(88, 129, 87, 0.15) 100%);
        backdrop-filter: blur(20px);
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        border: 1px solid rgba(163, 177, 138, 0.2);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        animation: scaleIn 0.5s ease-out backwards;
    }
    
    @keyframes scaleIn {
        from { 
            opacity: 0; 
            transform: scale(0.9);
        }
        to { 
            opacity: 1; 
            transform: scale(1);
        }
    }
    
    .metric-card:hover {
        transform: translateY(-8px) scale(1.05);
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.25) 0%, rgba(88, 129, 87, 0.25) 100%);
        box-shadow: 0 16px 48px rgba(163, 177, 138, 0.3);
        border-color: rgba(163, 177, 138, 0.4);
    }
    
    .metric-card h2, .metric-card h3 {
        color: #E8DDB5;
        margin: 0.5rem 0;
    }
    
    /* Premium Buttons - Enhanced for better functionality */
    .stButton>button {
        background: linear-gradient(135deg, #A3B18A 0%, #588157 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        letter-spacing: 0.02em !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 16px rgba(163, 177, 138, 0.3) !important;
        cursor: pointer !important;
        pointer-events: auto !important;
        position: relative !important;
        z-index: 10 !important;
    }

    .stButton>button:hover {
        background: linear-gradient(135deg, #588157 0%, #3a5a40 100%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 24px rgba(163, 177, 138, 0.5) !important;
    }

    .stButton>button:active {
        transform: translateY(0) !important;
        box-shadow: 0 4px 16px rgba(163, 177, 138, 0.3) !important;
    }

    /* Ensure buttons are clickable and not blocked */
    .stButton {
        pointer-events: auto !important;
        z-index: 10 !important;
        position: relative !important;
    }

    /* Fix for buttons in containers */
    div.stButton > button {
        pointer-events: auto !important;
        z-index: 10 !important;
        position: relative !important;
    }
    
    /* Feedback Buttons */
    .feedback-button {
        width: 100%;
        margin: 0.3rem 0;
    }
    
    /* Expander Premium Style */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.1) 0%, rgba(88, 129, 87, 0.1) 100%);
        border-radius: 12px;
        border: 1px solid rgba(163, 177, 138, 0.2);
        transition: all 0.3s ease;
    }
    
    .streamlit-expanderHeader:hover {
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%);
        border-color: rgba(163, 177, 138, 0.4);
    }
    
    /* Category Badge */
    .category-badge {
        display: inline-block;
        padding: 0.5rem 1rem;
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%);
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        color: #E8DDB5;
        margin: 0.5rem 0;
        border: 1px solid rgba(163, 177, 138, 0.3);
    }
    
    /* Rank Badge */
    .rank-badge {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 32px;
        height: 32px;
        background: linear-gradient(135deg, #A3B18A 0%, #588157 100%);
        border-radius: 50%;
        font-weight: 700;
        font-size: 0.9rem;
        color: white;
        margin-right: 0.75rem;
        box-shadow: 0 2px 8px rgba(163, 177, 138, 0.4);
    }
    
    /* Smooth Scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.03);
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, #A3B18A 0%, #588157 100%);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(180deg, #588157 0%, #3a5a40 100%);
    }
    
    /* Tab Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: rgba(255, 255, 255, 0.03);
        border-radius: 12px;
        padding: 0.75rem 1.5rem;
        color: rgba(255, 255, 255, 0.7);
        transition: all 0.3s ease;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: rgba(163, 177, 138, 0.1);
        color: #E8DDB5;
        border-color: rgba(163, 177, 138, 0.3);
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%) !important;
        color: #E8DDB5 !important;
        border-color: rgba(163, 177, 138, 0.4) !important;
    }
    
    /* Form Elements */
    .stTextInput>div>div>input,
    .stSelectbox>div>div>select {
        background: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        color: white !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextInput>div>div>input:focus,
    .stSelectbox>div>div>select:focus {
        border-color: rgba(163, 177, 138, 0.5) !important;
        box-shadow: 0 0 0 2px rgba(163, 177, 138, 0.2) !important;
    }

    /* Enhanced Dynamic Slider Design */
    .stSlider label {
        color: #000000 !important;
        font-weight: 700 !important;
        font-size: 1.1rem !important;
        margin-bottom: 15px !important;
    }

    /* Default slider track with animated gradient */
    .stSlider > div[data-baseweb="slider"] > div > div {
        background: linear-gradient(90deg,
            #FF6B6B 0%,
            #FFE66D 25%,
            #95E77E 50%,
            #4ECDC4 75%,
            #45B7D1 100%) !important;
        height: 8px !important;
        border-radius: 4px !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.15) !important;
        transition: all 0.3s ease !important;
    }

    /* Enhanced slider thumb with hover effects */
    .stSlider > div[data-baseweb="slider"] > div > div > div[role="slider"] {
        background: linear-gradient(135deg, #a8b894 0%, #5B9C96 100%) !important;
        border: 3px solid #FFFFFF !important;
        box-shadow: 0 2px 12px rgba(0,0,0,0.3) !important;
        transition: all 0.3s ease !important;
        width: 24px !important;
        height: 24px !important;
    }

    .stSlider > div[data-baseweb="slider"] > div > div > div[role="slider"]:hover {
        transform: scale(1.15) !important;
        box-shadow: 0 4px 16px rgba(0,0,0,0.4) !important;
        background: linear-gradient(135deg, #5B9C96 0%, #a8b894 100%) !important;
    }

    /* Slider value text with enhanced styling */
    .stSlider > div[data-baseweb="slider"] > div > div > div {
        color: #000000 !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
    }

    /* Animated slider container - fixed to not interfere with other elements */
    .stSlider > div[data-baseweb="slider"] {
        transition: all 0.3s ease !important;
        pointer-events: auto !important;
        position: relative !important;
        z-index: 1 !important;
    }

    .stSlider > div[data-baseweb="slider"]:hover {
        transform: translateY(-2px) !important;
    }

    /* Pulse animation for focused sliders - contained */
    @keyframes sliderPulse {
        0% { box-shadow: 0 0 0 0 rgba(91, 156, 150, 0.7); }
        70% { box-shadow: 0 0 0 10px rgba(91, 156, 150, 0); }
        100% { box-shadow: 0 0 0 0 rgba(91, 156, 150, 0); }
    }

    .stSlider > div[data-baseweb="slider"]:focus-within {
        animation: sliderPulse 2s infinite !important;
        transform: none !important; /* Prevent transform conflicts */
    }

    /* Success/Error Messages */
    .stSuccess, .stError, .stInfo, .stWarning {
        border-radius: 12px;
        backdrop-filter: blur(20px);
        animation: slideInRight 0.4s ease-out;
    }
    
    @keyframes slideInRight {
        from { 
            opacity: 0; 
            transform: translateX(20px);
        }
        to { 
            opacity: 1; 
            transform: translateX(0);
        }
    }
    
    /* Loading Spinner */
    .stSpinner > div {
        border-color: #A3B18A transparent transparent transparent !important;
    }
    
    /* Premium Section Headers */
    h1, h2, h3 {
        color: #E8DDB5 !important;
        font-weight: 700 !important;
        letter-spacing: -0.02em !important;
    }
    
    /* Sidebar Styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(10, 10, 10, 0.95) 0%, rgba(20, 20, 20, 0.95) 100%);
        backdrop-filter: blur(20px);
    }
    
    [data-testid="stSidebar"] .stRadio > label {
        background: rgba(255, 255, 255, 0.03);
        border-radius: 12px;
        padding: 0.75rem;
        margin: 0.3rem 0;
        transition: all 0.3s ease;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }
    
    [data-testid="stSidebar"] .stRadio > label:hover {
        background: rgba(163, 177, 138, 0.1);
        border-color: rgba(163, 177, 138, 0.3);
    }
</style>
""", unsafe_allow_html=True)

# HELPER FUNCTIONS

def render_logo(size=150):
    """Render logo with premium animation"""
    if LOGO_PATH and LOGO_PATH.exists():
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.image(str(LOGO_PATH), width=size)
    else:
         st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&display=swap');

.typing-header {
    text-align: center;
    font-family: 'Playfair Display', serif;
    font-size: 3.5rem;
    color: #E8DDB5;
    letter-spacing: 2px;
    white-space: nowrap;
    overflow: hidden;
    border-right: 3px solid #E8DDB5;
    width: 0;
    margin: 2rem auto 1rem auto;
    animation: typing 3s steps(20, end) forwards, blink 0.7s step-end infinite;
}

/* typing animation */
@keyframes typing {
    from { width: 0; }
    to { width: 16ch; } /* length of "🎵 TheraMuse" */
}

/* blinking cursor */
@keyframes blink {
    50% { border-color: transparent; }
}

/* optional subtitle */
.sub-header {
    text-align: center;
    font-family: 'Poppins', sans-serif;
    font-size: 1rem;
    letter-spacing: 1px;
    color: rgba(232, 221, 181, 0.8);
    opacity: 0;
    animation: fadeIn 1s ease-in-out 3s forwards;
}

@keyframes fadeIn {
    from { opacity: 0; }
    to { opacity: 1; }
}
</style>

<h1 class='typing-header'> TheramuseRx</h1>
<p class='sub-header'> Music Therapy for Mind & Soul</p>
""", unsafe_allow_html=True)



def compute_age_from_dob(dob: date) -> int:
    """Calculate age from date of birth"""
    today = date.today()
    return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))

def reverse_1to7(score: float) -> float:
    """Reverse score for negatively keyed items"""
    return 8.0 - score

def slider_with_ticks(label: str, key: str) -> float:
    """Create a dynamic slider with track color that moves with the thumb"""
    # Store the previous value to detect changes
    if f'{key}_prev_value' not in st.session_state:
        st.session_state[f'{key}_prev_value'] = 4

    value = st.slider(label, 1, 7, st.session_state[f'{key}_prev_value'], key=key)

    # Calculate the position of the thumb (0-100%)
    position = ((value - 1) / 6) * 100  # Convert 1-7 range to 0-100%

    # Dynamic color scheme based on value
    if value <= 2:
        left_color = "#FF6B6B"  # Red for low values
        right_color = "#4ECDC4"  # Teal for high values
        thumb_color = "#FF6B6B"
    elif value <= 4:
        left_color = "#FFE66D"  # Yellow for medium-low values
        right_color = "#4ECDC4"  # Teal for high values
        thumb_color = "#FFE66D"
    elif value <= 6:
        left_color = "#95E77E"  # Light green for medium-high values
        right_color = "#4ECDC4"  # Teal for high values
        thumb_color = "#95E77E"
    else:
        left_color = "#4ECDC4"  # Teal for very high values
        right_color = "#45B7D1"  # Blue for maximum values
        thumb_color = "#4ECDC4"

    # Create dynamic track color with smooth gradient
    st.markdown(f"""
    <style>
    .dynamic-slider-{key} > div > div {{
        background: linear-gradient(to right,
            {left_color} 0%,
            {left_color} {position}%.1f%,
            {right_color} {position}%.1f%,
            {right_color} 100%) !important;
        height: 8px !important;
        border-radius: 4px !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.15) !important;
        transition: all 0.3s ease !important;
    }}

    /* Dynamic thumb color */
    .dynamic-slider-{key} > div > div > div[role="slider"] {{
        background-color: {thumb_color} !important;
        border: 3px solid #FFFFFF !important;
        box-shadow: 0 2px 12px rgba(0,0,0,0.3) !important;
        transition: all 0.3s ease !important;
    }}

    .dynamic-slider-{key} > div > div > div[role="slider"]:hover {{
        transform: scale(1.1) !important;
        box-shadow: 0 4px 16px rgba(0,0,0,0.4) !important;
    }}

    /* Animated glow effect for active slider */
    @keyframes glow-{key} {{
        0% {{ box-shadow: 0 0 5px {thumb_color}40; }}
        50% {{ box-shadow: 0 0 20px {thumb_color}60; }}
        100% {{ box-shadow: 0 0 5px {thumb_color}40; }}
    }}

    .dynamic-slider-{key}:focus-within > div > div {{
        animation: glow-{key} 2s ease-in-out infinite !important;
    }}
    </style>
    """, unsafe_allow_html=True)

    # Apply the class to the slider container with improved, non-interfering JavaScript
    st.markdown(f"""
    <script>
    (function() {{
        function initSlider_{key}() {{
            // Try multiple selectors for different Streamlit versions
            var sliderElements = document.querySelectorAll('[data-testid="stSlider-{key}"], .stSlider-{key}, div[data-testid="stSlider"] input[type="range"]');
            if (sliderElements.length > 0) {{
                sliderElements.forEach(function(element) {{
                    // Only add to direct slider container, not parent elements
                    if (element.closest('[data-testid="stSlider-{key}"]') === element) {{
                        element.classList.add('dynamic-slider-{key}');
                    }}
                }});
                return;
            }}

            // More careful fallback to avoid interfering with other elements
            var allSliders = document.querySelectorAll('[data-testid*="stSlider"]');
            allSliders.forEach(function(slider) {{
                // Only target sliders that actually contain our key
                var hasCorrectKey = slider.querySelector('[data-testid*="{key}"]') ||
                                 slider.innerHTML.includes('data-testid="stSlider-{key}"') ||
                                 (slider.querySelector('input[type="range"]') &&
                                  slider.querySelector('input[type="range"]').value == "{value}");

                if (hasCorrectKey) {{
                    slider.classList.add('dynamic-slider-{key}');
                }}
            }});
        }}

        // Initialize after DOM is ready
        if (document.readyState === 'loading') {{
            document.addEventListener('DOMContentLoaded', initSlider_{key});
        }} else {{
            initSlider_{key}();
        }}

        // Also try after a short delay for Streamlit rendering
        setTimeout(initSlider_{key}, 200);
    }})();
    </script>
    """, unsafe_allow_html=True)

    # Store current value for next interaction
    st.session_state[f'{key}_prev_value'] = value

    # Add enhanced number labels below the slider with dynamic colors
    number_colors = []
    for i in range(1, 8):
        if i <= 2:
            color = "#FF6B6B"
        elif i <= 4:
            color = "#FFE66D"
        elif i <= 6:
            color = "#95E77E"
        else:
            color = "#4ECDC4"
        number_colors.append(color)

    st.markdown(f"""
    <div style='display: flex; justify-content: space-between; margin-top: -5px; margin-bottom: 15px; padding: 0 5px;'>
        <span style='color: {number_colors[0]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>1</span>
        <span style='color: {number_colors[1]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>2</span>
        <span style='color: {number_colors[2]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>3</span>
        <span style='color: {number_colors[3]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>4</span>
        <span style='color: {number_colors[4]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>5</span>
        <span style='color: {number_colors[5]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>6</span>
        <span style='color: {number_colors[6]}; font-weight: 700; font-size: 0.9rem; text-shadow: 0 1px 2px rgba(0,0,0,0.1);'>7</span>
    </div>

    <!-- Value indicator -->
    <div style='text-align: center; margin-top: -10px; margin-bottom: 10px;'>
        <span style='display: inline-block; background: {thumb_color}; padding: 4px 12px; border-radius: 20px; font-weight: 600; font-size: 0.85rem; box-shadow: 0 2px 8px rgba(0,0,0,0.2);'>
            <span style='color: #000000; font-weight: 600;'>Current Value: {value}</span>
        </span>
    </div>
    """, unsafe_allow_html=True)

    return value

def get_condition_code(condition: str) -> str:
    """Map condition names to codes used in main.py"""
    mapping = {
        "Dementia / Alzheimer's": "dementia",
        "Down Syndrome": "down_syndrome",
        "ADHD": "adhd"
    }
    return mapping.get(condition, "dementia")

# PATIENT DATABASE FUNCTIONS
def get_patient_db_connection():
    """Get connection to patient database"""
    db_path = Path("/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app/theramuse.db")
    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()

    # Create patients table if it doesn't exist
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patients (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            age INTEGER,
            sex TEXT,
            birthplace_city TEXT,
            birthplace_country TEXT,
            favorite_genre TEXT,
            favorite_musician TEXT,
            favorite_season TEXT,
            instruments TEXT,
            natural_elements TEXT,
            condition TEXT,
            difficulty_sleeping BOOLEAN,
            trouble_remembering BOOLEAN,
            forgets_everyday_things BOOLEAN,
            difficulty_recalling_old_memories BOOLEAN,
            memory_worse_than_year_ago BOOLEAN,
            visited_mental_health_professional BOOLEAN,
            extraversion REAL,
            agreeableness REAL,
            conscientiousness REAL,
            neuroticism REAL,
            openness REAL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Create sessions table if it doesn't exist
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS therapy_sessions (
            id TEXT PRIMARY KEY,
            patient_id TEXT,
            session_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            recommendations_count INTEGER,
            session_data TEXT,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )
    ''')

    # Check if table exists and has correct structure
    cursor.execute("PRAGMA table_info(patients)")
    columns = cursor.fetchall()

    if not columns:
        # Table doesn't exist, commit the creation
        conn.commit()

    return conn

def save_patient_to_database(patient_info: Dict, big5_scores: Dict, recommendations: Dict, session_id: str):
    """Save patient information to database"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    patient_id = f"patient_{datetime.now().strftime('%Y%m%d%H%M%S')}"

    # Insert patient data
    cursor.execute('''
        INSERT OR REPLACE INTO patients (
            id, name, age, sex, birthplace_city, birthplace_country,
            favorite_genre, favorite_musician, favorite_season,
            instruments, natural_elements, condition,
            difficulty_sleeping, trouble_remembering, forgets_everyday_things,
            difficulty_recalling_old_memories, memory_worse_than_year_ago,
            visited_mental_health_professional, extraversion, agreeableness,
            conscientiousness, neuroticism, openness, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        patient_id,
        patient_info.get('name', 'Anonymous'),
        patient_info.get('age', 0),
        patient_info.get('sex', ''),
        patient_info.get('birthplace_city', ''),
        patient_info.get('birthplace_country', ''),
        patient_info.get('favorite_genre', ''),
        patient_info.get('favorite_musician', ''),
        patient_info.get('favorite_season', ''),
        json.dumps(patient_info.get('instruments', [])),
        json.dumps(patient_info.get('natural_elements', [])),
        patient_info.get('condition', ''),
        patient_info.get('difficulty_sleeping', False),
        patient_info.get('trouble_remembering', False),
        patient_info.get('forgets_everyday_things', False),
        patient_info.get('difficulty_recalling_old_memories', False),
        patient_info.get('memory_worse_than_year_ago', False),
        patient_info.get('visited_mental_health_professional', False),
        big5_scores.get('extraversion', 0),
        big5_scores.get('agreeableness', 0),
        big5_scores.get('conscientiousness', 0),
        big5_scores.get('neuroticism', 0),
        big5_scores.get('openness', 0),
        datetime.now()
    ))

    # Insert session data
    cursor.execute('''
        INSERT INTO therapy_sessions (
            id, patient_id, recommendations_count, session_data
        ) VALUES (?, ?, ?, ?)
    ''', (
        session_id,
        patient_id,
        recommendations.get('total_songs', 0),
        json.dumps(recommendations)
    ))

    conn.commit()
    conn.close()
    return patient_id

def get_all_patients():
    """Get all patients from database with Big Five scores and reinforcement learning"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # First try the enhanced schema with id column
        cursor.execute('''
            SELECT p.id, p.name, p.age, p.condition, p.created_at,
                   p.extraversion, p.agreeableness, p.conscientiousness,
                   p.neuroticism, p.openness, p.reinforcement_learning
            FROM patients p
            ORDER BY p.created_at DESC
        ''')
        patients = cursor.fetchall()
        conn.close()
        return patients
    except sqlite3.OperationalError as e:
        # If the enhanced schema doesn't exist, try the basic schema
        if "no such column: p.id" in str(e):
            try:
                cursor.execute('''
                    SELECT p.patient_id, p.name, p.age, p.condition, p.created_at,
                           0 as extraversion, 0 as agreeableness, 0 as conscientiousness,
                           0 as neuroticism, 0 as openness, 0 as reinforcement_learning
                    FROM patients p
                    ORDER BY p.created_at DESC
                ''')
                patients = cursor.fetchall()
                conn.close()
                return patients
            except sqlite3.Error as e2:
                st.error(f"Database query error (basic schema): {str(e2)}")
                conn.close()
                return []
        else:
            st.error(f"Database query error: {str(e)}")
            conn.close()
            return []
    except sqlite3.Error as e:
        st.error(f"Database query error: {str(e)}")
        conn.close()
        return []

def get_patient_details(patient_id: str):
    """Get detailed information for a specific patient with Big Five scores and reinforcement learning"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # Try to get patient with enhanced schema first
        try:
            cursor.execute('SELECT * FROM patients WHERE id = ?', (patient_id,))
            patient = cursor.fetchone()
        except sqlite3.OperationalError:
            # Fall back to basic schema
            cursor.execute('SELECT * FROM patients WHERE patient_id = ?', (patient_id,))
            patient = cursor.fetchone()

        # Get Big Five scores with reinforcement learning
        cursor.execute('''
            SELECT * FROM big5_scores
            WHERE patient_id = ?
            ORDER BY created_at DESC
            LIMIT 1
        ''', (patient_id,))
        big5_scores = cursor.fetchone()

        cursor.execute('SELECT * FROM therapy_sessions WHERE patient_id = ? ORDER BY session_date DESC', (patient_id,))
        sessions = cursor.fetchall()

        # Combine patient data with Big Five scores if available
        if patient and big5_scores:
            patient = list(patient) + list(big5_scores[2:])  # Skip patient_id and session_id from big5_scores

        conn.close()
        return patient, sessions
    except sqlite3.Error as e:
        st.error(f"Database query error: {str(e)}")
        conn.close()
        return None, []

def delete_patient(patient_id: str):
    """Delete a patient from database"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    cursor.execute('DELETE FROM therapy_sessions WHERE patient_id = ?', (patient_id,))

    # Try both possible column names
    try:
        cursor.execute('DELETE FROM patients WHERE id = ?', (patient_id,))
    except sqlite3.OperationalError:
        cursor.execute('DELETE FROM patients WHERE patient_id = ?', (patient_id,))

    conn.commit()
    conn.close()

def get_comprehensive_patient_data():
    """Get comprehensive patient data with therapy recommendations"""
    conn = get_patient_db_connection()
    cursor = conn.cursor()

    try:
        # Get basic patient info with backward compatibility
        try:
            cursor.execute('''
                SELECT id, name, age, condition, created_at
                FROM patients ORDER BY created_at DESC
            ''')
            patients = cursor.fetchall()
            id_col = "id"
        except sqlite3.OperationalError:
            cursor.execute('''
                SELECT patient_id, name, age, condition, created_at
                FROM patients ORDER BY created_at DESC
            ''')
            patients = cursor.fetchall()
            id_col = "patient_id"

        patient_data = []

        for patient in patients:
            patient_id = patient[0]

            # Get therapy sessions and recommendations
            cursor.execute('''
                SELECT id, session_date, recommendations_count
                FROM therapy_sessions WHERE patient_id = ?
                ORDER BY session_date DESC
            ''', (patient_id,))
            sessions = cursor.fetchall()

            # Get therapy recommendations with song details
            cursor.execute('''
                SELECT category, song_title, video_id, channel, rank
                FROM therapy_recommendations WHERE patient_id = ?
                ORDER BY category, rank
            ''', (patient_id,))
            recommendations = cursor.fetchall()

            # Get feedback data
            cursor.execute('''
                SELECT feedback_type, reward, created_at
                FROM therapy_feedback WHERE patient_id = ?
                ORDER BY created_at DESC
            ''', (patient_id,))
            feedback = cursor.fetchall()

            # Get Big Five scores if available
            cursor.execute('''
                SELECT openness, conscientiousness, extraversion,
                       agreeableness, neuroticism, reinforcement_learning
                FROM big5_scores WHERE patient_id = ?
                ORDER BY created_at DESC LIMIT 1
            ''', (patient_id,))
            big5_scores = cursor.fetchone()

            patient_data.append({
                'patient_info': patient,
                'sessions': sessions,
                'recommendations': recommendations,
                'feedback': feedback,
                'big5_scores': big5_scores or (0, 0, 0, 0, 0, 0)
            })

        conn.close()
        return patient_data

    except Exception as e:
        st.error(f"Database error: {str(e)}")
        conn.close()
        return []

def page_patient_database():
    """Advanced Patient Database Management Page"""
    render_logo()

    st.markdown("""
    <div class='glass-card'>
        <h2> Advanced Patient Database</h2>
        <p style='color: rgba(255,255,255,0.7);'>Comprehensive patient records with therapy analytics and recommendations</p>
    </div>
    """, unsafe_allow_html=True)

    # Database connection status
    try:
        conn = get_patient_db_connection()
        cursor = conn.cursor()

        # Get database statistics
        cursor.execute("SELECT COUNT(*) FROM patients")
        total_patients = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_sessions")
        total_sessions = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_recommendations")
        total_recommendations = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(*) FROM therapy_feedback")
        total_feedback = cursor.fetchone()[0]

        conn.close()

        # Display statistics
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown(f"""
            <div class='metric-card'>
                <h3>Total Patients</h3>
                <h2>{total_patients}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown(f"""
            <div class='metric-card'>
                <h3>Therapy Sessions</h3>
                <h2>{total_sessions}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col3:
            st.markdown(f"""
            <div class='metric-card'>
                <h3>Song Recommendations</h3>
                <h2>{total_recommendations}</h2>
            </div>
            """, unsafe_allow_html=True)

        with col4:
            st.markdown(f"""
            <div class='metric-card'>
                <h3>Feedback Recorded</h3>
                <h2>{total_feedback}</h2>
            </div>
            """, unsafe_allow_html=True)

    except Exception as e:
        st.error(f"Database connection error: {str(e)}")
        return

    # Advanced filters and search
    st.markdown("###  Advanced Search & Filters")

    col1, col2, col3 = st.columns(3)

    with col1:
        search_term = st.text_input(" Search patients by name...", placeholder="Enter patient name...")

    with col2:
        condition_filter = st.selectbox(" Filter by condition",
                                       ["All Conditions", "Dementia / Alzheimer's", "Down Syndrome", "ADHD"])

    with col3:
        sort_by = st.selectbox(" Sort by",
                              ["Latest First", "Oldest First", "Most Sessions", "Most Feedback"])

    # Get comprehensive patient data
    patient_data = get_comprehensive_patient_data()

    if not patient_data:
        st.info("No patient records found in the database.")
        return

    # Apply filters
    filtered_data = patient_data.copy()

    if search_term:
        filtered_data = [p for p in filtered_data
                        if search_term.lower() in p['patient_info'][1].lower()]

    if condition_filter != "All Conditions":
        condition_map = {
            "Dementia / Alzheimer's": "dementia",
            "Down Syndrome": "down_syndrome",
            "ADHD": "adhd"
        }
        filtered_data = [p for p in filtered_data
                        if condition_map.get(condition_filter, "") in str(p['patient_info'][3]).lower()]

    # Apply sorting
    if sort_by == "Latest First":
        filtered_data.sort(key=lambda x: x['patient_info'][4] if len(x['patient_info']) > 4 else "", reverse=True)
    elif sort_by == "Oldest First":
        filtered_data.sort(key=lambda x: x['patient_info'][4] if len(x['patient_info']) > 4 else "")
    elif sort_by == "Most Sessions":
        filtered_data.sort(key=lambda x: len(x['sessions']), reverse=True)
    elif sort_by == "Most Feedback":
        filtered_data.sort(key=lambda x: len(x['feedback']), reverse=True)

    if not filtered_data:
        st.warning("No patients match the selected filters.")
        return

    st.success(f"Found {len(filtered_data)} patient(s) matching your criteria")

    # Enhanced patient display
    for idx, patient in enumerate(filtered_data):
        patient_info = patient['patient_info']
        sessions = patient['sessions']
        recommendations = patient['recommendations']
        feedback = patient['feedback']
        big5_scores = patient['big5_scores']

        with st.expander(f"👤 {patient_info[1]} ({patient_info[2]} years) - {patient_info[3].title()}", expanded=False):

            # Patient Overview
            col1, col2, col3 = st.columns([2, 1, 1])

            with col1:
                st.markdown("#### 📋 Patient Information")
                st.write(f"**ID:** {patient_info[0]}")
                st.write(f"**Age:** {patient_info[2]}")
                st.write(f"**Condition:** {patient_info[3].title()}")
                if len(patient_info) > 4:
                    st.write(f"**Registered:** {patient_info[4]}")

            with col2:
                st.markdown("#### 🎭 Personality Profile")
                if big5_scores:
                    for trait, score in zip(['Openness', 'Conscientiousness', 'Extraversion', 'Agreeableness', 'Neuroticism'], big5_scores[:5]):
                        if score > 0:
                            st.metric(trait, f"{score:.1f}/7.0")
                    st.metric("RL Interactions", big5_scores[5] if len(big5_scores) > 5 else 0)
                else:
                    st.info("No personality data")

            with col3:
                st.markdown("#### 📊 Activity Summary")
                st.metric("Sessions", len(sessions))
                st.metric("Songs", len(recommendations))
                st.metric("Feedback", len(feedback))

            # Therapy Sessions
            if sessions:
                st.markdown("#### 🏥 Therapy Sessions")
                session_df = pd.DataFrame(sessions,
                                         columns=['Session ID', 'Date', 'Songs Recommended'])
                if not session_df.empty:
                    session_df['Date'] = pd.to_datetime(session_df['Date']).dt.strftime('%Y-%m-%d %H:%M')
                    st.dataframe(session_df, hide_index=True, use_container_width=True)

            # Song Recommendations
            if recommendations:
                st.markdown("#### 🎵 Song Recommendations")

                # Group recommendations by category
                categories = {}
                for rec in recommendations:
                    category = rec[0]
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(rec)

                for category, songs in categories.items():
                    with st.expander(f"🎼 {category.replace('_', ' ').title()} ({len(songs)} songs)", expanded=False):
                        for song in songs[:5]:  # Show first 5 songs per category
                            col_song, col_video = st.columns([3, 1])
                            with col_song:
                                st.write(f"**{song[1]}**")
                                if song[2]:
                                    st.caption(f"Channel: {song[3] if len(song) > 3 else 'N/A'}")
                            with col_video:
                                if song[2]:
                                    st.markdown(f"[▶️ Watch](https://www.youtube.com/watch?v={song[2]})")

                        if len(songs) > 5:
                            st.caption(f"... and {len(songs) - 5} more songs in this category")

            # Feedback Analysis
            if feedback:
                st.markdown("#### 💬 Feedback Analysis")

                feedback_counts = {}
                for fb in feedback:
                    fb_type = fb[0]
                    feedback_counts[fb_type] = feedback_counts.get(fb_type, 0) + 1

                col1, col2, col3 = st.columns(3)
                with col1:
                    if 'like' in feedback_counts:
                        st.metric("👍 Likes", feedback_counts['like'])
                with col2:
                    if 'dislike' in feedback_counts:
                        st.metric("👎 Dislikes", feedback_counts['dislike'])
                with col3:
                    if 'skip' in feedback_counts:
                        st.metric("⏭️ Skips", feedback_counts['skip'])

                # Recent feedback
                st.markdown("**Recent Feedback:**")
                for fb in feedback[:3]:
                    emoji = {"like": "👍", "dislike": "👎", "skip": "⏭️", "neutral": "😐"}.get(fb[0], "📝")
                    st.write(f"{emoji} {fb[0].title()} - {fb[2] if len(fb) > 2 else 'N/A'}")

            # Action buttons
            st.markdown("---")
            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button(f"📥 Export Data", key=f"export_{patient_info[0]}", use_container_width=True):
                    export_data = {
                        'patient_info': patient_info,
                        'sessions': sessions,
                        'recommendations': recommendations,
                        'feedback': feedback,
                        'big5_scores': big5_scores
                    }
                    st.download_button(
                        label="📄 Download JSON",
                        data=json.dumps(export_data, indent=2, default=str),
                        file_name=f"patient_{patient_info[1]}_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json",
                        use_container_width=True
                    )

            with col2:
                if st.button(f"🔄 Refresh Data", key=f"refresh_{patient_info[0]}", use_container_width=True):
                    st.rerun()

            with col3:
                if st.button(f"🗑️ Delete Patient", key=f"delete_{patient_info[0]}", type="secondary", use_container_width=True):
                    st.session_state[f"confirm_delete_{patient_info[0]}"] = True

            # Confirmation dialog
            if st.session_state.get(f"confirm_delete_{patient_info[0]}", False):
                st.warning("⚠️ **Confirm Deletion**: This will permanently delete all patient data including sessions and recommendations.")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✅ Yes, Delete", key=f"confirm_yes_{patient_info[0]}", type="primary"):
                        delete_patient(patient_info[0])
                        st.success("Patient deleted successfully!")
                        st.session_state[f"confirm_delete_{patient_info[0]}"] = False
                        st.rerun()
                with col2:
                    if st.button("❌ Cancel", key=f"confirm_no_{patient_info[0]}"):
                        st.session_state[f"confirm_delete_{patient_info[0]}"] = False
                        st.rerun()

def create_docx_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> bytes:
    """Create a DOCX file with patient information and recommendations"""
    doc = Document()

    # Add title
    title = doc.add_heading('TheraMuse Therapy Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add patient information
    doc.add_heading('Patient Information', level=1)
    patient_table = doc.add_table(rows=0, cols=2)
    patient_table.style = 'Table Grid'

    patient_data = [
        ('Name', patient_info.get('name', 'N/A')),
        ('Age', str(patient_info.get('age', 'N/A'))),
        ('Sex', patient_info.get('sex', 'N/A')),
        ('Birthplace City', patient_info.get('birthplace_city', 'N/A')),
        ('Birthplace Country', patient_info.get('birthplace_country', 'N/A')),
        ('Condition', str(patient_info.get('condition', 'N/A')).title()),
        ('Favorite Genres', patient_info.get('favorite_genre', 'N/A')),
        ('Favorite Musician', patient_info.get('favorite_musician', 'N/A')),
        ('Favorite Season', patient_info.get('favorite_season', 'N/A')),
        ('Preferred Instruments', ', '.join(patient_info.get('instruments', []))),
        ('Natural Elements', ', '.join(patient_info.get('natural_elements', [])))
    ]

    for key, value in patient_data:
        row_cells = patient_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    # Add Big-5 personality scores
    doc.add_heading('Personality Profile (Big-5)', level=1)
    personality_table = doc.add_table(rows=0, cols=2)
    personality_table.style = 'Table Grid'

    personality_labels = {
        'extraversion': 'Extraversion',
        'agreeableness': 'Agreeableness',
        'conscientiousness': 'Conscientiousness',
        'neuroticism': 'Neuroticism',
        'openness': 'Openness'
    }

    for key, label in personality_labels.items():
        score = big5_scores.get(key, 0)
        row_cells = personality_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{score:.2f}/7.00"

    # Add recommendations
    doc.add_heading('Music Recommendations', level=1)

    if 'categories' in recommendations:
        category_labels = {
            "birthplace_country": "From Your Country",
            "birthplace_city": "From Your City",
            "instruments": "Instrumental Favorites",
            "seasonal": "Seasonal Music",
            "natural_elements": "Nature-Inspired",
            "favorite_genre": "Favorite Genres",
            "favorite_musician": "Favorite Musician",
            "therapeutic": "Therapeutic Selections",
            "personality_based": "Personality Match",
            "calming_sensory": "Calming Sensory",
            "concentration": "Focus & Concentration",
            "binaural_beats": "Binaural Beats",
            "relief_study": "Study & Relief",
            "additional_calm": "Additional Calming",
            "additional_focus": "Additional Focus"
        }

        for category, data in recommendations['categories'].items():
            label = category_labels.get(category, category.replace('_', ' ').title())
            songs = data.get('songs', [])

            if songs:
                doc.add_heading(label, level=2)
                for idx, song in enumerate(songs, 1):
                    title = song.get('title', 'Unknown Title')
                    channel = song.get('channel', 'Unknown Channel')
                    url = song.get('url', 'No URL available')

                    # Add song paragraph with title and channel
                    song_para = doc.add_paragraph()
                    song_para.add_run(f"{idx}. {title} - {channel}").bold = True

                    # Add YouTube URL
                    if url and url.startswith('https://www.youtube.com/'):
                        url_para = doc.add_paragraph()
                        url_para.add_run("YouTube Link: ").bold = True
                        url_para.add_run(f"{url}")
                    else:
                        url_para = doc.add_paragraph()
                        url_para.add_run("URL: ").bold = True
                        url_para.add_run(url if url else "No URL available")

                    doc.add_paragraph()  # Add spacing between songs

    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Generated on: ').bold = True
    footer.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Save document to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

def create_pdf_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> str:
    """Create a simple HTML-based PDF content"""
    html_content = f"""
    <html>
    <head>
        <title>TheraMuse Therapy Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ color: #A3B18A; text-align: center; }}
            h2 {{ color: #588157; border-bottom: 2px solid #A3B18A; }}
            h3 {{ color: #E8DDB5; }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            .song {{ margin: 5px 0; padding: 5px; background-color: #f9f9f9; }}
        </style>
    </head>
    <body>
        <h1>TheraMuse Therapy Report</h1>
        <p><strong>Generated on:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

        <h2>Patient Information</h2>
        <table>
            <tr><th>Field</th><th>Value</th></tr>
            <tr><td>Name</td><td>{patient_info.get('name', 'N/A')}</td></tr>
            <tr><td>Age</td><td>{patient_info.get('age', 'N/A')}</td></tr>
            <tr><td>Sex</td><td>{patient_info.get('sex', 'N/A')}</td></tr>
            <tr><td>Birthplace City</td><td>{patient_info.get('birthplace_city', 'N/A')}</td></tr>
            <tr><td>Birthplace Country</td><td>{patient_info.get('birthplace_country', 'N/A')}</td></tr>
            <tr><td>Condition</td><td>{patient_info.get('condition', 'N/A').title()}</td></tr>
            <tr><td>Favorite Genres</td><td>{patient_info.get('favorite_genre', 'N/A')}</td></tr>
            <tr><td>Favorite Musician</td><td>{patient_info.get('favorite_musician', 'N/A')}</td></tr>
            <tr><td>Favorite Season</td><td>{patient_info.get('favorite_season', 'N/A')}</td></tr>
            <tr><td>Preferred Instruments</td><td>{', '.join(patient_info.get('instruments', []))}</td></tr>
            <tr><td>Natural Elements</td><td>{', '.join(patient_info.get('natural_elements', []))}</td></tr>
        </table>

        <h2>Personality Profile (Big-5)</h2>
        <table>
            <tr><th>Trait</th><th>Score</th></tr>
            <tr><td>Extraversion</td><td>{big5_scores.get('extraversion', 0):.2f}/7.00</td></tr>
            <tr><td>Agreeableness</td><td>{big5_scores.get('agreeableness', 0):.2f}/7.00</td></tr>
            <tr><td>Conscientiousness</td><td>{big5_scores.get('conscientiousness', 0):.2f}/7.00</td></tr>
            <tr><td>Neuroticism</td><td>{big5_scores.get('neuroticism', 0):.2f}/7.00</td></tr>
            <tr><td>Openness</td><td>{big5_scores.get('openness', 0):.2f}/7.00</td></tr>
        </table>

        <h2>Music Recommendations</h2>
    """

    if 'categories' in recommendations:
        category_labels = {
            "birthplace_country": "From Your Country",
            "birthplace_city": "From Your City",
            "instruments": "Instrumental Favorites",
            "seasonal": "Seasonal Music",
            "natural_elements": "Nature-Inspired",
            "favorite_genre": "Favorite Genres",
            "favorite_musician": "Favorite Musician",
            "therapeutic": "Therapeutic Selections",
            "personality_based": "Personality Match",
            "calming_sensory": "Calming Sensory",
            "concentration": "Focus & Concentration",
            "binaural_beats": "Binaural Beats",
            "relief_study": "Study & Relief",
            "additional_calm": "Additional Calming",
            "additional_focus": "Additional Focus"
        }

        for category, data in recommendations['categories'].items():
            label = category_labels.get(category, category.replace('_', ' ').title())
            songs = data.get('songs', [])

            if songs:
                html_content += f"<h3>{label}</h3>"
                for idx, song in enumerate(songs, 1):
                    title = song.get('title', 'Unknown Title')
                    channel = song.get('channel', 'Unknown Channel')
                    url = song.get('url', 'No URL available')

                    html_content += f'<div class="song">'
                    html_content += f'<strong>{idx}. {title} - {channel}</strong><br>'

                    # Add YouTube link if available
                    if url and url.startswith('https://www.youtube.com/'):
                        html_content += f'<a href="{url}" target="_blank" style="color: #007bff; text-decoration: underline;">📺 YouTube Link</a><br>'
                    else:
                        html_content += f'<span style="color: #666;">URL: {url}</span><br>'

                    html_content += '</div>'

    html_content += """
    </body>
    </html>
    """

    return html_content

def create_json_download(patient_info: Dict, recommendations: Dict, big5_scores: Dict) -> str:
    """Create a JSON file with all data"""
    export_data = {
        "patient_info": patient_info,
        "recommendations": recommendations,
        "big5_scores": big5_scores,
        "export_timestamp": datetime.now().isoformat(),
        "export_version": "TheraMuse v9.0"
    }
    return json.dumps(export_data, indent=2, ensure_ascii=False)

def render_download_options(patient_info: Dict, recommendations: Dict, big5_scores: Dict):
    """Render download options at the end of the page"""
    st.markdown("---")

    # Research Evidence Section
    st.markdown("## Research Evidence")
    st.write("Scientific evidence supporting music therapy for various conditions and personality-based music preferences")

    if st.button(" View Research Evidence", key="research_evidence_btn", width='stretch'):
        st.session_state.page = "Research Evidence"
        st.rerun()

    st.markdown("""
    <div class='glass-card' style='margin-top: 1rem;'>
        <h2>Download Your Therapy Report</h2>
        <p style='color: rgba(255,255,255,0.7);'>Save your personalized music therapy recommendations in different formats</p>
    </div>
    """, unsafe_allow_html=True)

    # Create download data
    docx_data = create_docx_download(patient_info, recommendations, big5_scores)
    json_data = create_json_download(patient_info, recommendations, big5_scores)
    pdf_data = create_pdf_download(patient_info, recommendations, big5_scores)

    # Download buttons in columns
    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            label=" Download DOCX",
            data=docx_data,
            file_name=f"theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width='stretch'
        )

    with col2:
        st.markdown(f"""
        <a href="data:text/html;base64,{base64.b64encode(pdf_data.encode()).decode()}"
           download="theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
           style="
               display: block;
               width: 100%;
               padding: 0.75rem 1.5rem;
               background: linear-gradient(135deg, #A3B18A 0%, #588157 100%);
               color: white;
               text-decoration: none;
               border-radius: 12px;
               font-weight: 600;
               text-align: center;
               transition: all 0.3s ease;
               box-shadow: 0 4px 16px rgba(163, 177, 138, 0.3);
               margin: 0;
        ">
             Download PDF 
        </a>
        """, unsafe_allow_html=True)

    with col3:
        st.download_button(
            label=" Download JSON",
            data=json_data,
            file_name=f"theramuse_report_{patient_info.get('name', 'patient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            width='stretch'
        )

    st.markdown("""
    <div style='text-align: center; margin-top: 1rem; color: rgba(255,255,255,0.6); font-size: 0.9rem;'>
         <em>DOCX format includes formatted tables and is best for printing. JSON contains all raw data for developers. PDF (HTML) can be printed to PDF from your browser.</em>
    </div>
    """, unsafe_allow_html=True)

# PERSONALITY & RAGA FUNCTIONS

def create_personality_radar(base_values: Dict, adjusted_values: Dict = None):
    """Create premium radar chart for Big-5 personality"""
    categories = ['Openness', 'Conscientiousness', 'Extraversion', 'Agreeableness', 'Neuroticism']
    
    fig = go.Figure()
    
    # Base personality with premium colors
    base_scores = [base_values.get(k.lower(), 4) for k in categories]
    fig.add_trace(go.Scatterpolar(
        r=base_scores,
        theta=categories,
        fill='toself',
        name='Base Personality',
        line=dict(color='#A3B18A', width=3),
        fillcolor='rgba(163, 177, 138, 0.2)'
    ))
    
    # Adjusted personality
    if adjusted_values:
        adj_scores = [adjusted_values.get(k.lower(), 4) for k in categories]
        fig.add_trace(go.Scatterpolar(
            r=adj_scores,
            theta=categories,
            fill='toself',
            name='Adjusted (Nature)',
            line=dict(color='#E8DDB5', width=3),
            fillcolor='rgba(232, 221, 181, 0.2)'
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True, 
                range=[0, 7],
                gridcolor='rgba(255, 255, 255, 0.1)',
                linecolor='rgba(255, 255, 255, 0.2)'
            ),
            angularaxis=dict(
                gridcolor='rgba(255, 255, 255, 0.1)',
                linecolor='rgba(255, 255, 255, 0.2)'
            ),
            bgcolor='rgba(0, 0, 0, 0)'
        ),
        showlegend=True,
        title={
            'text': "Big-5 Personality Profile",
            'font': {'size': 24, 'color': '#E8DDB5', 'family': 'Inter'}
        },
        paper_bgcolor='rgba(0, 0, 0, 0)',
        plot_bgcolor='rgba(0, 0, 0, 0)',
        font=dict(color='white', family='Inter')
    )
    
    return fig

def get_generational_ragas(birth_year: int) -> List[str]:
    """Get therapeutic ragas based on birth year"""
    matrix = BangladeshiGenerationalMatrix()
    context = matrix.get_generational_context(birth_year)
    return context.get("therapeutic_ragas", [])

# PREMIUM RECOMMENDATION DISPLAY

def extract_youtube_id(url: str) -> Optional[str]:
    """Best-effort extraction of a YouTube video ID from various URL shapes."""
    if not url:
        return None
    patterns = [
        r"youtu\.be/([A-Za-z0-9_-]{6,})",
        r"v=([A-Za-z0-9_-]{6,})",
        r"/embed/([A-Za-z0-9_-]{6,})",
        r"/shorts/([A-Za-z0-9_-]{6,})",
    ]
    for pat in patterns:
        m = re.search(pat, url)
        if m:
            vid = m.group(1)
            # Strip any trailing params just in case
            for sep in ("?", "&", "#"):
                vid = vid.split(sep)[0]
            return vid
    return None

def normalize_youtube_from_song(song: Dict) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Return (video_id, youtube_url, embed_url) using id dict, id str, or URL fallbacks."""
    video_id: Optional[str] = None
    # 1) id might be dict with videoId
    _id = song.get('id')
    if isinstance(_id, dict):
        video_id = _id.get('videoId') or _id.get('video_id')
    # 2) sometimes videoId is at top level
    if not video_id:
        video_id = song.get('videoId') or song.get('video_id')
    # 3) id might be a string containing either the id or a URL
    if not video_id and isinstance(_id, str):
        if re.fullmatch(r"[A-Za-z0-9_-]{6,}", _id):
            video_id = _id
        else:
            video_id = extract_youtube_id(_id)
    # 4) URL candidates
    url_candidates = [
        song.get('url'), song.get('youtube_url'), song.get('link'),
        song.get('webpage_url'), song.get('watch_url')
    ]
    youtube_url: Optional[str] = None
    for u in url_candidates:
        if not u:
            continue
        if not youtube_url and ("youtube" in u or "youtu.be" in u):
            youtube_url = u
        if not video_id:
            video_id = extract_youtube_id(u)
    # 5) Construct URLs from ID if needed
    if video_id and not youtube_url:
        youtube_url = f"https://www.youtube.com/watch?v={video_id}"
    embed_url = f"https://www.youtube.com/embed/{video_id}" if video_id else None
    return video_id, youtube_url, embed_url

def display_song_card(song: Dict, category: str, rank: int):
    """Display a premium song card with embedded YouTube player and link below title."""
    title = song.get('title', 'Unknown Title')
    channel = song.get('channel', 'Unknown Channel')
    description = song.get('description', '')[:200] + '...' if song.get('description') else ''

    video_id, youtube_url, youtube_embed = normalize_youtube_from_song(song)

    # Build link markup if we have any usable URL
    youtube_link_markup = (
        f'<a href="{youtube_url}" target="_blank" class="youtube-link" '
        f'style="font-size: 0.9rem; margin-top: 0.25rem; display: inline-block;"> Watch on YouTube</a>'
        if youtube_url else ''
    )

    # Always show the card with the link under the title; add iframe only if we have a video id
    ifr = (
        f"<div class='youtube-embed-container'>\n"
        f"  <iframe src=\"{youtube_embed}\" title=\"{title}\" frameborder=\"0\" "
        f"allow=\"accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture\" "
        f"referrerpolicy=\"strict-origin-when-cross-origin\" loading=\"lazy\" allowfullscreen></iframe>\n"
        f"</div>"
        if video_id and youtube_embed else ''
    )

    st.markdown(f"""
    <div class='song-card' style='animation-delay: {rank * 0.05}s;'>
        <div style='display: flex; align-items: center; margin-bottom: 1rem;'>
            <span class='rank-badge'>#{rank}</span>
            <div>
                <h4 style='margin: 0;'>{title}</h4>
                {youtube_link_markup}
            </div>
        </div>
        <p><strong>Channel:</strong> {channel}</p>
        <p><strong>Category:</strong> <span class='category-badge'>{category}</span></p>
        {f'<p style="font-size: 0.85rem; color: rgba(255,255,255,0.5);">{description}</p>' if description else ''}
        {ifr}
    </div>
    """, unsafe_allow_html=True)

def render_recommendations_with_feedback(recommendations: Dict, patient_info: Dict, 
                                        session_id: str, patient_id: str):
    """Render premium recommendations with feedback options"""
    
    if not recommendations or 'categories' not in recommendations:
        st.warning("No recommendations available.")
        return
    
    st.markdown(f"""
    <div class='glass-card'>
        <h2 style='margin-bottom: 0.5rem;'> Personalized Music Recommendations</h2>
        <p style='color: rgba(255,255,255,0.7); font-size: 1.1rem;'>
            <strong style='color: #A3B18A;'>{recommendations.get('total_songs', 0)}</strong> songs curated for your therapy
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Get TheraMuse instance
    if 'theramuse' not in st.session_state:
        st.session_state.theramuse = TheraMuse(db_path="/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app/theramuse.db")

    theramuse = st.session_state.theramuse
    
    # Display by category
    category_labels = {
        "birthplace_country": " From Your Country",
        "birthplace_city": " From Your City",
        "instruments": " Instrumental Favorites",
        "seasonal": " Seasonal Music",
        "natural_elements": " Nature-Inspired",
        "favorite_genre": " Favorite Genres",
        "favorite_musician": " Favorite Musician",
        "therapeutic": " Therapeutic Selections",
        "personality_based": " Personality Match",
        "calming_sensory": " Calming Sensory",
        "concentration": " Focus & Concentration",
        "binaural_beats": " Binaural Beats",
        "relief_study": " Study & Relief",
        "additional_calm": " Additional Calming",
        "additional_focus": "Additional Focus"
    }
    
    for category, data in recommendations['categories'].items():
        label = category_labels.get(category, category.replace('_', ' ').title())
        songs = data.get('songs', [])
        
        if not songs:
            continue
            
        with st.expander(f"{label} ({len(songs)} songs)", expanded=True):
                  
            for idx, song in enumerate(songs, 1):
                col1, col2 = st.columns([5, 1])
                
                with col1:
                    display_song_card(song, label, idx)
                
                with col2:
                    feedback_key = f"feedback_{session_id}_{category}_{idx}"
                    
                    st.markdown("<div class='feedback-button'>", unsafe_allow_html=True)
                    
                    # Feedback buttons
                    if st.button(" Like", key=f"like_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id, 
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "like", patient_info
                        )
                        st.success(" Feedback recorded!")
                        st.toast("Thank you for your feedback!")
                    
                    if st.button(" Dislike", key=f"dislike_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id,
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "dislike", patient_info
                        )
                        st.info("Feedback recorded!")
                    
                    if st.button(" Skip", key=f"skip_{feedback_key}", width='stretch'):
                        theramuse.record_feedback(
                            patient_id, session_id,
                            get_condition_code(patient_info.get('condition', 'dementia')),
                            song, "skip", patient_info
                        )
                        st.info("Skipped!")
                    
                    st.markdown("</div>", unsafe_allow_html=True)

# STREAMLIT PAGES

def page_intake():
    """Patient intake page with premium design"""
    render_logo(size=220)
    
    st.markdown("""
    <div class='glass-card'>
        <h2>Patient Intake</h2>
        <p style='color: rgba(255,255,255,0.7); font-size: 1.05rem;'>
            Complete the assessment to receive TheramuseRX personalized music recommendations from YouTube.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize session state flags
    if 'show_results' not in st.session_state:
        st.session_state.show_results = False
    if 'processing_complete' not in st.session_state:
        st.session_state.processing_complete = False

    # Only show form if not currently showing results
    if not st.session_state.show_results:
        with st.form("intake_form", clear_on_submit=False):
            st.markdown("###  Demographics")
            col1, col2 = st.columns(2)

            name = col1.text_input("Name", placeholder="John Doe")
            birthplace_city = col1.text_input("Birthplace City", placeholder="e.g., Dhaka")
            dob = col1.date_input("Date of Birth *", value=None, min_value=date(1900, 1, 1), max_value=date.today())
            birthplace_country = col2.text_input("Birthplace Country *", placeholder="e.g., Bangladesh, India, USA")
            sex = col1.selectbox("Sex", ["Male", "Female", "Other"])
            
            st.markdown("###  Musical Preferences")
            preferred_instruments = st.multiselect(
                "Preferred Instruments",
                ["Piano", "Guitar", "Violin", "Flute", "Sitar", "Tabla", "Drums", "Saxophone"]
            )

            col1, col2 = st.columns(2)
            with col1:
                # Support multiple favorite genres
                st.markdown("**Favorite Music Genres** (select up to 3)")
                favorite_genre1 = st.selectbox(
                    "Favorite Genre 1 *",
                    ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                    key="favorite_genre1"
                )
                favorite_genre2 = st.selectbox(
                    "Favorite Genre 2 (optional)",
                    ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                    key="favorite_genre2"
                )
                favorite_genre3 = st.selectbox(
                    "Favorite Genre 3 (optional)",
                    ["", "Classical", "Jazz", "Rock", "Pop", "R&B", "Hip-Hop", "Country", "Folk", "Electronic", "Metal", "Indie", "Blues", "Reggae", "Classical Fusion", "Bengali Folk"],
                    key="favorite_genre3"
                )

                # Combine favorite genres into a comma-separated string
                favorite_genres = [genre for genre in [favorite_genre1, favorite_genre2, favorite_genre3] if genre]
                favorite_genre = ", ".join(favorite_genres) if favorite_genres else ""

            with col2:
                favorite_musician = st.text_input(
                    "Favorite Musician/Artist",
                    placeholder="e.g., Beatles, Mozart, A.R. Rahman",
                    key="favorite_musician_input",
                    on_change=None
                )

            st.markdown("###  Environmental Preferences")
            col1, col2 = st.columns(2)
            with col1:
                favorite_season = st.selectbox(
                    "Favorite Season",
                    ["", "Spring", "Summer", "Monsoon", "Autumn", "Winter"]
                )
            with col2:
                natural_elements = st.multiselect(
                    "Natural Elements You Connect With",
                    ["Rain", "Forest", "Ocean", "Mountains", "Desert", "Rivers", "Sunset", "Sunrise"]
                )
            
            st.markdown("###  Health & Wellness Profile")
            st.markdown("Select the condition for therapy:")
            
            condition = st.selectbox(
                "Primary Condition *",
                ["Dementia / Alzheimer's", "Down Syndrome", "ADHD"]
            )
            
            # Condition-specific questions
            if condition == "Dementia / Alzheimer's":
                st.markdown("#### Memory & Sleep Assessment")
                difficulty_sleeping = st.checkbox("Difficulty sleeping?")
                trouble_remembering = st.checkbox("Trouble remembering recent things?")
                forgets_everyday_things = st.checkbox("Often forgets everyday things?")
                difficulty_recalling_old_memories = st.checkbox("Difficulty recalling older memories?")
                memory_worse_than_year_ago = st.checkbox("Memory worse than a year ago?")
                visited_mental_health = st.checkbox("Visited mental health professional?")
            else:
                difficulty_sleeping = False
                trouble_remembering = False
                forgets_everyday_things = False
                difficulty_recalling_old_memories = False
                memory_worse_than_year_ago = False
                visited_mental_health = False
            
            st.markdown("###  Personality Assessment (Big-5)")
            st.markdown("Rate each statement from 1 (Strongly Disagree) to 7 (Strongly Agree)")
            
            statements = [
                "I see myself as extraverted, enthusiastic",
                "I see myself as critical, quarrelsome",
                "I see myself as dependable, self-disciplined",
                "I see myself as anxious, easily upset",
                "I see myself as open to new experiences, complex",
                "I see myself as reserved, quiet",
                "I see myself as sympathetic, warm",
                "I see myself as disorganized, careless",
                "I see myself as calm, emotionally stable",
                "I see myself as conventional, uncreative"
            ]
            
            big5_items = [slider_with_ticks(f"{i+1}. {s}", key=f"big5_{i+1}") for i, s in enumerate(statements)]

            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                submitted = st.form_submit_button("TheramusRX Recommendations", type="primary", width='stretch')

        # ONLY process when form is explicitly submitted
        if submitted:
            # Validate required fields
            if not dob or not birthplace_country:
                st.error(" Please fill in Date of Birth and Birthplace Country.")
                st.stop()  # Stop execution here

            # Validate at least one favorite genre is selected
            if not favorite_genre1:
                st.error(" Please select at least one favorite music genre.")
                st.stop()  # Stop execution here

            # Calculate age and birth year
            age = compute_age_from_dob(dob)
            birth_year = dob.year

            # Calculate Big-5 scores
            E_base = (big5_items[0] + reverse_1to7(big5_items[5])) / 2.0
            A_base = (reverse_1to7(big5_items[1]) + big5_items[6]) / 2.0
            C_base = (big5_items[2] + reverse_1to7(big5_items[7])) / 2.0
            N_base = (big5_items[3] + reverse_1to7(big5_items[8])) / 2.0
            O_base = (big5_items[4] + reverse_1to7(big5_items[9])) / 2.0

            big5_scores = {
                "extraversion": E_base,
                "agreeableness": A_base,
                "conscientiousness": C_base,
                "neuroticism": N_base,
                "openness": O_base
            }

            # Determine actual condition
            actual_condition = get_condition_code(condition)

            # Build patient info dictionary
            patient_info = {
                "name": name or "Anonymous",
                "age": age,
                "birth_year": birth_year,
                "sex": sex,
                "birthplace_city": birthplace_city.strip(),
                "birthplace_country": birthplace_country.strip(),
                "instruments": preferred_instruments,
                "favorite_genre": favorite_genre,
                "favorite_musician": favorite_musician.strip() if favorite_musician else "",
                "favorite_season": favorite_season,
                "natural_elements": natural_elements,
                "condition": actual_condition,
                "difficulty_sleeping": difficulty_sleeping,
                "trouble_remembering": trouble_remembering,
                "forgets_everyday_things": forgets_everyday_things,
                "difficulty_recalling_old_memories": difficulty_recalling_old_memories,
                "memory_worse_than_year_ago": memory_worse_than_year_ago,
                "visited_mental_health_professional": visited_mental_health,
                "big5_scores": big5_scores
            }

            # Generate patient ID
            patient_id = f"patient_{datetime.now().strftime('%Y%m%d%H%M%S')}"

            # Initialize TheraMuse
            with st.spinner(f" TheramuseRX is generating personalized recommendations for age {age}..."):
                try:
                    if 'theramuse' not in st.session_state:
                        st.session_state.theramuse = TheraMuse(db_path="/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app/theramuse.db")

                    theramuse = st.session_state.theramuse

                    # Get recommendations
                    recommendations = theramuse.get_therapy_recommendations(
                        patient_info, actual_condition, patient_id
                    )
                except Exception as db_error:
                    st.error(f"Database initialization error: {str(db_error)}")
                    st.error("Please try again or check database permissions.")
                    st.stop()

                # Save to our enhanced patient database
                try:
                    db_patient_id = save_patient_to_database(
                        patient_info, big5_scores, recommendations,
                        recommendations.get('session_id', f"session_{datetime.now().strftime('%Y%m%d%H%M%S')}")
                    )
                    st.success(f" 📝 Patient data saved to database! ID: {db_patient_id}")
                    st.info("💾 You can view all patient records in the '🗂️ Patient Database' section")
                except Exception as e:
                    st.warning(f" ⚠️ Could not save to enhanced database: {str(e)}")
                    st.info("Recommendations will still be displayed, but enhanced patient features won't be available.")
                    st.info("📝 Check the '🗂️ Patient Database' section to see if other patients are stored there.")

                # Store in session state
                st.session_state.tm_recs = recommendations
                st.session_state.tm_patient_data = patient_info
                st.session_state.tm_patient_id = patient_id
                st.session_state.tm_session_id = recommendations.get('session_id')
                st.session_state.tm_b5_scores = big5_scores
                st.session_state.show_results = True
                st.session_state.processing_complete = True

            st.success(f" Patient intake completed! Patient ID: {patient_id}")
            st.snow()
            st.rerun()  # Rerun to show results
    
    # Display results ONLY if processing is complete and show_results is True
    if st.session_state.show_results and st.session_state.processing_complete and 'tm_recs' in st.session_state:
        st.markdown("---")
        st.success(f" TheramuseRX's {st.session_state.tm_recs.get('total_songs', 0)}  recommendations!")

        # Start new intake button at top
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button(" Start New Patient Intake", type="secondary", width='stretch'):
                # Clear all session state
                for key in ['tm_recs', 'tm_patient_data', 'tm_patient_id', 'tm_session_id', 'tm_b5_scores', 'show_results', 'processing_complete']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

        tab1, tab2, tab3 = st.tabs([" Recommendations", " Personality Profile", " Bandit Statistics"])

        with tab1:
            render_recommendations_with_feedback(
                st.session_state.tm_recs,
                st.session_state.tm_patient_data,
                st.session_state.tm_session_id,
                st.session_state.tm_patient_id
            )

        with tab2:
            st.markdown("###  Big-5 Personality Profile")
            fig = create_personality_radar(st.session_state.tm_b5_scores)
            st.plotly_chart(fig, use_container_width=True)

            # Show generational context if dementia
            if st.session_state.tm_patient_data.get('condition') == 'dementia':
                birth_year = st.session_state.tm_patient_data.get('birth_year')
                if birth_year:
                    ragas = get_generational_ragas(birth_year)
                    if ragas:
                        st.markdown("""
                        <div class='glass-card'>
                            <h3>🎼 Therapeutic Ragas for Your Generation</h3>
                        </div>
                        """, unsafe_allow_html=True)
                        for raga in ragas:
                            st.markdown(f"<span class='category-badge'>{raga}</span>", unsafe_allow_html=True)

        with tab3:
            st.markdown("###  Thompson Sampling Statistics")
            bandit_stats = st.session_state.tm_recs.get('bandit_stats', {})

            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3>Total Interactions</h3>
                    <h2>{bandit_stats.get('n_interactions', 0)}</h2>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3>Average Reward</h3>
                    <h2>{bandit_stats.get('avg_reward', 0):.3f}</h2>
                </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                <div class='metric-card'>
                    <h3>Exploration Rate</h3>
                    <h2>{bandit_stats.get('exploration_rate', 0):.3f}</h2>
                </div>
                """, unsafe_allow_html=True)

            st.info(" The system uses Linear Thompson Sampling to learn from your feedback and improve recommendations over time.")

        # Add download options at the end
        render_download_options(
            st.session_state.tm_patient_data,
            st.session_state.tm_recs,
            st.session_state.tm_b5_scores
        )

def page_analytics():
    """Analytics dashboard page with premium design"""
    render_logo()
    st.markdown("""
    <div class='glass-card'>
        <h2> Analytics Dashboard</h2>
        <p style='color: rgba(255,255,255,0.7);'>Monitor system performance and user engagement</p>
    </div>
    """, unsafe_allow_html=True)
    
    if 'theramuse' not in st.session_state:
        st.session_state.theramuse = TheraMuse(db_path="/home/spectre-rosamund/Documents/ubuntu/thera/theramuse_app/theramuse.db")

    theramuse = st.session_state.theramuse
    analytics = theramuse.get_analytics()
    
    # Top metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f"""
        <div class='metric-card'>
            <h3>Total Sessions</h3>
            <h2>{analytics.get('total_sessions', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class='metric-card'>
            <h3>Total Feedback</h3>
            <h2>{analytics.get('total_feedback', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown(f"""
        <div class='metric-card'>
            <h3>Total Patients</h3>
            <h2>{analytics.get('total_patients', 0)}</h2>
        </div>
        """, unsafe_allow_html=True)
    
    # Rewards by condition
    st.markdown("###  Average Reward by Condition")
    rewards_data = analytics.get('rewards_by_condition', [])
    
    if rewards_data:
        df_rewards = pd.DataFrame(rewards_data)
        fig = px.bar(df_rewards, x='condition', y='avg_reward', 
                    title='Average Reward by Medical Condition',
                    color='count', color_continuous_scale='Viridis')
        fig.update_layout(
            paper_bgcolor='rgba(0, 0, 0, 0)',
            plot_bgcolor='rgba(0, 0, 0, 0)',
            font=dict(color='white', family='Inter')
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No feedback data available yet. Start collecting feedback to see analytics!")
    
    # YouTube API Health
    st.markdown("###  System Health")
    health = theramuse.check_api_health()
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.json(health.get('youtube_api', {}))
        st.markdown("</div>", unsafe_allow_html=True)
    with col2:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.json(health.get('database', {}))
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Cache management
    st.markdown("###  Cache Management")
    cache_status = theramuse.get_youtube_cache_status()
    st.metric("YouTube Cache Size", cache_status.get('cache_size', 0))
    
    if st.button("Clear YouTube Cache", type="secondary"):
        theramuse.clear_youtube_cache()
        st.success(" YouTube cache cleared!")
        st.snow()

def page_research_evidence():
    """Research Evidence page with comprehensive scientific literature"""
    render_logo()

    st.markdown("""
    <div class='glass-card'>
        <h2>Research Evidence</h2>
        <p style='color: rgba(255,255,255,0.7); font-size: 1.05rem;'>
            Comprehensive scientific evidence supporting music therapy for various conditions and personality-based music preferences
        </p>
    </div>
    """, unsafe_allow_html=True)

    # # Back button
    if st.button("← Back to Previous Page", type="secondary"):
        st.session_state.page = " Patient Intake"
        st.rerun()

    st.markdown("## Research Evidence of Music Therapy for Dementia, ADHD, Down Syndrome, and Big Five Personality Traits & Music Genre Preferences")

    st.markdown("### Music Therapy for Dementia (Older People)")
    st.markdown("#### Systematic Reviews and Meta-Analyses:")
    st.write("Moreno-Morales et al. (2020) conducted a comprehensive systematic review and meta-analysis examining music therapy's effects on dementia patients. The study analyzed eight studies with 816 participants and found that music therapy significantly improved cognitive function (SMD = -0.23, 95% CI: -0.44, -0.02), with passive interventions (listening to music) showing greater effects than active interventions. Shorter intervention periods (<20 weeks) were more effective than longer ones.")

    st.write("Recent NHS Research (2025) published findings from the MELODIC (Music therapy Embedded in the Life Of Dementia Inpatient Care) protocol, a feasibility study involving 28 patients, 13 family members, and 48 staff members across two NHS mental health dementia wards. The intervention showed high treatment adherence with no increase in distress symptoms or safety incidents during the intervention period.")

    st.write("Evidence from systematic review (2024) demonstrated that music therapy reduces distress and improves well-being for people with advanced dementia in institutional settings by meeting unmet needs and increasing communication between staff and family members.")

    st.write("Meta-analysis on quality of life (2024) examining eight studies with 605 subjects (330 receiving music therapy) showed positive effects on depression scores, though with wide confidence intervals, emphasizing the need for more robust research.")

    st.markdown("### Music Therapy for ADHD (Any Age)")
    st.markdown("#### Meta-Analyses and Systematic Reviews:")
    st.write("de Oliveira Goes et al. (2025) conducted a meta-analysis examining music therapy outcomes in children and adolescents with ADHD. The study showed a trend toward efficacy (effect size: 1.18; CI: -3.8 - 0.21; p = 0.08) with significant heterogeneity among trials (I² = 92%). Active music therapy improved social skills, self-esteem, and reduced aggressive behavior.")

    st.write("Kasuya-Ueba et al. (2024) published a systematic review examining music's effects on ADHD, finding that both active (playing instruments) and passive (listening) music therapy reduced symptomatology, with rock music specifically reducing motor activity and improving attention.")

    st.write("Quispe et al. (2025) systematic review of 24 documents found music as effective non-pharmacological therapy improving concentration, self-esteem, social skills development, and reducing hyperactivity and anxiety states.")

    st.write("Multi-dimensional computational study (2025) analyzed over 9,215 tracks from r/ADHD community, finding that focus music was characterized by higher valence and instrumentalness, suggesting preference for uplifting and instrumental tracks that aid concentration.")

    st.write("Neurocognitive mechanisms review (2025) identified seven potential mechanisms through which music interventions help ADHD: executive function enhancement, timing improvement, arousal regulation, default mode network modulation, neural entrainment, affective management, and social bonding facilitation.")

    st.markdown("### Music Therapy for Children with Down Syndrome")
    st.markdown("#### Research Studies:")
    st.write("He et al. (2024) conducted a quasi-experimental study with 18 children with Down syndrome (ages 7-10) using Orff-Schulwerk-based music engagement and movement activities. Results showed statistically significant improvements in attentiveness (t = 9.0+) and memory compared to control groups.")

    st.write("Zhang (2024) literature review and interviews found that active music therapy promotes socialization development in children under 12 with Down syndrome in three aspects: language skills, social-emotional development, and prosocial behavior.")

    st.write("Music therapy research (2020) demonstrated that music plays a crucial role in development of children with Down syndrome, improving social skills, communication, and emotional expression.")

    st.write("Systematic review of music-based interventions (2025) in pediatric populations found music therapy valuable as adjunct to conventional neurorehabilitation, reducing anxiety, pain, and depressive symptoms while improving well-being and caregiver interactions.")

    st.markdown("### Big Five Personality Traits and Music Genre Preferences")
    st.markdown("#### Research Findings:")
    st.write("Boccia (2020) study of 175 participants found that individuals high in Openness, Extraversion, or Agreeableness enjoyed more music genres. Those low in Neuroticism also enjoyed more genres. High Conscientiousness and Neuroticism were associated with preference for mellow music (classical, new age).")

    st.write("Ferwerda et al. (2017) analyzed 1,415 Last.fm users across 83 countries, finding:")
    st.write("Openness: Correlated with new age (r=.101), classical (r=.136), blues (r=.120), country (r=.106), world (r=.134), folk (r=.214), and jazz (r=.139)")
    st.write("Extraversion: Positively correlated with R&B (r=.103) and rap (r=.129)")
    st.write("Agreeableness: Correlated with country (r=.104) and folk (r=.104)")
    st.write("Neuroticism: Positively correlated with alternative music (r=.101)")

    st.write("Nave et al. (2018) study of 8,097 participants found:")
    st.write("Openness: Associated with sophisticated music (r=.16) and less liking of mellow (r=-.12) and contemporary music (r=-.11)")
    st.write("Extraversion: Associated with unpretentious music preference (r=.13)")

    st.write("Cambridge University Research (2022) found consistent global correlations between:")
    st.write("Extraversion and contemporary music")
    st.write("Conscientiousness and unpretentious music")
    st.write("Agreeableness and mellow and unpretentious music")
    st.write("Openness and mellow, contemporary, intense, and sophisticated music")

    st.write("Spotify Research (2025) analyzing three months of listening data found:")
    st.write("Openness: Linked to Classical, Afropop, and \"Sentimental\" music")
    st.write("Emotional Stability: Inversely related to Blues and \"Brooding\" music, positively to Soul and \"Lively\" music")
    st.write("Agreeableness: Inversely related to Death Metal and \"Aggressive\" music, positively to Jazz and Country")

    st.markdown("### Additional Research Papers (Continuing toward 50+ papers)")
    st.write("Rentfrow & Gosling (2011) - Five-factor model of musical preferences based on affective reactions")
    st.write("Setti (2022) - Influence of openness facets on music preference in 478 undergraduate students")
    st.write("Li (2024) - Cross-cultural study on personality influence on music preference across geographic locations")
    st.write("Hansen et al. (2025) - Role of Big Five personality domains in musical preferences using Nature publication")
    st.write("Mendoza Garay et al. (2023) - Exploring relations between Big Five traits and embodied musical emotions")
    st.write("Vella (2017) - Influence of openness and extraversion on music preference mediation")
    st.write("Systematic review (2025) - Music intervention for neurodevelopment in pediatric populations")
    st.write("Meta-analysis (2025) - Effectiveness of music-based therapy on adolescents and children")
    st.write("Research (2022) - Effectiveness of music therapy in children with autism spectrum disorder")
    st.write("Study (2025) - Music therapy for pain and anxiety reduction in pediatric settings")
    st.write("Clinical trial (2012) - Randomized controlled trial of improvisational music therapy for autism spectrum disorders")
    st.write("Research (2024) - Use and effectiveness of musical social story therapy in developmental disorders")

    st.markdown("""
    <div style='background: linear-gradient(135deg, rgba(163, 177, 138, 0.2) 0%, rgba(88, 129, 87, 0.2) 100%);
                padding: 1.5rem;
                border-radius: 12px;
                border-left: 4px solid #A3B18A;
                margin: 1.5rem 0;'>
        <h3 style='color: #E8DDB5; margin-top: 0; margin-bottom: 1rem; font-weight: 600;'>Research Summary</h3>
        <p style='color: #FFFFFF; line-height: 1.6; margin: 0;'>
            The research evidence demonstrates <strong style='color: #A3B18A;'>consistent beneficial effects</strong> of music therapy across all three populations studied.
            For dementia patients, music therapy <strong style='color: #588157;'>improves cognitive function, reduces distress, and enhances quality of life</strong>.
            For individuals with ADHD, music therapy <strong style='color: #E8DDB5;'>shows promise in improving attention, reducing hyperactivity, and enhancing social skills</strong>.
            For children with Down syndrome, music therapy <strong style='color: #A3B18A;'>significantly improves socialization, communication, and cognitive development</strong>.
            Regarding personality and music preferences, research consistently shows that the Big Five traits predict musical preferences, with <strong style='color: #588157;'>Openness and Extraversion</strong> showing the strongest correlations with diverse musical tastes, while <strong style='color: #E8DDB5;'>Conscientiousness and Neuroticism</strong> are associated with more specific genre preferences.
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='color: rgba(255,255,255,0.6); font-size: 0.85rem; margin: 1rem 0; font-style: italic;'>
        This evidence base, comprising systematic reviews, meta-analyses, randomized controlled trials, and large-scale observational studies,
        provides strong support for the therapeutic applications of music across these populations and confirms the relationship between
        personality traits and musical preferences.
    </div>
    """, unsafe_allow_html=True)

    # Research Links Section
    st.markdown("""
    <div style='margin: 2rem 0;'>
        <h3 style='color: #E8DDB5; margin-bottom: 1rem; font-size: 1.5rem;'>🔗 Research Paper Links</h3>
        <p style='color: #A3B18A; font-size: 1.1rem; margin-bottom: 1rem; font-weight: 500;'>
            Access to full research papers and additional resources:
        </p>
    </div>
    """, unsafe_allow_html=True)

    with st.expander(" View Research Links", expanded=False):
        research_links = [
            "https://www.frontiersin.org/journals/medicine/articles/10.3389/fmed.2020.00160/full",
            "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1618324/full",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12307461/",
            "https://www.nature.com/articles/s44220-024-00342-x",
            "https://jpalliativecare.com/effect-of-music-therapy-on-quality-of-life-in-geriatric-population-a-systematic-review-and-meta-analysis/",
            "https://pubmed.ncbi.nlm.nih.gov/40680190/",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC10221503/",
            "https://journals.sapienzaeditorial.com/index.php/SIJIS/article/view/e25044",
            "https://journals.plos.org/plosone/article?id=10.1371%2Fjournal.pone.0324369",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12316199/",
            "https://archive.conscientiabeam.com/index.php/61/article/download/3626/7913",
            "https://www.scholarlyreview.org/api/v1/articles/121698-music-as-a-language-assessing-the-extent-to-which-active-music-therapy-promotes-socialization-development-for-children-under-12-with-down-syndrome.pdf",
            "https://journals.sagepub.com/doi/abs/10.1080/03080188.2020.1755556",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12191269/",
            "https://digitalcommons.lindenwood.edu/cgi/viewcontent.cgi?article=1000&context=psych_journals",
            "https://www.cp.jku.at/people/schedl/Research/Publications/pdf/ferwerda_umap_2017.pdf",
            "https://www.davidmgreenberg.com/wp-content/uploads/2018/11/Nave-et-al-2018-music-preferences-from-fb-likes.pdf",
            "https://www.cam.ac.uk/stories/musical-preferences-unite-personalities-worldwide",
            "https://research.atspotify.com/just-the-way-you-are-music-listening-and-personality",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC3138530/",
            "https://ir.library.illinoisstate.edu/cgi/viewcontent.cgi?article=2574&context=etd",
            "https://www.pioneerpublisher.com/jrssh/article/download/1045/948/1097",
            "https://www.nature.com/articles/s41598-025-95661-z",
            "https://journals.sagepub.com/doi/10.1177/03057356221135355",
            "https://journals.sagepub.com/doi/pdf/10.1177/0305735616658957",
            "https://www.nature.com/articles/s41598-025-93795-8",
            "https://www.sciencedirect.com/science/article/abs/pii/S0190740925001343",
            "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2022.905113/full",
            "https://journals.sagepub.com/doi/10.1177/23320249241265240",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC3280156/",
            "https://www.auctoresonline.org/article/use-and-effectiveness-of-musical-social-story-therapy-in-children-with-developmental-disorders-down-syndrome-autism-spectrum-disorder-fragile-x-syndrome-fetal-alcohol-spectrum-disorder-cerebral-palsy-and-adhd",
            "https://baylor-ir.tdl.org/bitstreams/81c66bd2-202f-4051-b3b5-3aaee0e48fa9/download",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC12235852/",
            "https://files.eric.ed.gov/fulltext/EJ976663.pdf",
            "https://www.sciencedirect.com/science/article/pii/S002239562200231X",
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC6481398/",
            "https://www.pagepress.org/medicine/gimle/article/view/595"
        ]

        # Display all research links as black colored text
        st.markdown("""
        <div style='color: #000000; line-height: 1.8;'>
        """ + "<br>".join([f"• <a href='{link}' target='_blank' style='color: #000000; text-decoration: underline;'>{link}</a>" for link in research_links]) + """
        </div>
        """, unsafe_allow_html=True)

def page_about():
    """About page with premium design"""
    render_logo()

    st.markdown("## About TheramuseRX")

    st.markdown("""
    **TheramuseRX is an our own AI powered model music therapy recommendation system**
    designed to provide personalized and therapeutically relevant music selections from YouTube with
    **automatic reinforcement learning**.
    """)

    st.markdown("### Key Features")

    features = [
        "**Condition-Specific Therapy:** Dedicated modules for Dementia, Down Syndrome, and ADHD",
        "**Big-5 Personality:** Personality-based music genre recommendations",
        "**Reinforcement Learnin:** Reinforcement learning algorithm for personalized learning",

    ]

    for feature in features:
        st.markdown(f"• {feature}")

    st.markdown("### Therapy Modules")

    with st.container():
        st.markdown("#### Developed for:")
        dementia_features = [
            "**Working on:** Dementia, Down Sysndrome,ADHD",
            "**Memory Enhancement:** Music to stimulate memory recall and cognitive function",
            "**Sleep Improvement:** Calming tracks to aid sleep quality",
            "**Emotional Well-being:** Uplifting music to enhance mood and reduce anxiety",
            "**Personalized Selection:** Incorporates patient's birthplace, age, and preferences"

        ]
        for feature in dementia_features:
            st.markdown(f"  - {feature}")



    # st.markdown("### Reinforcement Learning")

    # rl_features = [
    #     "**Linear Thompson Sampling:** Bayesian approach for exploration-exploitation",
    #     "**Contextual Features:** 20-dimensional feature space",
    #     "**Continuous Learning:** Updates with every feedback submission",
    #     "**Non-stationary:** Decay factor for adapting to changing preferences"
    # ]

    # for feature in rl_features:
    #     st.markdown(f"• {feature}")
    st.markdown("### Developer")
    st.markdown("**Moodifai-Ashiq Sazid**")
 
# MAIN APP

def main():
    """Main application"""
    
    # Sidebar navigation
    # Sidebar navigation
    # Sidebar navigation with logo
    st.sidebar.image("b.png", width=200)
    st.sidebar.markdown("<h1 style='text-align: center; margin-bottom: 2rem; color: #E8DDB5;'></h1>", unsafe_allow_html=True)

    pages = {
    " Patient Intake": page_intake,
    "Patient Database": page_patient_database,
    " Analytics": page_analytics,
    "Research Evidence": page_research_evidence,
    "About": page_about
}

    # Always show the sidebar navigation
    selection = st.sidebar.radio("Navigation", list(pages.keys()))

    # Check for session state navigation override
    if 'page' in st.session_state and st.session_state.page in pages:
        selection = st.session_state.page
        # Clear the session state after using it
        del st.session_state.page

    # Display selected page
    pages[selection]()
    
    # Footer
    st.sidebar.markdown("---")
    st.sidebar.markdown("""
    <div style='text-align: center; padding: 1rem;'>
        <p style='font-weight: 700; font-size: 1.1rem; color: #E8DDB5; margin-bottom: 0.5rem;'>“Where words fail, music speaks — not to the mind, but to the soul.”</p>
        <p style='color: rgba(255,255,255,0.6); font-size: 0.9rem;'>— Hans Christian Andersen </p>
        <p style='color: rgba(255,255,255,0.4); font-size: 0.8rem; margin-top: 0.5rem;'></p>
    </div>
    """, unsafe_allow_html=True)
    
    # System status in sidebar

if __name__ == "__main__":
    main()
